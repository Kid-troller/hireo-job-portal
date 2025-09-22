from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.forms import PasswordChangeForm
from django.contrib import messages
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods, require_POST
from django.core.paginator import Paginator
from django.db.models import Q, Count, Avg
from django.utils import timezone
from django.urls import reverse
from django.core.mail import send_mail
from django.conf import settings
from django.template.loader import render_to_string
from django.utils.html import strip_tags
from django.contrib.auth.models import User
from django.db import transaction
from django.core.exceptions import ValidationError
from django.utils.crypto import get_random_string
import json
import uuid
import secrets
import string
from datetime import datetime, timedelta
import re
from collections import defaultdict
import io
import os
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from applications.models import Application, ApplicationStatus, Interview, Message, Notification
from applications.notification_utils import NotificationManager
from .forms import JobSeekerProfileForm, EmployerProfileForm, UserProfileForm, CustomAuthenticationForm, UserRegistrationForm, ForgotPasswordRequestForm, SecurityQuestionsForm, NewPasswordForm, SecurityQuestionsSetupForm
from .decorators import jobseeker_required, employer_required
from .ats_engine import ATSOptimizationEngine
from .resume_ai import ResumeAI
from jobs.models import JobPost, JobCategory, JobLocation, SavedJob, JobAlert
from employers.models import Company, EmployerProfile
from .models import *
from django.db import models
import secrets
import json
from datetime import datetime, timedelta
import re
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage

def home(request):
    """Home page view"""
    featured_jobs = JobPost.objects.filter(
        is_featured=True, 
        status='active'
    ).order_by('-created_at')[:6]
    
    recent_jobs = JobPost.objects.filter(
        status='active'
    ).order_by('-created_at')[:8]
    
    job_categories = JobPost.objects.values('category__name').distinct()[:8]
    
    context = {
        'featured_jobs': featured_jobs,
        'recent_jobs': recent_jobs,
        'job_categories': job_categories,
    }
    return render(request, 'base/home.html', context)

def register(request):
    """User registration view"""
    if request.user.is_authenticated:
        # Redirect authenticated users to their appropriate dashboard
        if hasattr(request.user, 'userprofile'):
            if request.user.userprofile.user_type == 'employer':
                return redirect('employer:dashboard')
            else:
                return redirect('accounts:dashboard')
        return redirect('accounts:dashboard')
    
    if request.method == 'POST':
        form = UserRegistrationForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)
            user_type = form.cleaned_data['user_type']
            
            # Keep user active for immediate login
            user.is_active = True
            
            with transaction.atomic():
                user.save()
                
                # Create user profile
                user_profile = user.userprofile
                user_profile.user_type = user_type
                user_profile.is_verified = True  # Skip email verification for auto-login
                user_profile.save()
                
                # Create specific profile based on user type
                if user_type == 'jobseeker':
                    JobSeekerProfile.objects.create(user_profile=user_profile)
                elif user_type == 'employer':
                    # Employer will complete company profile later
                    pass
                
                # Automatically authenticate and login the user
                authenticated_user = authenticate(username=user.username, password=form.cleaned_data['password1'])
                if authenticated_user:
                    login(request, authenticated_user)
                    
                    # Redirect based on user type
                    if user_type == 'employer':
                        messages.success(request, 'Welcome! Your employer account has been created successfully.')
                        return redirect('employer:dashboard')
                    else:
                        messages.success(request, 'Welcome! Your account has been created successfully.')
                        return redirect('accounts:dashboard')
                else:
                    messages.error(request, 'Account created but login failed. Please try logging in manually.')
                    return redirect('accounts:login')
        else:
            # Add error messages for form validation
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f'{field.title()}: {error}')
    else:
        form = UserRegistrationForm()
    
    return render(request, 'accounts/register.html', {'form': form})

def send_verification_email(request, user, token):
    """Send email verification link to the user"""
    current_site = request.get_host()
    verify_url = f"http://{current_site}{reverse('accounts:verify_email', kwargs={'token': token.token})}"
    
    # Create email content
    subject = 'Verify your Hireo account'
    html_message = render_to_string('accounts/email/verification_email.html', {
        'user': user,
        'verify_url': verify_url,
        'expiry_days': 2,  # Token expires after 2 days
    })
    plain_message = strip_tags(html_message)
    
    # Send email
    send_mail(
        subject,
        plain_message,
        settings.EMAIL_HOST_USER,
        [user.email],
        html_message=html_message,
        fail_silently=False,
    )

def verify_email(request, token):
    """Verify user email with token"""
    from .models import EmailVerificationToken
    
    try:
        # Find the token
        verification_token = EmailVerificationToken.objects.get(token=token)
        
        # Check if token is valid
        if verification_token.is_valid():
            user = verification_token.user
            
            # Activate user account
            user.is_active = True
            user.save()
            
            # Mark token as used
            verification_token.is_used = True
            verification_token.save()
            
            # Update user profile
            user.userprofile.is_verified = True
            user.userprofile.save()
            
            messages.success(request, 'Your email has been verified! You can now log in to your account.')
            return redirect('accounts:login')
        else:
            messages.error(request, 'The verification link has expired or already been used.')
            return redirect('accounts:resend_verification')
    except EmailVerificationToken.DoesNotExist:
        messages.error(request, 'Invalid verification link.')
        return redirect('accounts:login')

def resend_verification(request):
    """Resend verification email"""
    if request.method == 'POST':
        email = request.POST.get('email')
        try:
            user = User.objects.get(email=email, is_active=False)
            from .models import EmailVerificationToken
            token = EmailVerificationToken.generate_token(user)
            send_verification_email(request, user, token)
            messages.success(request, 'Verification email has been resent. Please check your inbox.')
            return redirect('accounts:login')
        except User.DoesNotExist:
            messages.error(request, 'No pending verification found for this email address.')
    
    return render(request, 'accounts/resend_verification.html')

def user_login(request):
    """Optimized user login view with performance improvements"""
    if request.user.is_authenticated:
        # Fast redirect for already authenticated users
        try:
            user_type = request.user.userprofile.user_type
            if user_type == 'employer':
                return redirect('employer:dashboard')
            else:
                return redirect('accounts:dashboard')
        except:
            return redirect('accounts:dashboard')
    
    if request.method == 'POST':
        form = CustomAuthenticationForm(request, data=request.POST)
        if form.is_valid():
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            user = authenticate(username=username, password=password)
            
            if user is not None:
                # Quick active check
                if not user.is_active:
                    messages.error(request, 'Your account is not active. Please verify your email first.')
                    return redirect('accounts:resend_verification')
                
                # Login user immediately
                login(request, user)
                
                # Handle remember me functionality
                if not request.POST.get('remember_me'):
                    request.session.set_expiry(0)
                
                # Get next URL for fast redirect
                next_url = request.GET.get('next')
                if next_url:
                    return redirect(next_url)
                
                # Optimized role-based redirect with single query
                try:
                    # Use select_related to get userprofile in single query
                    from django.contrib.auth.models import User
                    user_with_profile = User.objects.select_related('userprofile').get(id=user.id)
                    user_type = user_with_profile.userprofile.user_type
                    
                    if user_type == 'employer':
                        messages.success(request, f'Welcome back, {user.get_full_name()}!')
                        return redirect('employer:dashboard')
                    else:
                        # Skip profile completion check for faster login
                        # Profile completion will be checked on dashboard if needed
                        messages.success(request, f'Welcome back, {user.get_full_name()}!')
                        return redirect('accounts:dashboard')
                        
                except Exception:
                    # Fallback - create profile if needed
                    UserProfile.objects.get_or_create(user=user)
                    messages.success(request, f'Welcome back, {user.get_full_name()}!')
                    return redirect('accounts:dashboard')
            else:
                messages.error(request, 'Invalid username or password. Please try again.')
        else:
            # Simplified error handling
            messages.error(request, 'Please check your username and password.')
    else:
        form = CustomAuthenticationForm()
    
    return render(request, 'accounts/login.html', {'form': form})

@login_required
def user_logout(request):
    """User logout view with enhanced security"""
    # Get user info before logout for personalized message
    user_name = request.user.get_full_name() or request.user.username
    
    # Clear all session data
    request.session.flush()
    
    # Logout user
    logout(request)
    
    # Add success message
    messages.success(request, f'Goodbye {user_name}! You have been logged out successfully. We hope to see you again soon.')
    
    # Redirect to login page or home
    next_url = request.GET.get('next', 'accounts:login')
    if next_url == 'home':
        return redirect('home')
    return redirect('accounts:login')

@login_required
def dashboard(request):
    """User dashboard view"""
    user_profile = request.user.userprofile
    
    if user_profile.user_type == 'employer':
        return redirect('employer:dashboard')
    
    # Job seeker dashboard - create profile if it doesn't exist
    try:
        job_seeker_profile = JobSeekerProfile.objects.get(user_profile=user_profile)
    except JobSeekerProfile.DoesNotExist:
        job_seeker_profile = JobSeekerProfile.objects.create(user_profile=user_profile)
        messages.info(request, 'Your job seeker profile has been created. Please complete your profile for better job recommendations.')
    
    # Get user's applications
    applications = Application.objects.filter(applicant=job_seeker_profile).order_by('-applied_at')[:5]
    
    # Get notifications for job seeker
    notifications = NotificationManager.get_user_notifications(request.user, limit=10)
    unread_notifications = NotificationManager.get_user_notifications(request.user, limit=5, unread_only=True)
    notification_counts = NotificationManager.get_notification_counts(request.user)
    
    # Get job alerts
    job_alerts = JobAlert.objects.filter(user=job_seeker_profile, is_active=True)
    
    # Get recommended jobs based on skills
    user_skills = job_seeker_profile.skills.lower().split(',') if job_seeker_profile.skills else []
    recommended_jobs = []
    
    if user_skills:
        recommended_jobs = JobPost.objects.filter(
            Q(status='active') &
            (Q(required_skills__icontains=user_skills[0]) | Q(preferred_skills__icontains=user_skills[0]))
        ).order_by('-created_at')[:3]
    
    context = {
        'job_seeker_profile': job_seeker_profile,
        'applications': applications,
        'job_alerts': job_alerts,
        'recommended_jobs': recommended_jobs,
        'notifications': notifications,
        'unread_notifications': unread_notifications,
        'notification_counts': notification_counts,
    }
    
    return render(request, 'accounts/dashboard.html', context)

@login_required
def complete_profile(request):
    """Complete job seeker profile view"""
    user_profile = request.user.userprofile
    
    if user_profile.user_type != 'jobseeker':
        return redirect('accounts:dashboard')
    
    try:
        job_seeker_profile = JobSeekerProfile.objects.create(user_profile=user_profile)
    except:
        job_seeker_profile = JobSeekerProfile.objects.get(user_profile=user_profile)
    
    if request.method == 'POST':
        profile_form = JobSeekerProfileForm(request.POST, request.FILES, instance=job_seeker_profile)
        user_form = UserProfileForm(request.POST, request.FILES, instance=user_profile)
        
        if profile_form.is_valid() and user_form.is_valid():
            profile_form.save()
            user_form.save()
            messages.success(request, 'Profile updated successfully!')
            return redirect('accounts:dashboard')
    else:
        profile_form = JobSeekerProfileForm(instance=job_seeker_profile)
        user_form = UserProfileForm(instance=user_profile)
    
    context = {
        'profile_form': profile_form,
        'user_form': user_form,
        'job_seeker_profile': job_seeker_profile,
    }
    
    return render(request, 'accounts/complete_profile.html', context)

@login_required
def profile_detail(request):
    """User profile detail view"""
    user_profile = request.user.userprofile
    
    if user_profile.user_type == 'employer':
        return redirect('employer:profile')
    
    job_seeker_profile = get_object_or_404(JobSeekerProfile, user_profile=user_profile)
    
    # Get related data
    educations = Education.objects.filter(job_seeker=job_seeker_profile).order_by('-start_date')
    experiences = Experience.objects.filter(job_seeker=job_seeker_profile).order_by('-start_date')
    skills = Skill.objects.filter(job_seeker=job_seeker_profile).order_by('-years_of_experience')
    certifications = Certification.objects.filter(job_seeker=job_seeker_profile).order_by('-issue_date')
    
    context = {
        'user_profile': user_profile,
        'job_seeker_profile': job_seeker_profile,
        'educations': educations,
        'experiences': experiences,
        'skills': skills,
        'certifications': certifications,
    }
    
    return render(request, 'accounts/profile_detail.html', context)

@login_required
def edit_profile(request):
    """Edit profile view"""
    user_profile = request.user.userprofile
    
    if user_profile.user_type == 'employer':
        return redirect('employer:edit_profile')
    
    job_seeker_profile = get_object_or_404(JobSeekerProfile, user_profile=user_profile)
    
    if request.method == 'POST':
        profile_form = JobSeekerProfileForm(request.POST, request.FILES, instance=job_seeker_profile)
        user_form = UserProfileForm(request.POST, request.FILES, instance=user_profile)
        
        if profile_form.is_valid() and user_form.is_valid():
            profile_form.save()
            user_form.save()
            messages.success(request, 'Profile updated successfully!')
            return redirect('accounts:profile_detail')
    else:
        profile_form = JobSeekerProfileForm(instance=job_seeker_profile)
        user_form = UserProfileForm(instance=user_profile)
    
    context = {
        'user': request.user,
        'user_profile': request.user.userprofile,
        'profile_form': profile_form,
        'user_form': user_form,
        'job_seeker_profile': job_seeker_profile,
    }
    return render(request, 'accounts/profile_edit.html', context)

@login_required
def complete_profile(request):
    """Complete job seeker profile view"""
    user_profile = request.user.userprofile
    
    if user_profile.user_type != 'jobseeker':
        return redirect('accounts:dashboard')
    
    try:
        job_seeker_profile = JobSeekerProfile.objects.create(user_profile=user_profile)
    except:
        job_seeker_profile = JobSeekerProfile.objects.get(user_profile=user_profile)
    
    if request.method == 'POST':
        # Handle profile completion
        user = request.user
        user_profile = user.userprofile
        
        # Update basic info
        user.first_name = request.POST.get('first_name', '')
        user.last_name = request.POST.get('last_name', '')
        user.save()
        
        # Update profile info
        user_profile.phone = request.POST.get('phone', '')
        user_profile.location = request.POST.get('location', '')
        user_profile.bio = request.POST.get('bio', '')
        user_profile.save()
        
        messages.success(request, 'Profile completed successfully!')
        return redirect('accounts:dashboard')
    
    context = {
        'user': request.user,
        'user_profile': request.user.userprofile,
    }
    return render(request, 'accounts/complete_profile.html', context)

@login_required
def add_education(request):
    """Add education view"""
    user_profile = request.user.userprofile
    
    if user_profile.user_type != 'jobseeker':
        return redirect('dashboard')
    
    job_seeker_profile = get_object_or_404(JobSeekerProfile, user_profile=user_profile)
    
    if request.method == 'POST':
        form = EducationForm(request.POST)
        if form.is_valid():
            education = form.save(commit=False)
            education.job_seeker = job_seeker_profile
            education.save()
            messages.success(request, 'Education added successfully!')
            return redirect('accounts:profile_detail')
    else:
        form = EducationForm()
    
    return render(request, 'accounts/add_education.html', {'form': form})

@login_required
def add_experience(request):
    """Add experience view"""
    user_profile = request.user.userprofile
    
    if user_profile.user_type != 'jobseeker':
        return redirect('dashboard')
    
    job_seeker_profile = get_object_or_404(JobSeekerProfile, user_profile=user_profile)
    
    if request.method == 'POST':
        form = ExperienceForm(request.POST)
        if form.is_valid():
            experience = form.save(commit=False)
            experience.job_seeker = job_seeker_profile
            experience.save()
            messages.success(request, 'Experience added successfully!')
            return redirect('accounts:profile_detail')
    else:
        form = ExperienceForm()
    
    return render(request, 'accounts/add_experience.html', {'form': form})

@login_required
def add_skill(request):
    """Add skill view"""
    user_profile = request.user.userprofile
    
    if user_profile.user_type != 'jobseeker':
        return redirect('dashboard')
    
    job_seeker_profile = get_object_or_404(JobSeekerProfile, user_profile=user_profile)
    
    if request.method == 'POST':
        form = SkillForm(request.POST)
        if form.is_valid():
            skill = form.save(commit=False)
            skill.job_seeker = job_seeker_profile
            skill.save()
            messages.success(request, 'Skill added successfully!')
            return redirect('accounts:profile_detail')
    else:
        form = SkillForm()
    
    return render(request, 'accounts/add_skill.html', {'form': form})

@login_required
def add_certification(request):
    """Add certification view"""
    user_profile = request.user.userprofile
    
    if user_profile.user_type != 'jobseeker':
        return redirect('dashboard')
    
    job_seeker_profile = get_object_or_404(JobSeekerProfile, user_profile=user_profile)
    
    if request.method == 'POST':
        form = CertificationForm(request.POST)
        if form.is_valid():
            certification = form.save(commit=False)
            certification.job_seeker = job_seeker_profile
            certification.save()
            messages.success(request, 'Certification added successfully!')
            return redirect('accounts:profile_detail')
    else:
        form = CertificationForm()
    
    return render(request, 'accounts/add_certification.html', {'form': form})

@login_required
def change_password(request):
    """Change password view"""
    if request.method == 'POST':
        form = PasswordChangeForm(user=request.user, data=request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Password changed successfully! Please log in again.')
            return redirect('accounts:login')
    else:
        form = PasswordChangeForm(user=request.user)
    
    return render(request, 'accounts/change_password.html', {'form': form})

@login_required
def applications_list(request):
    """Enhanced job seeker applications list with advanced filtering and analytics - SQLite3 version"""
    from datetime import datetime, timedelta
    from django.utils import timezone
    from hireo.db_utils import db
    
    # Get user profile using raw SQL
    user_data = db.get_user_by_id(request.user.id)
    if not user_data or user_data['user_type'] != 'jobseeker':
        return redirect('dashboard')
    
    # Get or create JobSeekerProfile using raw SQL
    job_seeker_profile = db.get_jobseeker_profile(request.user.id)
    if not job_seeker_profile:
        # Create job seeker profile
        user_profile_id = user_data['id']  # This should be the profile ID, need to get it properly
        
        # Get user profile ID
        profile_query = "SELECT id FROM accounts_userprofile WHERE user_id = ?"
        profile_result = db.execute_single(profile_query, (request.user.id,))
        
        if profile_result:
            db.create_jobseeker_profile(
                profile_result['id'], 
                request.user.first_name or 'User', 
                request.user.last_name or 'Name'
            )
            job_seeker_profile = db.get_jobseeker_profile(request.user.id)
            messages.info(request, 'Your job seeker profile has been created. Please complete your profile for better job recommendations.')
    
    # Get applications with filtering using raw SQL
    # First try to get JobSeekerProfile ID, if not available, use user ID
    applicant_id = job_seeker_profile['id'] if job_seeker_profile else request.user.id
    
    # Build dynamic SQL query with filters
    base_query = """
    SELECT a.*, j.title as job_title, c.name as company_name, 
           j.min_salary, j.max_salary, j.employment_type,
           jc.name as category_name
    FROM applications_application a
    JOIN jobs_jobpost j ON a.job_id = j.id
    JOIN employers_company c ON j.company_id = c.id
    LEFT JOIN jobs_jobcategory jc ON j.category_id = jc.id
    WHERE a.applicant_id = ?
    """
    
    params = [applicant_id]
    conditions = []
    
    # Enhanced filtering
    status_filter = request.GET.get('status')
    if status_filter:
        conditions.append("a.status = ?")
        params.append(status_filter)
    
    # Date filtering
    date_filter = request.GET.get('date_filter')
    if date_filter:
        if date_filter == 'today':
            conditions.append("DATE(a.applied_at) = DATE('now')")
        elif date_filter == 'week':
            conditions.append("a.applied_at >= datetime('now', '-7 days')")
        elif date_filter == 'month':
            conditions.append("a.applied_at >= datetime('now', '-30 days')")
        elif date_filter == 'quarter':
            conditions.append("a.applied_at >= datetime('now', '-90 days')")
    
    # Job type filtering
    job_type_filter = request.GET.get('job_type')
    if job_type_filter:
        conditions.append("j.employment_type = ?")
        params.append(job_type_filter)
    
    # Company search
    company_search = request.GET.get('company')
    if company_search:
        conditions.append("c.name LIKE ?")
        params.append(f"%{company_search}%")
    
    # Job title search
    job_title_search = request.GET.get('job_title')
    if job_title_search:
        conditions.append("j.title LIKE ?")
        params.append(f"%{job_title_search}%")
    
    # Legacy search parameter support
    search = request.GET.get('search')
    if search:
        conditions.append("(j.title LIKE ? OR c.name LIKE ?)")
        params.extend([f"%{search}%", f"%{search}%"])
    
    # Add conditions to query
    if conditions:
        base_query += " AND " + " AND ".join(conditions)
    
    base_query += " ORDER BY a.applied_at DESC"
    
    # Get applications using raw SQL
    applications_data = db.execute_query(base_query, params)
    
    # Enhanced application statistics with percentages using raw SQL
    stats = db.get_application_stats_by_jobseeker(request.user.id)
    
    total_applications = stats['total_applications'] if stats else 0
    pending_applications = stats['applied'] if stats else 0
    reviewing_applications = stats['reviewing'] if stats else 0
    shortlisted_applications = stats['shortlisted'] if stats else 0
    interview_applications = stats['interviewing'] if stats else 0
    accepted_applications = stats['hired'] if stats else 0
    rejected_applications = stats['rejected'] if stats else 0
    
    # Calculate percentages for progress bars
    if total_applications > 0:
        pending_percentage = round((pending_applications / total_applications) * 100)
        shortlisted_percentage = round((shortlisted_applications / total_applications) * 100)
        accepted_percentage = round((accepted_applications / total_applications) * 100)
    else:
        pending_percentage = 0
        shortlisted_percentage = 0
        accepted_percentage = 0
    
    # Manual pagination for SQLite3 results
    page_number = request.GET.get('page', 1)
    try:
        page_number = int(page_number)
    except (ValueError, TypeError):
        page_number = 1
    
    per_page = 10
    offset = (page_number - 1) * per_page
    
    # Get paginated results
    paginated_query = base_query + f" LIMIT {per_page} OFFSET {offset}"
    applications_page = db.execute_query(paginated_query, params)
    
    # Convert datetime strings to datetime objects for template filters
    from datetime import datetime
    for application in applications_page:
        if application.get('applied_at') and isinstance(application['applied_at'], str):
            try:
                # Parse ISO format datetime string
                application['applied_at'] = datetime.fromisoformat(application['applied_at'].replace('Z', '+00:00'))
            except (ValueError, AttributeError):
                pass
        if application.get('updated_at') and isinstance(application['updated_at'], str):
            try:
                application['updated_at'] = datetime.fromisoformat(application['updated_at'].replace('Z', '+00:00'))
            except (ValueError, AttributeError):
                pass
        if application.get('reviewed_at') and isinstance(application['reviewed_at'], str):
            try:
                application['reviewed_at'] = datetime.fromisoformat(application['reviewed_at'].replace('Z', '+00:00'))
            except (ValueError, AttributeError):
                pass
    
    # Calculate pagination info
    total_pages = (total_applications + per_page - 1) // per_page
    has_previous = page_number > 1
    has_next = page_number < total_pages
    
    # Create pagination object-like structure
    class PaginationInfo:
        def __init__(self, data, page_num, total_items, per_page):
            self.object_list = data
            self.number = page_num
            self.paginator = type('Paginator', (), {
                'num_pages': (total_items + per_page - 1) // per_page,
                'count': total_items
            })()
            self.has_previous = lambda: page_num > 1
            self.has_next = lambda: page_num < self.paginator.num_pages
            self.previous_page_number = lambda: page_num - 1 if self.has_previous() else None
            self.next_page_number = lambda: page_num + 1 if self.has_next() else None
    
    page_obj = PaginationInfo(applications_page, page_number, total_applications, per_page)
    
    # Define status choices directly (replacing Django model reference)
    status_choices = [
        ('applied', 'Applied'),
        ('reviewing', 'Under Review'),
        ('shortlisted', 'Shortlisted'),
        ('interviewing', 'Interview'),
        ('hired', 'Hired'),
        ('rejected', 'Rejected'),
    ]
    
    context = {
        'applications': applications_page,
        'page_obj': page_obj,
        'user_profile': user_data,  # Use raw SQL user data
        'total_applications': total_applications,
        'pending_applications': pending_applications,
        'reviewing_applications': reviewing_applications,
        'shortlisted_applications': shortlisted_applications,
        'interview_applications': interview_applications,
        'accepted_applications': accepted_applications,
        'rejected_applications': rejected_applications,
        'pending_percentage': pending_percentage,
        'shortlisted_percentage': shortlisted_percentage,
        'accepted_percentage': accepted_percentage,
        'status_choices': status_choices,
        'current_status': status_filter,
        'current_date_filter': date_filter,
        'current_job_type': job_type_filter,
        'current_company': company_search,
        'current_job_title': job_title_search,
        'search_query': search,
    }
    
    return render(request, 'accounts/applications_list.html', context)


@login_required
def job_alerts_list(request):
    """User job alerts list view"""
    user_profile = request.user.userprofile
    
    if user_profile.user_type != 'jobseeker':
        return redirect('dashboard')
    
    job_seeker_profile = get_object_or_404(JobSeekerProfile, user_profile=user_profile)
    job_alerts = JobAlert.objects.filter(user=job_seeker_profile).order_by('-created_at')
    
    # Calculate statistics
    total_alerts = job_alerts.count()
    active_alerts = job_alerts.filter(is_active=True).count()
    
    context = {
        'job_alerts': job_alerts,
        'total_alerts': total_alerts,
        'active_alerts': active_alerts,
        'jobs_found': 0,  # Placeholder for now
        'emails_sent': 0,  # Placeholder for now
        'user_profile': user_profile,
    }
    
    return render(request, 'accounts/job_alerts_list.html', context)

@login_required
@require_http_methods(["POST"])
def save_job(request, job_id):
    """Save job view (AJAX)"""
    if request.user.userprofile.user_type != 'jobseeker':
        return JsonResponse({'success': False, 'message': 'Unauthorized'})
    
    job = get_object_or_404(JobPost, id=job_id)
    job_seeker_profile = get_object_or_404(JobSeekerProfile, user_profile=request.user.userprofile)
    
    saved_job, created = SavedJob.objects.get_or_create(
        job=job,
        user=job_seeker_profile
    )
    
    if created:
        return JsonResponse({'success': True, 'message': 'Job saved successfully!', 'action': 'saved'})
    else:
        saved_job.delete()
        return JsonResponse({'success': True, 'message': 'Job removed from saved jobs!', 'action': 'removed'})

@login_required
@require_http_methods(["POST"])
def create_job_alert(request):
    """Create job alert (AJAX)"""
    if request.user.userprofile.user_type != 'jobseeker':
        return JsonResponse({'success': False, 'message': 'Unauthorized'})
    
    try:
        job_seeker_profile = get_object_or_404(JobSeekerProfile, user_profile=request.user.userprofile)
        
        # Get form data
        alert_name = request.POST.get('alert_name', '').strip()
        keywords = request.POST.get('keywords', '').strip()
        location = request.POST.get('location', '').strip()
        frequency = request.POST.get('frequency', 'weekly')
        employment_type = request.POST.get('employment_type', '')
        experience_level = request.POST.get('experience_level', '')
        salary_min = request.POST.get('salary_min')
        salary_max = request.POST.get('salary_max')
        email_notifications = request.POST.get('email_notifications') == 'on'
        remote_only = request.POST.get('remote_only') == 'on'
        
        if not alert_name:
            return JsonResponse({'success': False, 'message': 'Alert name is required'})
        
        # Create job alert
        job_alert = JobAlert.objects.create(
            user=job_seeker_profile,
            title=alert_name,
            keywords=keywords or None,
            frequency=frequency,
            employment_type=employment_type or None,
            experience_level=experience_level or None,
            min_salary=float(salary_min) if salary_min else None,
            max_salary=float(salary_max) if salary_max else None,
            is_remote=remote_only,
            is_active=True
        )
        
        return JsonResponse({
            'success': True, 
            'message': 'Job alert created successfully!',
            'alert_id': job_alert.id
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)})

@login_required
@require_http_methods(["POST"])
def update_job_alert(request, alert_id):
    """Update job alert (AJAX)"""
    if request.user.userprofile.user_type != 'jobseeker':
        return JsonResponse({'success': False, 'message': 'Unauthorized'})
    
    try:
        job_seeker_profile = get_object_or_404(JobSeekerProfile, user_profile=request.user.userprofile)
        job_alert = get_object_or_404(JobAlert, id=alert_id, user=job_seeker_profile)
        
        # Update alert fields
        job_alert.title = request.POST.get('edit_alert_name', job_alert.title)
        job_alert.keywords = request.POST.get('edit_keywords', '') or None
        job_alert.frequency = request.POST.get('edit_frequency', job_alert.frequency)
        job_alert.employment_type = request.POST.get('edit_employment_type', '') or None
        job_alert.experience_level = request.POST.get('edit_experience_level', '') or None
        
        salary_min = request.POST.get('edit_salary_min')
        salary_max = request.POST.get('edit_salary_max')
        job_alert.min_salary = float(salary_min) if salary_min else None
        job_alert.max_salary = float(salary_max) if salary_max else None
        
        job_alert.is_remote = request.POST.get('edit_remote_only') == 'on'
        job_alert.save()
        
        return JsonResponse({'success': True, 'message': 'Job alert updated successfully!'})
        
    except JobAlert.DoesNotExist:
        return JsonResponse({'success': False, 'message': 'Job alert not found'})
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)})

@login_required
@require_http_methods(["POST"])
def delete_job_alert(request, alert_id):
    """Delete job alert (AJAX)"""
    if request.user.userprofile.user_type != 'jobseeker':
        return JsonResponse({'success': False, 'message': 'Unauthorized'})
    
    try:
        job_seeker_profile = get_object_or_404(JobSeekerProfile, user_profile=request.user.userprofile)
        job_alert = get_object_or_404(JobAlert, id=alert_id, user=job_seeker_profile)
        
        alert_name = job_alert.title
        job_alert.delete()
        
        return JsonResponse({
            'success': True, 
            'message': f'Job alert "{alert_name}" deleted successfully!'
        })
        
    except JobAlert.DoesNotExist:
        return JsonResponse({'success': False, 'message': 'Job alert not found'})
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)})

@login_required
@require_http_methods(["POST"])
def toggle_job_alert(request, alert_id):
    """Toggle job alert active status (AJAX)"""
    if request.user.userprofile.user_type != 'jobseeker':
        return JsonResponse({'success': False, 'message': 'Unauthorized'})
    
    try:
        job_seeker_profile = get_object_or_404(JobSeekerProfile, user_profile=request.user.userprofile)
        job_alert = get_object_or_404(JobAlert, id=alert_id, user=job_seeker_profile)
        
        job_alert.is_active = not job_alert.is_active
        job_alert.save()
        
        status = 'activated' if job_alert.is_active else 'paused'
        return JsonResponse({
            'success': True, 
            'message': f'Job alert {status} successfully!',
            'is_active': job_alert.is_active
        })
        
    except JobAlert.DoesNotExist:
        return JsonResponse({'success': False, 'message': 'Job alert not found'})
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)})

@login_required
@require_http_methods(["GET"])
def get_job_alert(request, alert_id):
    """Get job alert details for editing (AJAX)"""
    if request.user.userprofile.user_type != 'jobseeker':
        return JsonResponse({'success': False, 'message': 'Unauthorized'})
    
    try:
        job_seeker_profile = get_object_or_404(JobSeekerProfile, user_profile=request.user.userprofile)
        job_alert = get_object_or_404(JobAlert, id=alert_id, user=job_seeker_profile)
        
        return JsonResponse({
            'success': True,
            'alert': {
                'id': job_alert.id,
                'title': job_alert.title,
                'keywords': job_alert.keywords or '',
                'frequency': job_alert.frequency,
                'employment_type': job_alert.employment_type or '',
                'experience_level': job_alert.experience_level or '',
                'min_salary': str(job_alert.min_salary) if job_alert.min_salary else '',
                'max_salary': str(job_alert.max_salary) if job_alert.max_salary else '',
                'is_remote': job_alert.is_remote,
                'is_active': job_alert.is_active
            }
        })
        
    except JobAlert.DoesNotExist:
        return JsonResponse({'success': False, 'message': 'Job alert not found'})
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)})

@login_required
def delete_education(request, education_id):
    """Delete education view"""
    education = get_object_or_404(Education, id=education_id)
    
    if education.job_seeker.user_profile.user != request.user:
        messages.error(request, 'You are not authorized to delete this education.')
        return redirect('accounts:profile_detail')
    
    education.delete()
    messages.success(request, 'Education deleted successfully!')
    return redirect('accounts:profile_detail')

@login_required
def delete_experience(request, experience_id):
    """Delete experience view"""
    experience = get_object_or_404(Experience, id=experience_id)
    
    if experience.job_seeker.user_profile.user != request.user:
        messages.error(request, 'You are not authorized to delete this experience.')
        return redirect('accounts:profile_detail')
    
    experience.delete()
    messages.success(request, 'Experience deleted successfully!')
    return redirect('accounts:profile_detail')

@login_required
def delete_skill(request, skill_id):
    """Delete skill view"""
    skill = get_object_or_404(Skill, id=skill_id)
    
    if skill.job_seeker.user_profile.user != request.user:
        messages.error(request, 'You are not authorized to delete this skill.')
        return redirect('accounts:profile_detail')
    
    skill.delete()
    messages.success(request, 'Skill deleted successfully!')
    return redirect('accounts:profile_detail')

@login_required
def delete_certification(request, certification_id):
    """Delete certification view"""
    certification = get_object_or_404(Certification, id=certification_id)
    
    if certification.job_seeker.user_profile.user != request.user:
        messages.error(request, 'You are not authorized to delete this certification.')
        return redirect('accounts:profile_detail')
    
    certification.delete()
    messages.success(request, 'Certification deleted successfully!')
    return redirect('accounts:profile_detail')

# Employer registration is now handled by the unified register view

def password_reset(request):
    """Enhanced password reset view with rate limiting and security features"""
    if request.user.is_authenticated:
        return redirect('accounts:dashboard')
    
    if request.method == 'POST':
        email = request.POST.get('email')
        
        # Rate limiting check
        client_ip = request.META.get('REMOTE_ADDR')
        rate_limit_key = f'password_reset_rate_{client_ip}'
        rate_limit_data = request.session.get(rate_limit_key, {'count': 0, 'last_request': None})
        
        current_time = timezone.now()
        if rate_limit_data['last_request']:
            last_request = datetime.fromisoformat(rate_limit_data['last_request'])
            time_diff = (current_time - last_request).total_seconds()
            
            # Reset count if more than 5 minutes have passed
            if time_diff > 300:  # 5 minutes
                rate_limit_data = {'count': 0, 'last_request': None}
            elif rate_limit_data['count'] >= 3:
                messages.error(request, 'Too many password reset requests. Please wait 5 minutes before trying again.')
                return render(request, 'accounts/password_reset.html')
        
        # Update rate limit data
        rate_limit_data['count'] += 1
        rate_limit_data['last_request'] = current_time.isoformat()
        request.session[rate_limit_key] = rate_limit_data
        
        # Email validation
        if not email or '@' not in email:
            messages.error(request, 'Please enter a valid email address.')
            return render(request, 'accounts/password_reset.html')
        
        try:
            user = User.objects.get(email=email)
            
            # Check if user already has a pending reset token
            existing_tokens = [key for key in request.session.keys() if key.startswith('password_reset_') and key != rate_limit_key]
            for token_key in existing_tokens:
                token_data = request.session.get(token_key)
                if token_data and token_data.get('email') == email:
                    expiry = datetime.fromisoformat(token_data['expires'])
                    if timezone.now() < expiry:
                        messages.info(request, 'A password reset link has already been sent to this email. Please check your inbox or wait for the current link to expire.')
                        return render(request, 'accounts/password_reset.html')
            
            # Generate a unique token for password reset
            token = ''.join(secrets.choice('abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789') for _ in range(64))
            
            # Store token in session with expiry
            request.session[f'password_reset_{token}'] = {
                'user_id': user.id,
                'email': user.email,
                'expires': (timezone.now() + timezone.timedelta(hours=24)).isoformat(),
                'created_at': timezone.now().isoformat(),
                'ip_address': client_ip
            }
            
            # Send password reset email
            current_site = request.get_host()
            reset_url = f"http://{current_site}{reverse('accounts:password_reset_confirm', kwargs={'token': token})}"
            
            # Create email content
            subject = '🔐 Reset your Hireo password - Secure Link Inside'
            html_message = render_to_string('accounts/email/password_reset_email.html', {
                'user': user,
                'reset_url': reset_url,
                'expiry_hours': 24,
                'ip_address': client_ip,
                'timestamp': current_time.strftime('%Y-%m-%d %H:%M:%S UTC')
            })
            plain_message = strip_tags(html_message)
            
            # Send email
            send_mail(
                subject,
                plain_message,
                settings.EMAIL_HOST_USER,
                [user.email],
                html_message=html_message,
                fail_silently=False,
            )
            
            messages.success(request, '✅ If an account with that email exists, you will receive a secure password reset link shortly. Please check your inbox and spam folder.')
        except User.DoesNotExist:
            # Don't reveal if email exists or not for security
            messages.success(request, '✅ If an account with that email exists, you will receive a secure password reset link shortly. Please check your inbox and spam folder.')
        except Exception as e:
            messages.error(request, 'An error occurred while sending the reset email. Please try again later.')
        
        return redirect('accounts:login')
    
    return render(request, 'accounts/password_reset.html')

def password_reset_confirm(request, token):
    """Enhanced password reset confirmation with advanced security and validation"""
    # Check if token exists and is valid
    session_key = f'password_reset_{token}'
    token_data = request.session.get(session_key)
    
    if not token_data:
        messages.error(request, '❌ Invalid or expired password reset link.')
        return redirect('accounts:password_reset')
    
    # Check if token is expired
    expiry = datetime.fromisoformat(token_data['expires'])
    if timezone.now() > expiry:
        # Remove expired token
        del request.session[session_key]
        messages.error(request, '⏰ Your password reset link has expired. Please request a new one.')
        return redirect('accounts:password_reset')
    
    # Additional security check - verify IP if available
    current_ip = request.META.get('REMOTE_ADDR')
    if token_data.get('ip_address') and current_ip != token_data.get('ip_address'):
        messages.warning(request, '🔒 Security Notice: This reset link was requested from a different IP address. If this wasn\'t you, please contact support.')
    
    if request.method == 'POST':
        password1 = request.POST.get('password1')
        password2 = request.POST.get('password2')
        
        # Basic validation
        if not password1 or not password2:
            messages.error(request, '❌ Please enter both password fields.')
            return render(request, 'accounts/password_reset_confirm.html', {'token': token})
        
        if password1 != password2:
            messages.error(request, '❌ The two password fields didn\'t match.')
            return render(request, 'accounts/password_reset_confirm.html', {'token': token})
        
        # Enhanced password strength validation
        password_errors = []
        
        if len(password1) < 8:
            password_errors.append('Password must be at least 8 characters long')
        
        if not any(c.isupper() for c in password1):
            password_errors.append('Password must contain at least one uppercase letter')
        
        if not any(c.islower() for c in password1):
            password_errors.append('Password must contain at least one lowercase letter')
        
        if not any(c.isdigit() for c in password1):
            password_errors.append('Password must contain at least one number')
        
        if not any(c in '!@#$%^&*()_+-=[]{}|;:,.<>?' for c in password1):
            password_errors.append('Password must contain at least one special character')
        
        # Check for common weak passwords
        common_passwords = ['password', '12345678', 'qwerty123', 'admin123', 'welcome123']
        if password1.lower() in common_passwords:
            password_errors.append('This password is too common. Please choose a more unique password')
        
        # Check if password contains email or username
        try:
            user = User.objects.get(id=token_data['user_id'])
            if user.email.split('@')[0].lower() in password1.lower():
                password_errors.append('Password should not contain your email address')
            if user.username.lower() in password1.lower():
                password_errors.append('Password should not contain your username')
        except User.DoesNotExist:
            pass
        
        if password_errors:
            for error in password_errors:
                messages.error(request, f'❌ {error}')
            return render(request, 'accounts/password_reset_confirm.html', {'token': token})
        
        try:
            # Get user and update password
            user = User.objects.get(id=token_data['user_id'], email=token_data['email'])
            user.set_password(password1)
            user.save()
            
            # Remove token from session
            del request.session[session_key]
            
            # Remove any other pending reset tokens for this user
            tokens_to_remove = []
            for key in request.session.keys():
                if key.startswith('password_reset_') and key != session_key:
                    session_token_data = request.session.get(key)
                    if session_token_data and session_token_data.get('user_id') == user.id:
                        tokens_to_remove.append(key)
            
            for key in tokens_to_remove:
                del request.session[key]
            
            messages.success(request, '✅ Your password has been reset successfully! You can now log in with your new secure password.')
            return redirect('accounts:login')
        except User.DoesNotExist:
            messages.error(request, '❌ An error occurred. Please try again.')
            return redirect('accounts:password_reset')
    
    # Get user info for display
    context = {'token': token}
    try:
        user = User.objects.get(id=token_data['user_id'])
        context['user_email'] = user.email
        context['reset_requested_at'] = datetime.fromisoformat(token_data.get('created_at', token_data['expires']))
    except (User.DoesNotExist, KeyError):
        pass
    
    return render(request, 'accounts/password_reset_confirm.html', context)


@login_required
def cover_letter(request):
    """Cover Letter Creator tool view"""
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    return render(request, 'accounts/cover_letter.html')

@login_required
def interview_prep(request):
    """Interview Preparation Tools view"""
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    return render(request, 'accounts/interview_prep.html')

@login_required
def gamification(request):
    """Career achievements and gamification view"""
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    user_profile = request.user.userprofile
    jobseeker_profile = get_object_or_404(JobSeekerProfile, user_profile=user_profile)
    
    # Calculate achievements based on user profile data
    achievements = calculate_user_achievements(request.user, jobseeker_profile)
    
    # Calculate user stats
    user_stats = calculate_user_stats(request.user, jobseeker_profile)
    
    context = {
        'user': request.user,
        'profile': user_profile,
        'jobseeker_profile': jobseeker_profile,
        'achievements': achievements,
        'user_stats': user_stats,
    }
    return render(request, 'accounts/gamification.html', context)

@login_required
def skill_assessment(request):
    """AI-powered skill assessment platform"""
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    user_profile = request.user.userprofile
    jobseeker_profile = get_object_or_404(JobSeekerProfile, user_profile=user_profile)
    
    context = {
        'user': request.user,
        'profile': user_profile,
        'jobseeker_profile': jobseeker_profile,
    }
    return render(request, 'accounts/skill_assessment.html', context)

def calculate_user_achievements(user, jobseeker_profile):
    """Calculate user achievements based on profile data"""
    achievements = []
    
    # Profile Completion Achievements
    profile_completion = calculate_profile_completion(user, jobseeker_profile)
    
    achievements.append({
        'id': 'profile_complete',
        'name': 'Profile Master',
        'description': 'Complete your profile 100%',
        'icon': 'fas fa-user-check',
        'category': 'Profile',
        'completed': profile_completion >= 100,
        'progress': profile_completion,
        'max_progress': 100,
        'xp_reward': 100,
        'badge_color': 'success' if profile_completion >= 100 else 'secondary'
    })
    
    achievements.append({
        'id': 'profile_basic',
        'name': 'Getting Started',
        'description': 'Complete basic profile information',
        'icon': 'fas fa-user-plus',
        'category': 'Profile',
        'completed': profile_completion >= 50,
        'progress': min(profile_completion, 50),
        'max_progress': 50,
        'xp_reward': 50,
        'badge_color': 'success' if profile_completion >= 50 else 'secondary'
    })
    
    # Application Achievements
    applications_count = Application.objects.filter(applicant=jobseeker_profile).count()
    
    achievements.append({
        'id': 'first_application',
        'name': 'Job Hunter',
        'description': 'Submit your first job application',
        'icon': 'fas fa-paper-plane',
        'category': 'Applications',
        'completed': applications_count >= 1,
        'progress': min(applications_count, 1),
        'max_progress': 1,
        'xp_reward': 75,
        'badge_color': 'success' if applications_count >= 1 else 'secondary'
    })
    
    achievements.append({
        'id': 'application_streak',
        'name': 'Application Master',
        'description': 'Submit 10 job applications',
        'icon': 'fas fa-fire',
        'category': 'Applications',
        'completed': applications_count >= 10,
        'progress': min(applications_count, 10),
        'max_progress': 10,
        'xp_reward': 200,
        'badge_color': 'success' if applications_count >= 10 else 'secondary'
    })
    
    # Experience Achievements
    experience_years = jobseeker_profile.experience_years or 0
    
    achievements.append({
        'id': 'experienced_professional',
        'name': 'Experienced Professional',
        'description': 'Have 3+ years of experience',
        'icon': 'fas fa-briefcase',
        'category': 'Experience',
        'completed': experience_years >= 3,
        'progress': min(experience_years, 3),
        'max_progress': 3,
        'xp_reward': 150,
        'badge_color': 'success' if experience_years >= 3 else 'secondary'
    })
    
    # Skills Achievements
    skills_count = len(jobseeker_profile.skills.split(',')) if jobseeker_profile.skills else 0
    
    achievements.append({
        'id': 'skill_collector',
        'name': 'Skill Collector',
        'description': 'Add 5+ skills to your profile',
        'icon': 'fas fa-cogs',
        'category': 'Skills',
        'completed': skills_count >= 5,
        'progress': min(skills_count, 5),
        'max_progress': 5,
        'xp_reward': 100,
        'badge_color': 'success' if skills_count >= 5 else 'secondary'
    })
    
    # Education Achievements
    education_count = Education.objects.filter(job_seeker=jobseeker_profile).count()
    
    achievements.append({
        'id': 'educated',
        'name': 'Scholar',
        'description': 'Add education information',
        'icon': 'fas fa-graduation-cap',
        'category': 'Education',
        'completed': education_count >= 1,
        'progress': min(education_count, 1),
        'max_progress': 1,
        'xp_reward': 75,
        'badge_color': 'success' if education_count >= 1 else 'secondary'
    })
    
    # Recent Activity Achievements
    recent_applications = Application.objects.filter(
        applicant=jobseeker_profile,
        applied_at__gte=timezone.now() - timedelta(days=7)
    ).count()
    
    achievements.append({
        'id': 'active_seeker',
        'name': 'Active Job Seeker',
        'description': 'Apply to jobs this week',
        'icon': 'fas fa-rocket',
        'category': 'Activity',
        'completed': recent_applications >= 1,
        'progress': min(recent_applications, 1),
        'max_progress': 1,
        'xp_reward': 50,
        'badge_color': 'success' if recent_applications >= 1 else 'secondary'
    })
    
    return achievements

def calculate_profile_completion(user, jobseeker_profile):
    """Calculate profile completion percentage"""
    completion_score = 0
    total_fields = 10
    
    # Basic user info (3 points)
    if user.first_name and user.last_name:
        completion_score += 1
    if user.email:
        completion_score += 1
    if user.userprofile.phone:
        completion_score += 1
    
    # JobSeeker profile info (7 points)
    if jobseeker_profile.skills:
        completion_score += 1
    if jobseeker_profile.experience_years:
        completion_score += 1
    if jobseeker_profile.expected_salary:
        completion_score += 1
    if jobseeker_profile.preferred_location:
        completion_score += 1
    if jobseeker_profile.resume:
        completion_score += 1
    if jobseeker_profile.linkedin_url:
        completion_score += 1
    if user.userprofile.bio:
        completion_score += 1
    
    return int((completion_score / total_fields) * 100)

def calculate_user_stats(user, jobseeker_profile):
    """Calculate user statistics for gamification"""
    achievements = calculate_user_achievements(user, jobseeker_profile)
    completed_achievements = [a for a in achievements if a['completed']]
    
    # Calculate total XP
    total_xp = sum(a['xp_reward'] for a in completed_achievements)
    
    # Calculate level (every 200 XP = 1 level)
    level = max(1, total_xp // 200)
    xp_for_next_level = ((level + 1) * 200) - total_xp
    
    # Calculate profile completion
    profile_completion = calculate_profile_completion(user, jobseeker_profile)
    
    # Calculate streak (simplified - days since last application)
    last_application = Application.objects.filter(applicant=jobseeker_profile).order_by('-applied_at').first()
    if last_application:
        days_since = (timezone.now() - last_application.applied_at).days
        streak = max(0, 7 - days_since) if days_since <= 7 else 0
    else:
        streak = 0
    
    return {
        'level': level,
        'total_xp': total_xp,
        'xp_for_next_level': xp_for_next_level,
        'completed_achievements': len(completed_achievements),
        'total_achievements': len(achievements),
        'profile_completion': profile_completion,
        'streak': streak,
        'level_progress': ((total_xp % 200) / 200) * 100
    }

@login_required
def resume_templates(request):
    """Resume templates gallery view"""
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    # Get user's existing resumes
    from .models import Resume
    user_resumes = Resume.objects.filter(user_profile=request.user.userprofile)
    
    context = {
        'user': request.user,
        'profile': request.user.userprofile,
        'user_resumes': user_resumes,
    }
    return render(request, 'accounts/resume_templates.html', context)


@login_required
def resume_builder(request, resume_id=None):
    """Resume builder view for creating/editing resumes"""
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    # Simple view that just renders the template for live preview
    context = {
        'user': request.user,
    }
    return render(request, 'accounts/resume_builder.html', context)


def populate_resume_from_profile(resume, user_profile):
    """Pre-populate resume with user profile data"""
    try:
        jobseeker_profile = user_profile.jobseekerprofile
        
        # Basic info
        resume.full_name = f"{user_profile.user.first_name} {user_profile.user.last_name}".strip()
        resume.email = user_profile.user.email
        resume.phone = getattr(jobseeker_profile, 'phone', '')
        resume.address = getattr(jobseeker_profile, 'location', '')
        resume.linkedin_url = getattr(jobseeker_profile, 'linkedin_url', '')
        resume.github_url = getattr(jobseeker_profile, 'github_url', '')
        
        # Professional summary from bio
        if hasattr(jobseeker_profile, 'bio') and jobseeker_profile.bio:
            resume.professional_summary = jobseeker_profile.bio
        
        resume.save()
        
        # Import skills
        from .models import ResumeSkill, ResumeEducation, ResumeWorkExperience
        for skill in jobseeker_profile.skill_set.all():
            ResumeSkill.objects.create(
                resume=resume,
                name=skill.name,
                category='technical',
                level='intermediate'
            )
        
        # Import education
        for education in jobseeker_profile.education_set.all():
            ResumeEducation.objects.create(
                resume=resume,
                institution=education.institution,
                degree_type='bachelor',  # Default
                field_of_study=education.field_of_study,
                start_date=education.start_date or timezone.now().date(),
                end_date=education.end_date,
                description=education.description or ''
            )
        
        # Import work experience
        for experience in jobseeker_profile.workexperience_set.all():
            ResumeWorkExperience.objects.create(
                resume=resume,
                company_name=experience.company,
                position=experience.position,
                start_date=experience.start_date,
                end_date=experience.end_date,
                is_current=experience.is_current,
                description=experience.description or ''
            )
            
    except Exception as e:
        # If profile doesn't exist or has issues, continue with empty resume
        pass


def handle_resume_form_submission(request, resume):
    """Handle resume form submission and updates"""
    from django.http import JsonResponse
    
    action = request.POST.get('action')
    
    if action == 'update_basic_info':
        # Update basic information
        resume.full_name = request.POST.get('full_name', '')
        resume.email = request.POST.get('email', '')
        resume.phone = request.POST.get('phone', '')
        resume.address = request.POST.get('address', '')
        resume.website = request.POST.get('website', '')
        resume.linkedin_url = request.POST.get('linkedin_url', '')
        resume.github_url = request.POST.get('github_url', '')
        resume.professional_summary = request.POST.get('professional_summary', '')
        resume.save()
        
        return JsonResponse({'success': True, 'message': 'Basic information updated'})
    
    elif action == 'update_template':
        # Update template and styling
        resume.template = request.POST.get('template', resume.template)
        resume.color_scheme = request.POST.get('color_scheme', resume.color_scheme)
        resume.font_family = request.POST.get('font_family', resume.font_family)
        resume.save()
        
        return JsonResponse({'success': True, 'message': 'Template updated'})
    
    elif action == 'add_work_experience':
        # Add new work experience
        ResumeWorkExperience.objects.create(
            resume=resume,
            company_name=request.POST.get('company_name', ''),
            position=request.POST.get('position', ''),
            location=request.POST.get('location', ''),
            start_date=request.POST.get('start_date'),
            end_date=request.POST.get('end_date') if request.POST.get('end_date') else None,
            is_current=request.POST.get('is_current') == 'on',
            description=request.POST.get('description', '')
        )
        
        return JsonResponse({'success': True, 'message': 'Work experience added'})
    
    elif action == 'add_education':
        # Add new education
        ResumeEducation.objects.create(
            resume=resume,
            institution=request.POST.get('institution', ''),
            degree_type=request.POST.get('degree_type', ''),
            field_of_study=request.POST.get('field_of_study', ''),
            location=request.POST.get('location', ''),
            start_date=request.POST.get('start_date'),
            end_date=request.POST.get('end_date') if request.POST.get('end_date') else None,
            gpa=request.POST.get('gpa') if request.POST.get('gpa') else None,
            description=request.POST.get('description', '')
        )
        
        return JsonResponse({'success': True, 'message': 'Education added'})
    
    elif action == 'add_skill':
        # Add new skill
        skill_name = request.POST.get('skill_name', '')
        if skill_name and not resume.skills.filter(name=skill_name).exists():
            ResumeSkill.objects.create(
                resume=resume,
                name=skill_name,
                level=request.POST.get('skill_level', 'intermediate'),
                category=request.POST.get('skill_category', 'technical'),
                years_experience=request.POST.get('years_experience') if request.POST.get('years_experience') else None
            )
            return JsonResponse({'success': True, 'message': 'Skill added'})
        else:
            return JsonResponse({'success': False, 'message': 'Skill already exists or invalid'})
    
    elif action == 'export_resume':
        # Handle resume export
        export_format = request.POST.get('format', 'pdf')
        return export_resume(resume, export_format)
    
    return JsonResponse({'success': False, 'message': 'Invalid action'})


def export_resume(resume, format_type):
    """Export resume in specified format"""
    from django.http import HttpResponse
    import json
    
    if format_type == 'pdf':
        # For now, return JSON data - PDF generation would require additional libraries
        response_data = {
            'success': True,
            'message': 'PDF export functionality will be implemented with reportlab/weasyprint',
            'download_url': f'/accounts/resume/{resume.id}/export/pdf/'
        }
        return JsonResponse(response_data)
    
    elif format_type == 'html':
        # Generate HTML version
        from django.template.loader import render_to_string
        html_content = render_to_string('accounts/resume_export.html', {'resume': resume})
        
        response = HttpResponse(html_content, content_type='text/html')
        response['Content-Disposition'] = f'attachment; filename="{resume.title}.html"'
        return response
    
    elif format_type == 'docx':
        # DOCX export would require python-docx library
        response_data = {
            'success': True,
            'message': 'DOCX export functionality will be implemented with python-docx',
            'download_url': f'/accounts/resume/{resume.id}/export/docx/'
        }
        return JsonResponse(response_data)
    
    return JsonResponse({'success': False, 'message': 'Invalid export format'})

@login_required
def career_roadmap(request):
    """Career development roadmap view"""
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    # Handle career path selection
    if request.method == 'POST':
        career_path_id = request.POST.get('career_path_id')
        if career_path_id:
            career_path = get_object_or_404(CareerPath, id=career_path_id)
            user_progress, created = UserCareerProgress.objects.get_or_create(
                user_profile=request.user.userprofile,
                career_path=career_path,
                defaults={'is_active': True}
            )
            if created:
                # Create milestone progress entries for all milestones
                for milestone in career_path.milestones.all():
                    MilestoneProgress.objects.create(
                        user_career_progress=user_progress,
                        milestone=milestone
                    )
                messages.success(request, f'Successfully enrolled in {career_path.name} career path!')
            else:
                user_progress.is_active = True
                user_progress.save()
                messages.info(request, f'Resumed progress in {career_path.name} career path!')
            return redirect('accounts:career_progress', career_id=career_path.id)
    
    # Get all available career paths
    career_paths = CareerPath.objects.all().order_by('category', 'name')
    
    # Get user's current career progress
    user_career_progress = UserCareerProgress.objects.filter(
        user_profile=request.user.userprofile,
        is_active=True
    ).select_related('career_path')
    
    # Group career paths by category and process skills
    career_categories = {}
    for path in career_paths:
        category = path.get_category_display()
        if category not in career_categories:
            career_categories[category] = []
        
        # Process required skills for template display
        if path.required_skills:
            skills_list = [skill.strip() for skill in path.required_skills.split(',')]
            path.skills_display = skills_list[:3]  # First 3 skills
        else:
            path.skills_display = []
            
        career_categories[category].append(path)
    
    context = {
        'user': request.user,
        'profile': request.user.userprofile,
        'career_categories': career_categories,
        'user_career_progress': user_career_progress,
    }
    return render(request, 'accounts/career_roadmap.html', context)







@login_required
def nepal_salary_guide(request):
    if not request.user.is_authenticated:
        return redirect('accounts:login')
    
    if hasattr(request.user, 'employerprofile'):
        return redirect('employer:dashboard')
    
    # Comprehensive Nepal salary data with detailed job information
    nepal_salary_data = {
        'it_tech': [
            {
                'position': 'Software Developer',
                'salary_range': '35,000 - 80,000',
                'annual_range': '4.2L - 9.6L',
                'responsibilities': 'Design, develop, and maintain software applications. Write clean, efficient code. Debug and troubleshoot issues. Collaborate with cross-functional teams.',
                'requirements': 'Bachelor\'s in Computer Science or related field. Proficiency in programming languages (Java, Python, JavaScript). Problem-solving skills. 0-3 years experience.',
                'skills': 'Programming, Software Architecture, Version Control (Git), Database Management, API Development'
            },
            {
                'position': 'Data Scientist',
                'salary_range': '45,000 - 120,000',
                'annual_range': '5.4L - 14.4L',
                'responsibilities': 'Analyze complex datasets to extract insights. Build predictive models and machine learning algorithms. Create data visualizations and reports. Present findings to stakeholders.',
                'requirements': 'Master\'s in Data Science, Statistics, or related field. Strong analytical and mathematical skills. Experience with Python/R, SQL, and ML frameworks. 2-5 years experience.',
                'skills': 'Machine Learning, Statistical Analysis, Python/R, SQL, Data Visualization, Big Data Technologies'
            },
            {
                'position': 'DevOps Engineer',
                'salary_range': '40,000 - 90,000',
                'annual_range': '4.8L - 10.8L',
                'responsibilities': 'Manage CI/CD pipelines. Automate deployment processes. Monitor system performance and reliability. Manage cloud infrastructure and containerization.',
                'requirements': 'Bachelor\'s in Engineering or Computer Science. Experience with cloud platforms (AWS, Azure). Knowledge of containerization and automation tools. 2-4 years experience.',
                'skills': 'Cloud Platforms, Docker/Kubernetes, CI/CD, Infrastructure as Code, Monitoring Tools, Scripting'
            },
            {
                'position': 'UI/UX Designer',
                'salary_range': '30,000 - 70,000',
                'annual_range': '3.6L - 8.4L',
                'responsibilities': 'Design user interfaces and user experiences. Create wireframes, prototypes, and mockups. Conduct user research and usability testing. Collaborate with development teams.',
                'requirements': 'Bachelor\'s in Design or related field. Proficiency in design tools (Figma, Adobe Creative Suite). Understanding of user-centered design principles. Portfolio required.',
                'skills': 'UI/UX Design, Prototyping, User Research, Design Tools, HTML/CSS Basics, Design Systems'
            },
            {
                'position': 'Product Manager',
                'salary_range': '50,000 - 150,000',
                'annual_range': '6L - 18L',
                'responsibilities': 'Define product strategy and roadmap. Gather and prioritize requirements. Coordinate with engineering, design, and marketing teams. Analyze market trends and user feedback.',
                'requirements': 'Bachelor\'s/Master\'s in Business, Engineering, or related field. Strong analytical and communication skills. Experience in product development lifecycle. 3-7 years experience.',
                'skills': 'Product Strategy, Market Analysis, Project Management, Data Analysis, Stakeholder Management'
            },
            {
                'position': 'Cybersecurity Specialist',
                'salary_range': '45,000 - 100,000',
                'annual_range': '5.4L - 12L',
                'responsibilities': 'Monitor and protect systems from security threats. Conduct security assessments and audits. Implement security policies and procedures. Respond to security incidents.',
                'requirements': 'Bachelor\'s in Cybersecurity, IT, or related field. Security certifications (CISSP, CEH, CompTIA Security+). Knowledge of security tools and frameworks. 2-5 years experience.',
                'skills': 'Network Security, Ethical Hacking, Risk Assessment, Security Tools, Incident Response, Compliance'
            },
            {
                'position': 'Mobile App Developer',
                'salary_range': '35,000 - 85,000',
                'annual_range': '4.2L - 10.2L',
                'responsibilities': 'Develop mobile applications for iOS and Android platforms. Optimize app performance and user experience. Integrate APIs and third-party services. Test and debug applications.',
                'requirements': 'Bachelor\'s in Computer Science or related field. Experience with mobile development frameworks (React Native, Flutter, Swift, Kotlin). App store deployment experience.',
                'skills': 'Mobile Development, React Native/Flutter, iOS/Android Development, API Integration, App Store Optimization'
            },
            {
                'position': 'Cloud Architect',
                'salary_range': '60,000 - 180,000',
                'annual_range': '7.2L - 21.6L',
                'responsibilities': 'Design and implement cloud infrastructure solutions. Ensure scalability, security, and cost optimization. Lead cloud migration projects. Provide technical guidance to teams.',
                'requirements': 'Bachelor\'s/Master\'s in Engineering or Computer Science. Cloud certifications (AWS Solutions Architect, Azure Architect). 5+ years of cloud experience. Leadership skills.',
                'skills': 'Cloud Architecture, AWS/Azure/GCP, Infrastructure Design, Cost Optimization, Security, Leadership'
            }
        ],
        'healthcare': [
            {
                'position': 'General Physician',
                'salary_range': '40,000 - 80,000',
                'annual_range': '4.8L - 9.6L',
                'responsibilities': 'Diagnose and treat common illnesses. Conduct routine check-ups and preventive care. Refer patients to specialists when needed. Maintain patient records and medical histories.',
                'requirements': 'MBBS degree (5.5 years). Medical license from Nepal Medical Council. Internship completion. Strong diagnostic and communication skills.',
                'skills': 'Clinical Diagnosis, Patient Care, Medical Knowledge, Communication, Emergency Response, Medical Ethics'
            },
            {
                'position': 'Specialist Doctor',
                'salary_range': '60,000 - 150,000',
                'annual_range': '7.2L - 18L',
                'responsibilities': 'Provide specialized medical care in specific fields. Perform advanced diagnostic procedures. Consult on complex cases. Train junior doctors and medical students.',
                'requirements': 'MBBS + MD/MS specialization (3-6 years). Board certification in specialty. 5+ years clinical experience. Research and publication experience preferred.',
                'skills': 'Specialized Medical Knowledge, Advanced Procedures, Research, Teaching, Leadership, Critical Thinking'
            },
            {
                'position': 'Surgeon',
                'salary_range': '80,000 - 250,000',
                'annual_range': '9.6L - 30L',
                'responsibilities': 'Perform surgical procedures. Pre and post-operative patient care. Emergency surgery when required. Collaborate with surgical teams and anesthesiologists.',
                'requirements': 'MBBS + MS in Surgery (6 years total). Surgical residency completion. Board certification. Manual dexterity and precision. High stress tolerance.',
                'skills': 'Surgical Techniques, Anatomy, Decision Making Under Pressure, Manual Dexterity, Team Leadership, Patient Care'
            },
            {
                'position': 'Nurse',
                'salary_range': '25,000 - 45,000',
                'annual_range': '3L - 5.4L',
                'responsibilities': 'Provide direct patient care. Administer medications and treatments. Monitor patient vital signs. Assist doctors during procedures. Educate patients and families.',
                'requirements': 'Bachelor of Nursing (4 years) or Diploma in Nursing (3 years). Nursing license. CPR certification. Compassionate and detail-oriented personality.',
                'skills': 'Patient Care, Medical Procedures, Medication Administration, Communication, Empathy, Attention to Detail'
            },
            {
                'position': 'Pharmacist',
                'salary_range': '30,000 - 60,000',
                'annual_range': '3.6L - 7.2L',
                'responsibilities': 'Dispense medications and provide drug information. Counsel patients on proper medication use. Monitor drug interactions. Maintain pharmacy inventory and records.',
                'requirements': 'Bachelor of Pharmacy (4 years). Pharmacy license. Knowledge of drug interactions and pharmacology. Attention to detail and accuracy.',
                'skills': 'Pharmacology, Drug Interactions, Patient Counseling, Inventory Management, Regulatory Compliance, Attention to Detail'
            },
            {
                'position': 'Medical Technologist',
                'salary_range': '28,000 - 50,000',
                'annual_range': '3.36L - 6L',
                'responsibilities': 'Perform laboratory tests and analyses. Operate medical equipment and instruments. Prepare samples and specimens. Maintain quality control and safety standards.',
                'requirements': 'Bachelor in Medical Laboratory Technology or related field. Laboratory certification. Technical skills with medical equipment. Precision and analytical thinking.',
                'skills': 'Laboratory Techniques, Medical Equipment Operation, Quality Control, Data Analysis, Safety Protocols, Technical Skills'
            },
            {
                'position': 'Physiotherapist',
                'salary_range': '35,000 - 65,000',
                'annual_range': '4.2L - 7.8L',
                'responsibilities': 'Assess and treat patients with physical disabilities. Design rehabilitation programs. Provide manual therapy and exercise guidance. Monitor patient progress.',
                'requirements': 'Bachelor in Physiotherapy (4-5 years). License from physiotherapy council. Manual therapy skills. Good physical fitness and communication skills.',
                'skills': 'Manual Therapy, Exercise Prescription, Patient Assessment, Rehabilitation Planning, Communication, Physical Fitness'
            },
            {
                'position': 'Dentist',
                'salary_range': '45,000 - 120,000',
                'annual_range': '5.4L - 14.4L',
                'responsibilities': 'Diagnose and treat dental problems. Perform dental procedures and surgeries. Provide preventive dental care. Educate patients on oral hygiene.',
                'requirements': 'Bachelor of Dental Surgery (5 years). Dental license. Manual dexterity and precision. Good communication and patient management skills.',
                'skills': 'Dental Procedures, Oral Surgery, Patient Care, Manual Dexterity, Diagnostic Skills, Preventive Care'
            },
            {
                'position': 'Radiologist',
                'salary_range': '70,000 - 200,000',
                'annual_range': '8.4L - 24L',
                'responsibilities': 'Interpret medical images (X-rays, CT, MRI). Diagnose conditions through imaging. Perform image-guided procedures. Consult with other physicians on cases.',
                'requirements': 'MBBS + MD in Radiology (8-9 years total). Specialized training in imaging techniques. Strong analytical and diagnostic skills. Technology proficiency.',
                'skills': 'Medical Imaging, Diagnostic Analysis, Technology Proficiency, Attention to Detail, Consultation, Report Writing'
            },
            {
                'position': 'Anesthesiologist',
                'salary_range': '75,000 - 220,000',
                'annual_range': '9L - 26.4L',
                'responsibilities': 'Administer anesthesia during surgeries. Monitor patients during operations. Manage pain control and sedation. Handle emergency airway management.',
                'requirements': 'MBBS + MD in Anesthesiology (8-9 years). Board certification. ICU experience preferred. Strong decision-making under pressure.',
                'skills': 'Anesthesia Administration, Patient Monitoring, Emergency Response, Pain Management, Critical Care, Pharmacology'
            },
            {
                'position': 'Cardiologist',
                'salary_range': '90,000 - 300,000',
                'annual_range': '10.8L - 36L',
                'responsibilities': 'Diagnose and treat heart conditions. Perform cardiac procedures and interventions. Interpret cardiac tests and imaging. Provide preventive cardiac care.',
                'requirements': 'MBBS + DM in Cardiology (11+ years total). Fellowship training preferred. Board certification. Research experience valuable.',
                'skills': 'Cardiac Diagnosis, Interventional Procedures, ECG Interpretation, Cardiac Imaging, Patient Care, Research'
            },
            {
                'position': 'Neurologist',
                'salary_range': '85,000 - 280,000',
                'annual_range': '10.2L - 33.6L',
                'responsibilities': 'Diagnose and treat neurological disorders. Perform neurological examinations. Interpret brain imaging and tests. Manage chronic neurological conditions.',
                'requirements': 'MBBS + DM in Neurology (11+ years total). Board certification. Strong analytical and diagnostic skills. Research background preferred.',
                'skills': 'Neurological Diagnosis, Brain Imaging, Neurological Examination, Patient Management, Research, Critical Thinking'
            }
        ],
        'finance': [
            {
                'position': 'Financial Analyst',
                'salary_range': '35,000 - 70,000',
                'annual_range': '4.2L - 8.4L',
                'responsibilities': 'Analyze financial data and market trends. Prepare financial reports and forecasts. Evaluate investment opportunities. Support decision-making with data insights.',
                'requirements': 'Bachelor\'s in Finance, Economics, or Accounting. Strong analytical and mathematical skills. Proficiency in Excel and financial software. 1-3 years experience.',
                'skills': 'Financial Analysis, Excel Modeling, Data Analysis, Report Writing, Market Research, Financial Software'
            },
            {
                'position': 'Investment Banker',
                'salary_range': '50,000 - 150,000',
                'annual_range': '6L - 18L',
                'responsibilities': 'Facilitate mergers and acquisitions. Raise capital for companies. Provide financial advisory services. Conduct due diligence and valuations.',
                'requirements': 'Bachelor\'s/Master\'s in Finance or Business. Strong analytical and communication skills. High stress tolerance. 2-5 years experience in finance.',
                'skills': 'Financial Modeling, Valuation, M&A Analysis, Client Relations, Presentation Skills, Market Knowledge'
            },
            {
                'position': 'Accountant',
                'salary_range': '25,000 - 55,000',
                'annual_range': '3L - 6.6L',
                'responsibilities': 'Maintain financial records and books. Prepare tax returns and financial statements. Ensure compliance with regulations. Conduct audits and reconciliations.',
                'requirements': 'Bachelor\'s in Accounting or Finance. CA/ACCA certification preferred. Knowledge of accounting software. Attention to detail and accuracy.',
                'skills': 'Bookkeeping, Tax Preparation, Financial Statements, Audit, Compliance, Accounting Software, Attention to Detail'
            },
            {
                'position': 'Risk Manager',
                'salary_range': '45,000 - 100,000',
                'annual_range': '5.4L - 12L',
                'responsibilities': 'Identify and assess financial risks. Develop risk mitigation strategies. Monitor market and credit risks. Prepare risk reports for management.',
                'requirements': 'Bachelor\'s/Master\'s in Finance or Risk Management. Risk certifications (FRM, PRM). Strong analytical skills. 3-7 years experience.',
                'skills': 'Risk Assessment, Risk Modeling, Regulatory Knowledge, Data Analysis, Report Writing, Strategic Planning'
            },
            {
                'position': 'Credit Analyst',
                'salary_range': '30,000 - 65,000',
                'annual_range': '3.6L - 7.8L',
                'responsibilities': 'Evaluate creditworthiness of loan applicants. Analyze financial statements and credit histories. Recommend loan approvals or rejections. Monitor existing loan portfolios.',
                'requirements': 'Bachelor\'s in Finance, Economics, or Accounting. Strong analytical skills. Knowledge of credit analysis techniques. 1-4 years experience.',
                'skills': 'Credit Analysis, Financial Statement Analysis, Risk Assessment, Loan Processing, Regulatory Compliance, Decision Making'
            },
            {
                'position': 'Portfolio Manager',
                'salary_range': '60,000 - 180,000',
                'annual_range': '7.2L - 21.6L',
                'responsibilities': 'Manage investment portfolios for clients. Make investment decisions and asset allocation. Monitor market performance and adjust strategies. Provide investment advice and reports.',
                'requirements': 'Bachelor\'s/Master\'s in Finance or Economics. CFA certification preferred. Strong analytical skills. 5+ years investment experience.',
                'skills': 'Investment Analysis, Portfolio Management, Market Research, Risk Management, Client Relations, Financial Modeling'
            },
            {
                'position': 'Insurance Underwriter',
                'salary_range': '35,000 - 75,000',
                'annual_range': '4.2L - 9L',
                'responsibilities': 'Evaluate insurance applications and determine coverage terms. Assess risk factors and set premiums. Review policy renewals and claims. Ensure compliance with regulations.',
                'requirements': 'Bachelor\'s in Finance, Business, or related field. Insurance certifications preferred. Strong analytical and decision-making skills. 2-5 years experience.',
                'skills': 'Risk Assessment, Insurance Knowledge, Data Analysis, Decision Making, Regulatory Compliance, Attention to Detail'
            },
            {
                'position': 'Tax Consultant',
                'salary_range': '32,000 - 68,000',
                'annual_range': '3.84L - 8.16L',
                'responsibilities': 'Prepare tax returns for individuals and businesses. Provide tax planning advice. Ensure compliance with tax laws. Represent clients in tax matters.',
                'requirements': 'Bachelor\'s in Accounting or Finance. Tax certification preferred. Knowledge of tax laws and regulations. Strong attention to detail.',
                'skills': 'Tax Preparation, Tax Law Knowledge, Client Advisory, Compliance, Problem Solving, Communication'
            },
            {
                'position': 'Financial Planner',
                'salary_range': '38,000 - 85,000',
                'annual_range': '4.56L - 10.2L',
                'responsibilities': 'Develop comprehensive financial plans for clients. Provide investment and retirement planning advice. Analyze financial situations and goals. Monitor and adjust financial strategies.',
                'requirements': 'Bachelor\'s in Finance or related field. CFP certification preferred. Strong communication and analytical skills. 2-6 years experience.',
                'skills': 'Financial Planning, Investment Advisory, Retirement Planning, Client Relations, Risk Assessment, Communication'
            },
            {
                'position': 'Audit Manager',
                'salary_range': '45,000 - 95,000',
                'annual_range': '5.4L - 11.4L',
                'responsibilities': 'Lead audit teams and manage audit engagements. Review financial statements and internal controls. Ensure compliance with auditing standards. Communicate findings to clients.',
                'requirements': 'Bachelor\'s in Accounting. CA/CPA certification required. Strong leadership and analytical skills. 5+ years audit experience.',
                'skills': 'Audit Management, Team Leadership, Financial Analysis, Compliance, Risk Assessment, Communication'
            }
        ],
        'education': [
            {
                'position': 'School Teacher',
                'salary_range': '20,000 - 40,000',
                'annual_range': '2.4L - 4.8L',
                'responsibilities': 'Plan and deliver lessons to students. Assess student progress and provide feedback. Manage classroom behavior and create learning environment. Communicate with parents and colleagues.',
                'requirements': 'Bachelor\'s in Education or subject area. Teaching license required. Strong communication and patience. Classroom management skills.',
                'skills': 'Teaching, Classroom Management, Curriculum Development, Student Assessment, Communication, Patience'
            },
            {
                'position': 'University Professor',
                'salary_range': '45,000 - 120,000',
                'annual_range': '5.4L - 14.4L',
                'responsibilities': 'Conduct research and publish academic papers. Teach undergraduate and graduate courses. Supervise student research. Participate in academic committees and conferences.',
                'requirements': 'PhD in relevant field. Research experience and publications. Strong teaching and communication skills. Academic excellence.',
                'skills': 'Research, Teaching, Academic Writing, Critical Thinking, Public Speaking, Mentoring'
            },
            {
                'position': 'Principal',
                'salary_range': '35,000 - 80,000',
                'annual_range': '4.2L - 9.6L',
                'responsibilities': 'Lead and manage school operations. Supervise teachers and staff. Develop school policies and programs. Ensure educational standards and student welfare.',
                'requirements': 'Master\'s in Education or Administration. Teaching experience required. Strong leadership and management skills. Educational leadership certification.',
                'skills': 'Leadership, School Management, Policy Development, Staff Supervision, Educational Planning, Communication'
            },
            {
                'position': 'Educational Consultant',
                'salary_range': '40,000 - 90,000',
                'annual_range': '4.8L - 10.8L',
                'responsibilities': 'Provide educational guidance and advice to students and institutions. Develop educational programs and curricula. Conduct training workshops and seminars. Evaluate educational effectiveness.',
                'requirements': 'Master\'s in Education or related field. Teaching or educational administration experience. Strong communication and analytical skills. Knowledge of educational trends.',
                'skills': 'Educational Planning, Curriculum Design, Training, Consultation, Communication, Program Evaluation'
            },
            {
                'position': 'Training Coordinator',
                'salary_range': '30,000 - 60,000',
                'annual_range': '3.6L - 7.2L',
                'responsibilities': 'Plan and organize training programs. Coordinate with trainers and participants. Manage training schedules and resources. Evaluate training effectiveness and outcomes.',
                'requirements': 'Bachelor\'s in Education, HR, or related field. Training and coordination experience. Strong organizational and communication skills. Project management abilities.',
                'skills': 'Training Coordination, Project Management, Communication, Scheduling, Resource Management, Evaluation'
            },
            {
                'position': 'Curriculum Developer',
                'salary_range': '35,000 - 75,000',
                'annual_range': '4.2L - 9L',
                'responsibilities': 'Design and develop educational curricula and materials. Research educational best practices. Create learning objectives and assessments. Review and update existing curricula.',
                'requirements': 'Master\'s in Education or subject expertise. Curriculum development experience. Strong research and writing skills. Knowledge of learning theories.',
                'skills': 'Curriculum Design, Educational Research, Content Development, Learning Assessment, Writing, Educational Technology'
            },
            {
                'position': 'Academic Administrator',
                'salary_range': '40,000 - 85,000',
                'annual_range': '4.8L - 10.2L',
                'responsibilities': 'Manage academic programs and operations. Oversee student services and academic policies. Coordinate with faculty and staff. Ensure compliance with educational standards.',
                'requirements': 'Master\'s in Education or Administration. Academic or administrative experience. Strong leadership and management skills. Knowledge of educational regulations.',
                'skills': 'Academic Management, Policy Development, Leadership, Student Services, Compliance, Communication'
            },
            {
                'position': 'Research Coordinator',
                'salary_range': '38,000 - 78,000',
                'annual_range': '4.56L - 9.36L',
                'responsibilities': 'Coordinate research projects and activities. Manage research data and documentation. Facilitate collaboration between researchers. Ensure compliance with research protocols.',
                'requirements': 'Master\'s in relevant field. Research experience and methodology knowledge. Strong analytical and organizational skills. Data management abilities.',
                'skills': 'Research Coordination, Data Management, Project Management, Analysis, Documentation, Collaboration'
            }
        ],
        'engineering': [
            {
                'position': 'Civil Engineer',
                'salary_range': '35,000 - 75,000',
                'annual_range': '4.2L - 9L',
                'responsibilities': 'Design and oversee construction of infrastructure projects. Conduct site inspections and surveys. Prepare technical drawings and specifications. Ensure compliance with building codes.',
                'requirements': 'Bachelor\'s in Civil Engineering. Professional Engineer (PE) license preferred. Strong technical and mathematical skills. Knowledge of construction materials and methods.',
                'skills': 'Structural Design, Project Management, CAD Software, Construction Knowledge, Problem Solving, Technical Analysis'
            },
            {
                'position': 'Mechanical Engineer',
                'salary_range': '40,000 - 80,000',
                'annual_range': '4.8L - 9.6L',
                'responsibilities': 'Design and develop mechanical systems and products. Conduct testing and analysis of mechanical components. Oversee manufacturing processes. Troubleshoot mechanical problems.',
                'requirements': 'Bachelor\'s in Mechanical Engineering. Strong analytical and problem-solving skills. Knowledge of manufacturing processes. CAD software proficiency.',
                'skills': 'Mechanical Design, CAD/CAM, Manufacturing, Testing, Problem Solving, Project Management'
            },
            {
                'position': 'Electrical Engineer',
                'salary_range': '38,000 - 85,000',
                'annual_range': '4.56L - 10.2L',
                'responsibilities': 'Design electrical systems and components. Develop power distribution systems. Test electrical equipment and systems. Ensure electrical safety and compliance.',
                'requirements': 'Bachelor\'s in Electrical Engineering. Knowledge of electrical codes and standards. Strong analytical skills. Experience with electrical design software.',
                'skills': 'Electrical Design, Power Systems, Circuit Analysis, Testing, Safety Standards, Technical Documentation'
            },
            {'position': 'Project Manager', 'salary_range': '50,000 - 120,000', 'annual_range': '6L - 14.4L'},
            {'position': 'Quality Assurance Engineer', 'salary_range': '35,000 - 70,000', 'annual_range': '4.2L - 8.4L'},
            {'position': 'Structural Engineer', 'salary_range': '42,000 - 88,000', 'annual_range': '5.04L - 10.56L'},
            {'position': 'Environmental Engineer', 'salary_range': '38,000 - 78,000', 'annual_range': '4.56L - 9.36L'},
            {'position': 'Construction Manager', 'salary_range': '45,000 - 95,000', 'annual_range': '5.4L - 11.4L'},
            {'position': 'Site Engineer', 'salary_range': '32,000 - 65,000', 'annual_range': '3.84L - 7.8L'},
        ],
        'marketing': [
            {'position': 'Digital Marketing Manager', 'salary_range': '40,000 - 90,000', 'annual_range': '4.8L - 10.8L'},
            {'position': 'Content Creator', 'salary_range': '25,000 - 55,000', 'annual_range': '3L - 6.6L'},
            {'position': 'Brand Manager', 'salary_range': '45,000 - 100,000', 'annual_range': '5.4L - 12L'},
            {'position': 'SEO Specialist', 'salary_range': '30,000 - 65,000', 'annual_range': '3.6L - 7.8L'},
            {'position': 'Social Media Manager', 'salary_range': '28,000 - 60,000', 'annual_range': '3.36L - 7.2L'},
            {'position': 'Marketing Analyst', 'salary_range': '35,000 - 72,000', 'annual_range': '4.2L - 8.64L'},
            {'position': 'Public Relations Manager', 'salary_range': '38,000 - 82,000', 'annual_range': '4.56L - 9.84L'},
            {'position': 'Event Manager', 'salary_range': '32,000 - 68,000', 'annual_range': '3.84L - 8.16L'},
        ],
        'government': [
            {'position': 'Civil Service Officer', 'salary_range': '30,000 - 70,000', 'annual_range': '3.6L - 8.4L'},
            {'position': 'Police Officer', 'salary_range': '25,000 - 50,000', 'annual_range': '3L - 6L'},
            {'position': 'Government Administrator', 'salary_range': '35,000 - 80,000', 'annual_range': '4.2L - 9.6L'},
            {'position': 'Public Policy Analyst', 'salary_range': '40,000 - 90,000', 'annual_range': '4.8L - 10.8L'},
            {'position': 'Tax Officer', 'salary_range': '32,000 - 65,000', 'annual_range': '3.84L - 7.8L'},
            {'position': 'Customs Officer', 'salary_range': '35,000 - 72,000', 'annual_range': '4.2L - 8.64L'},
            {'position': 'Foreign Service Officer', 'salary_range': '45,000 - 95,000', 'annual_range': '5.4L - 11.4L'},
        ],
        'hospitality': [
            {'position': 'Hotel Manager', 'salary_range': '35,000 - 80,000', 'annual_range': '4.2L - 9.6L'},
            {'position': 'Chef', 'salary_range': '25,000 - 65,000', 'annual_range': '3L - 7.8L'},
            {'position': 'Tour Guide', 'salary_range': '20,000 - 45,000', 'annual_range': '2.4L - 5.4L'},
            {'position': 'Event Coordinator', 'salary_range': '28,000 - 58,000', 'annual_range': '3.36L - 6.96L'},
            {'position': 'Restaurant Manager', 'salary_range': '30,000 - 70,000', 'annual_range': '3.6L - 8.4L'},
            {'position': 'Travel Consultant', 'salary_range': '25,000 - 55,000', 'annual_range': '3L - 6.6L'},
            {'position': 'Travel Agent', 'salary_range': '20,000 - 40,000', 'annual_range': '2.4L - 4.8L'},
            {'position': 'Event Coordinator', 'salary_range': '30,000 - 55,000', 'annual_range': '3.6L - 6.6L'},
            {'position': 'Chef (Head)', 'salary_range': '40,000 - 70,000', 'annual_range': '4.8L - 8.4L'},
            {'position': 'Front Desk Manager', 'salary_range': '25,000 - 45,000', 'annual_range': '3L - 5.4L'},
            {'position': 'Tourism Marketing Specialist', 'salary_range': '35,000 - 60,000', 'annual_range': '4.2L - 7.2L'},
        ],
        'media_communications': [
            {'position': 'Journalist', 'salary_range': '25,000 - 50,000', 'annual_range': '3L - 6L'},
            {'position': 'News Anchor', 'salary_range': '60,000 - 120,000', 'annual_range': '7.2L - 14.4L'},
            {'position': 'Video Editor', 'salary_range': '30,000 - 55,000', 'annual_range': '3.6L - 6.6L'},
            {'position': 'Graphic Designer', 'salary_range': '25,000 - 45,000', 'annual_range': '3L - 5.4L'},
            {'position': 'Social Media Manager', 'salary_range': '30,000 - 60,000', 'annual_range': '3.6L - 7.2L'},
            {'position': 'Content Writer', 'salary_range': '20,000 - 40,000', 'annual_range': '2.4L - 4.8L'},
            {'position': 'Radio Jockey', 'salary_range': '25,000 - 50,000', 'annual_range': '3L - 6L'},
            {'position': 'Public Relations Manager', 'salary_range': '45,000 - 80,000', 'annual_range': '5.4L - 9.6L'},
        ],
        'agriculture_food': [
            {'position': 'Agricultural Engineer', 'salary_range': '35,000 - 65,000', 'annual_range': '4.2L - 7.8L'},
            {'position': 'Farm Manager', 'salary_range': '30,000 - 55,000', 'annual_range': '3.6L - 6.6L'},
            {'position': 'Food Technologist', 'salary_range': '40,000 - 70,000', 'annual_range': '4.8L - 8.4L'},
            {'position': 'Veterinarian', 'salary_range': '45,000 - 85,000', 'annual_range': '5.4L - 10.2L'},
            {'position': 'Agricultural Consultant', 'salary_range': '35,000 - 60,000', 'annual_range': '4.2L - 7.2L'},
            {'position': 'Food Quality Inspector', 'salary_range': '25,000 - 45,000', 'annual_range': '3L - 5.4L'},
            {'position': 'Livestock Specialist', 'salary_range': '30,000 - 50,000', 'annual_range': '3.6L - 6L'},
        ],
        'retail_sales': [
            {'position': 'Store Manager', 'salary_range': '35,000 - 65,000', 'annual_range': '4.2L - 7.8L'},
            {'position': 'Sales Executive', 'salary_range': '20,000 - 40,000', 'annual_range': '2.4L - 4.8L'},
            {'position': 'Regional Sales Manager', 'salary_range': '60,000 - 100,000', 'annual_range': '7.2L - 12L'},
            {'position': 'Retail Supervisor', 'salary_range': '25,000 - 45,000', 'annual_range': '3L - 5.4L'},
            {'position': 'Customer Service Manager', 'salary_range': '30,000 - 55,000', 'annual_range': '3.6L - 6.6L'},
            {'position': 'Merchandiser', 'salary_range': '20,000 - 35,000', 'annual_range': '2.4L - 4.2L'},
            {'position': 'E-commerce Manager', 'salary_range': '45,000 - 80,000', 'annual_range': '5.4L - 9.6L'},
        ],
        'manufacturing_production': [
            {'position': 'Production Manager', 'salary_range': '55,000 - 95,000', 'annual_range': '6.6L - 11.4L'},
            {'position': 'Quality Control Manager', 'salary_range': '45,000 - 75,000', 'annual_range': '5.4L - 9L'},
            {'position': 'Factory Supervisor', 'salary_range': '35,000 - 60,000', 'annual_range': '4.2L - 7.2L'},
            {'position': 'Industrial Engineer', 'salary_range': '40,000 - 70,000', 'annual_range': '4.8L - 8.4L'},
            {'position': 'Machine Operator', 'salary_range': '18,000 - 30,000', 'annual_range': '2.16L - 3.6L'},
            {'position': 'Safety Officer', 'salary_range': '30,000 - 50,000', 'annual_range': '3.6L - 6L'},
            {'position': 'Supply Chain Manager', 'salary_range': '50,000 - 85,000', 'annual_range': '6L - 10.2L'},
        ],
        'transportation_logistics': [
            {'position': 'Logistics Manager', 'salary_range': '45,000 - 80,000', 'annual_range': '5.4L - 9.6L'},
            {'position': 'Truck Driver', 'salary_range': '20,000 - 35,000', 'annual_range': '2.4L - 4.2L'},
            {'position': 'Warehouse Manager', 'salary_range': '35,000 - 65,000', 'annual_range': '4.2L - 7.8L'},
            {'position': 'Fleet Manager', 'salary_range': '40,000 - 70,000', 'annual_range': '4.8L - 8.4L'},
            {'position': 'Shipping Coordinator', 'salary_range': '25,000 - 45,000', 'annual_range': '3L - 5.4L'},
            {'position': 'Customs Officer', 'salary_range': '30,000 - 55,000', 'annual_range': '3.6L - 6.6L'},
        ],
        'construction_real_estate': [
            {'position': 'Construction Manager', 'salary_range': '60,000 - 110,000', 'annual_range': '7.2L - 13.2L'},
            {'position': 'Real Estate Agent', 'salary_range': '25,000 - 60,000', 'annual_range': '3L - 7.2L'},
            {'position': 'Architect', 'salary_range': '45,000 - 85,000', 'annual_range': '5.4L - 10.2L'},
            {'position': 'Site Engineer', 'salary_range': '30,000 - 55,000', 'annual_range': '3.6L - 6.6L'},
            {'position': 'Property Manager', 'salary_range': '35,000 - 65,000', 'annual_range': '4.2L - 7.8L'},
            {'position': 'Surveyor', 'salary_range': '25,000 - 45,000', 'annual_range': '3L - 5.4L'},
            {'position': 'Interior Designer', 'salary_range': '30,000 - 60,000', 'annual_range': '3.6L - 7.2L'},
        ],
    }
    
    # Comprehensive job insights with detailed career factors, advantages, and disadvantages
    job_insights = {
        'it_tech': {
            'growth_rate': 'High (15-20% annually)',
            'job_security': 'High - Technology skills are in constant demand across all industries',
            'work_life_balance': 'Moderate to Good - Flexible hours but project deadlines can be intense',
            'remote_work': 'Excellent - Most IT roles can be performed remotely with global opportunities',
            'skill_demand': 'Very High - Digital transformation drives continuous demand for tech skills',
            'career_progression': 'Fast - Merit-based advancement with opportunities to specialize or lead',
            'education_requirements': 'Bachelor\'s in Computer Science/IT or equivalent experience and certifications',
            'entry_barriers': 'Moderate - Requires technical aptitude and continuous learning mindset',
            'market_outlook': 'Excellent - Nepal\'s growing IT sector and global outsourcing opportunities',
            'advantages': [
                'High salary potential with performance-based increments and bonuses',
                'Remote work opportunities enabling global job market access',
                'Continuous learning keeps work intellectually stimulating',
                'Creative problem-solving and innovation opportunities',
                'Flexible working hours and modern work environments',
                'Strong job security due to high demand for technical skills',
                'Opportunity to work on cutting-edge technologies and projects',
                'Potential for freelancing and consulting income',
                'Fast career progression based on merit and skills',
                'Global networking opportunities through online communities'
            ],
            'disadvantages': [
                'High stress during project deadlines and system outages',
                'Constant need for upskilling to stay relevant with technology changes',
                'Long working hours especially during product launches',
                'Sedentary lifestyle leading to health issues if not managed',
                'Rapid technology changes can make skills obsolete quickly',
                'Imposter syndrome common due to vast knowledge requirements',
                'Eye strain and repetitive stress injuries from computer work',
                'Work-life balance challenges during critical project phases',
                'Burnout risk from high-pressure environments',
                'Communication challenges when working with non-technical stakeholders'
            ],
            'career_factors': {
                'skills_needed': 'Programming languages, problem-solving, logical thinking, continuous learning',
                'certifications': 'AWS, Google Cloud, Microsoft Azure, Cisco, CompTIA, specific technology certifications',
                'experience_levels': 'Entry (0-2 years), Mid (3-5 years), Senior (6-10 years), Lead/Architect (10+ years)',
                'specialization_paths': 'Full-stack development, DevOps, Data Science, Cybersecurity, Cloud Architecture',
                'industry_applications': 'All industries - Banking, Healthcare, E-commerce, Government, Startups'
            }
        },
        'healthcare': {
            'growth_rate': 'Stable (8-12% annually) - Consistent demand driven by aging population and health awareness',
            'job_security': 'Very High - Essential services with guaranteed demand regardless of economic conditions',
            'work_life_balance': 'Challenging - Irregular hours, on-call duties, and emergency situations',
            'remote_work': 'Limited - Most roles require physical presence for patient care',
            'skill_demand': 'High - Specialized medical knowledge and continuous professional development required',
            'career_progression': 'Steady - Clear hierarchy from junior to senior positions with specialization opportunities',
            'education_requirements': 'Extensive - MBBS (5.5 years), specialized degrees (3-6 additional years), continuous education',
            'entry_barriers': 'High - Rigorous education, licensing exams, and significant financial investment',
            'market_outlook': 'Excellent - Growing healthcare infrastructure and medical tourism in Nepal',
            'advantages': [
                'Exceptional job security with recession-proof career prospects',
                'Profound social impact and personal fulfillment from saving lives',
                'High social respect and professional prestige in community',
                'Excellent earning potential, especially for specialists and surgeons',
                'Continuous intellectual challenge and lifelong learning opportunities',
                'Comprehensive benefits including medical insurance and retirement plans',
                'Opportunity to work in diverse settings (hospitals, clinics, research)',
                'Global mobility with medical qualifications recognized worldwide',
                'Potential for private practice and entrepreneurial opportunities',
                'Strong professional networks and mentorship opportunities'
            ],
            'disadvantages': [
                'Extremely long and expensive education pathway (10+ years)',
                'High stress and emotional burden from patient care responsibilities',
                'Irregular working hours including nights, weekends, and holidays',
                'Risk of exposure to infectious diseases and workplace injuries',
                'Malpractice liability and legal risks from medical decisions',
                'Emotional trauma from patient deaths and difficult cases',
                'Limited work-life balance during residency and early career',
                'Continuous pressure to stay updated with medical advances',
                'Physical demands of long surgeries and standing for hours',
                'Bureaucratic challenges in healthcare systems and insurance'
            ],
            'career_factors': {
                'skills_needed': 'Medical knowledge, empathy, decision-making under pressure, communication, manual dexterity',
                'certifications': 'Medical license, board certifications, CPR, specialized procedure certifications',
                'experience_levels': 'Intern (1 year), Resident (3-6 years), Attending (licensed), Specialist (post-residency)',
                'specialization_paths': 'Internal Medicine, Surgery, Pediatrics, Cardiology, Neurology, Emergency Medicine',
                'industry_applications': 'Hospitals, private clinics, research institutions, pharmaceutical companies, telemedicine'
            }
        },
        'finance': {
            'growth_rate': 'Moderate (10-15% annually) - Steady growth driven by economic development and financial inclusion',
            'job_security': 'High - Essential for all businesses and economic stability',
            'work_life_balance': 'Moderate - Standard hours but intense during reporting periods and market volatility',
            'remote_work': 'Good - Many analytical and advisory roles can be performed remotely',
            'skill_demand': 'High - Strong analytical, mathematical, and regulatory knowledge required',
            'career_progression': 'Good - Clear advancement path from analyst to senior management roles',
            'education_requirements': 'Bachelor\'s in Finance/Economics/Accounting, professional certifications (CA, CFA, FRM)',
            'entry_barriers': 'Moderate - Requires strong analytical skills and understanding of financial markets',
            'market_outlook': 'Strong - Nepal\'s growing banking sector and capital market development',
            'advantages': [
                'Excellent earning potential with performance bonuses and incentives',
                'Clear career advancement opportunities to senior management',
                'Development of valuable analytical and quantitative skills',
                'Deep understanding of business operations and strategy',
                'Extensive networking opportunities with business leaders',
                'Transferable skills across industries and global markets',
                'Intellectual challenge of financial modeling and analysis',
                'Opportunity to influence major business decisions',
                'Professional recognition through certifications (CA, CFA)',
                'Potential for consulting and advisory income streams'
            ],
            'disadvantages': [
                'High-pressure environment with significant financial responsibility',
                'Stress from market volatility and economic uncertainty',
                'Long working hours during financial reporting and audit periods',
                'Complex regulatory compliance requirements and constant changes',
                'Career dependent on economic cycles and market conditions',
                'High accountability for financial accuracy and risk management',
                'Continuous learning required for evolving regulations',
                'Potential ethical dilemmas in investment and lending decisions',
                'Work can be highly quantitative with limited creativity',
                'Blame and scrutiny during financial losses or market downturns'
            ],
            'career_factors': {
                'skills_needed': 'Financial analysis, Excel modeling, risk assessment, regulatory knowledge, communication',
                'certifications': 'CA (Chartered Accountant), CFA (Chartered Financial Analyst), FRM (Financial Risk Manager)',
                'experience_levels': 'Analyst (0-3 years), Associate (3-6 years), Manager (6-10 years), Director (10+ years)',
                'specialization_paths': 'Investment Banking, Risk Management, Portfolio Management, Corporate Finance, Audit',
                'industry_applications': 'Banks, Insurance, Investment firms, Corporations, Government, Consulting'
            }
        },
        'education': {
            'growth_rate': 'Stable (5-8% annually) - Consistent demand driven by population growth and education reforms',
            'job_security': 'High - Essential service with government backing and tenure opportunities',
            'work_life_balance': 'Good - Regular hours, holidays, and vacation time align with academic calendar',
            'remote_work': 'Moderate - Online teaching expanded but hands-on instruction still preferred',
            'skill_demand': 'Moderate - Subject expertise and pedagogical skills with technology integration',
            'career_progression': 'Slow but Steady - Clear advancement through experience and additional qualifications',
            'education_requirements': 'Bachelor\'s degree minimum, Master\'s preferred, teaching certification required',
            'entry_barriers': 'Low to Moderate - Teaching aptitude and subject knowledge required',
            'market_outlook': 'Stable - Growing emphasis on quality education and skill development in Nepal',
            'advantages': [
                'Profound job satisfaction from shaping future generations',
                'Excellent work-life balance with regular hours and holidays',
                'Long vacation periods during school breaks and festivals',
                'Job security with tenure opportunities in government schools',
                'Intellectual stimulation through continuous learning and curriculum development',
                'Respected position in society with community recognition',
                'Pension and retirement benefits in government positions',
                'Opportunity to make lasting impact on students\' lives',
                'Collaborative work environment with fellow educators',
                'Potential for additional income through tutoring and coaching'
            ],
            'disadvantages': [
                'Lower salary potential compared to private sector careers',
                'Limited career mobility outside education sector',
                'Bureaucratic constraints in curriculum and administrative decisions',
                'Resource limitations affecting teaching effectiveness',
                'Slow salary growth and limited performance-based incentives',
                'Emotional stress from difficult students and parents',
                'Administrative burden reducing actual teaching time',
                'Outdated infrastructure in many educational institutions',
                'Political interference in educational policies and decisions',
                'Limited professional development opportunities in remote areas'
            ],
            'career_factors': {
                'skills_needed': 'Subject expertise, communication, patience, classroom management, technology integration',
                'certifications': 'Teaching license, subject-specific certifications, technology training',
                'experience_levels': 'Assistant Teacher (0-3 years), Teacher (3-8 years), Senior Teacher (8-15 years), Principal (15+ years)',
                'specialization_paths': 'Subject specialization, Educational administration, Curriculum development, Special education',
                'industry_applications': 'Schools, colleges, universities, training institutes, educational technology companies'
            }
        },
        'engineering': {
            'growth_rate': 'Moderate (8-12% annually) - Driven by infrastructure development and industrialization',
            'job_security': 'High - Essential for infrastructure and development projects',
            'work_life_balance': 'Good - Standard hours but project deadlines can be demanding',
            'remote_work': 'Moderate - Design work possible remotely but site supervision required',
            'skill_demand': 'High - Technical expertise and problem-solving skills in high demand',
            'career_progression': 'Steady - Clear path from junior engineer to project manager and consultant',
            'education_requirements': 'Bachelor\'s in Engineering, Professional Engineer (PE) license beneficial',
            'entry_barriers': 'Moderate - Strong technical and mathematical foundation required',
            'market_outlook': 'Strong - Nepal\'s infrastructure development and reconstruction projects',
            'advantages': [
                'Excellent problem-solving and analytical skill development',
                'Direct impact on infrastructure development and public welfare',
                'High technical expertise recognition and professional respect',
                'Diverse project variety from buildings to bridges to systems',
                'Good salary progression with experience and specialization',
                'Opportunity to work on landmark projects and infrastructure',
                'Strong job security due to essential nature of engineering',
                'Potential for consulting and independent practice',
                'Global mobility with engineering qualifications',
                'Intellectual satisfaction from creating tangible solutions'
            ],
            'disadvantages': [
                'High pressure from project deadlines and budget constraints',
                'Frequent site work requirements in challenging locations',
                'Weather-dependent outdoor work affecting project schedules',
                'Safety concerns and liability for structural integrity',
                'Complex regulatory compliance and approval processes',
                'Long project cycles with delayed gratification',
                'Physical demands of site inspections and supervision',
                'Responsibility for public safety in infrastructure projects',
                'Economic sensitivity affecting project funding and employment',
                'Continuous learning required for new technologies and codes'
            ],
            'career_factors': {
                'skills_needed': 'Technical analysis, CAD software, project management, problem-solving, attention to detail',
                'certifications': 'Professional Engineer (PE) license, project management (PMP), specialized software certifications',
                'experience_levels': 'Junior Engineer (0-3 years), Engineer (3-7 years), Senior Engineer (7-12 years), Principal (12+ years)',
                'specialization_paths': 'Civil, Mechanical, Electrical, Structural, Environmental, Project Management',
                'industry_applications': 'Construction, Infrastructure, Manufacturing, Energy, Government, Consulting'
            }
        },
        'marketing': {
            'growth_rate': 'High (12-18% annually) - Digital transformation and e-commerce growth driving demand',
            'job_security': 'Moderate - Performance-based with high demand for digital marketing skills',
            'work_life_balance': 'Moderate - Flexible but campaign deadlines and social media monitoring required',
            'remote_work': 'Excellent - Most marketing activities can be performed remotely with digital tools',
            'skill_demand': 'High - Digital marketing, analytics, and creative skills in high demand',
            'career_progression': 'Fast - Merit-based advancement with opportunities to specialize or lead teams',
            'education_requirements': 'Bachelor\'s in Marketing/Business/Communications, digital certifications preferred',
            'entry_barriers': 'Low to Moderate - Creativity and analytical thinking more important than formal qualifications',
            'market_outlook': 'Excellent - Digital marketing boom and growing e-commerce in Nepal',
            'advantages': [
                'Highly creative and dynamic work environment with variety',
                'Excellent remote work flexibility and global opportunities',
                'Extensive networking opportunities with diverse industries',
                'Exposure to diverse projects and business challenges',
                'Rapid development of valuable digital and analytical skills',
                'Fast career progression based on measurable results',
                'High earning potential with performance bonuses',
                'Opportunity to build personal brand and thought leadership',
                'Creative freedom in campaign development and execution',
                'Direct impact on business growth and revenue generation'
            ],
            'disadvantages': [
                'High pressure from campaign deadlines and launch schedules',
                'Results-driven environment with constant performance scrutiny',
                'Dependency on market trends and consumer behavior changes',
                'Stressful client management and expectation handling',
                'Irregular work hours during campaigns and product launches',
                'Constant need to stay updated with digital platform changes',
                'Budget constraints limiting creative execution',
                'Blame for poor campaign performance or market reception',
                'Burnout from always-on social media and digital monitoring',
                'Difficulty measuring ROI for brand and awareness campaigns'
            ],
            'career_factors': {
                'skills_needed': 'Creativity, analytics, digital tools, communication, strategic thinking, data interpretation',
                'certifications': 'Google Ads, Facebook Marketing, HubSpot, Google Analytics, content marketing certifications',
                'experience_levels': 'Coordinator (0-2 years), Specialist (2-5 years), Manager (5-8 years), Director (8+ years)',
                'specialization_paths': 'Digital Marketing, Content Marketing, Brand Management, SEO/SEM, Social Media Marketing',
                'industry_applications': 'All industries - FMCG, Technology, Healthcare, Education, E-commerce, Agencies'
            }
        },
        'government': {
            'growth_rate': 'Stable (3-6% annually) - Consistent with government expansion and public service needs',
            'job_security': 'Very High - Permanent positions with constitutional protection and tenure',
            'work_life_balance': 'Excellent - Fixed hours, weekends off, and generous leave policies',
            'remote_work': 'Limited - Most positions require physical presence for public service delivery',
            'skill_demand': 'Moderate - Administrative skills and subject expertise for specialized positions',
            'career_progression': 'Slow but Guaranteed - Structured promotion system based on seniority and performance',
            'education_requirements': 'Varies by position - Bachelor\'s for officer level, Master\'s for senior positions',
            'entry_barriers': 'High - Competitive exams (PSC) and merit-based selection process',
            'market_outlook': 'Stable - Consistent demand for public administration and specialized services',
            'advantages': [
                'Exceptional job security with constitutional protection',
                'Comprehensive benefits including health insurance and housing allowances',
                'Excellent work-life balance with fixed hours and holidays',
                'Generous pension plans and retirement security',
                'Meaningful public service impact on society and development',
                'Respect and social status in community',
                'Protection from arbitrary dismissal',
                'Opportunities for training and capacity building',
                'Transfer opportunities across different departments',
                'Additional allowances for remote and difficult postings'
            ],
            'disadvantages': [
                'Lower salary potential compared to private sector',
                'Slow and bureaucratic decision-making processes',
                'Limited scope for innovation and entrepreneurial thinking',
                'Political influences affecting transfers and promotions',
                'Extremely slow career advancement and promotion cycles',
                'Rigid hierarchical structure limiting flexibility',
                'Public scrutiny and accountability pressure',
                'Corruption allegations and ethical dilemmas',
                'Limited performance-based incentives',
                'Outdated systems and resistance to technological change'
            ],
            'career_factors': {
                'skills_needed': 'Administrative skills, policy analysis, public communication, regulatory knowledge, leadership',
                'certifications': 'Civil service qualifications, specialized training programs, language proficiency',
                'experience_levels': 'Officer (0-5 years), Section Officer (5-12 years), Under Secretary (12-20 years), Secretary (20+ years)',
                'specialization_paths': 'Administration, Finance, Development, Education, Health, Foreign Affairs, Legal',
                'industry_applications': 'Ministries, Departments, Local Government, Public Enterprises, Regulatory Bodies'
            }
        },
        'hospitality': {
            'growth_rate': 'High (10-15% annually) - Tourism recovery and growing domestic travel market',
            'job_security': 'Moderate - Seasonal variations and economic sensitivity affect stability',
            'work_life_balance': 'Challenging - Irregular hours, weekend work, and peak season demands',
            'remote_work': 'Limited - Service industry requiring physical presence and customer interaction',
            'skill_demand': 'Moderate - Customer service, language skills, and cultural awareness important',
            'career_progression': 'Fast - Merit-based advancement with opportunities to move up quickly',
            'education_requirements': 'Diploma/Certificate in hospitality preferred, language skills valuable',
            'entry_barriers': 'Low - Entry-level positions available with on-the-job training',
            'market_outlook': 'Strong - Nepal\'s tourism potential and growing hospitality infrastructure',
            'advantages': [
                'Rich customer interaction and relationship building opportunities',
                'Extensive cultural exposure and international guest interaction',
                'Additional income through tips and service charges',
                'International career opportunities and global mobility',
                'Fast career growth with performance-based advancement',
                'Development of valuable soft skills and communication abilities',
                'Networking opportunities with diverse clientele',
                'Creative opportunities in event planning and guest services',
                'Language skill development through international exposure',
                'Entrepreneurial opportunities in tourism and hospitality ventures'
            ],
            'disadvantages': [
                'Irregular and demanding working hours including nights and weekends',
                'Seasonal employment fluctuations affecting income stability',
                'High physical demands and long hours on feet',
                'Stressful customer service situations and complaint handling',
                'Lower base salaries with income dependent on tips and seasons',
                'Limited career advancement without specialized training',
                'Emotional labor dealing with difficult guests and situations',
                'High turnover rates and job instability',
                'Health impacts from irregular schedules and physical demands',
                'Economic sensitivity affecting tourism and hospitality demand'
            ],
            'career_factors': {
                'skills_needed': 'Customer service, communication, multitasking, cultural sensitivity, problem-solving, languages',
                'certifications': 'Hospitality management, food safety, first aid, language certifications, tourism guide license',
                'experience_levels': 'Staff (0-2 years), Supervisor (2-5 years), Manager (5-10 years), General Manager (10+ years)',
                'specialization_paths': 'Hotel Management, Restaurant Operations, Event Management, Tourism, Travel Operations',
                'industry_applications': 'Hotels, Restaurants, Resorts, Travel Agencies, Event Companies, Airlines, Tourism Boards'
            }
        },
        'media_communications': {
            'growth_rate': 'High (12-15% annually) - Digital media expansion and content creation boom',
            'job_security': 'Moderate - Project-based work with growing demand for digital content',
            'work_life_balance': 'Moderate - Flexible but deadline-driven with irregular schedules',
            'remote_work': 'Good - Many roles can be performed remotely with digital tools',
            'skill_demand': 'High - Digital literacy, content creation, and multimedia skills essential',
            'career_progression': 'Fast - Merit-based with opportunities for specialization and freelancing',
            'education_requirements': 'Bachelor\'s in Communications/Journalism/Media Studies, portfolio important',
            'entry_barriers': 'Moderate - Strong communication skills and portfolio required',
            'market_outlook': 'Excellent - Digital transformation and social media growth in Nepal',
            'advantages': [
                'High creative expression and artistic freedom in content creation',
                'Significant public influence and ability to shape opinions',
                'Extensive networking opportunities across industries and sectors',
                'Diverse content creation from news to entertainment to marketing',
                'Rapid digital platform growth creating new opportunities',
                'Flexibility to work as freelancer or independent creator',
                'Opportunity to build personal brand and thought leadership',
                'Access to events, interviews, and exclusive information',
                'Potential for viral content and rapid career advancement',
                'Multiple revenue streams through various platforms and clients'
            ],
            'disadvantages': [
                'Intense deadline pressure and time-sensitive content requirements',
                'Constant public scrutiny and criticism of work and opinions',
                'Irregular income with feast-or-famine cycles',
                'Continuous need for technology adaptation and platform changes',
                'High market competition and oversaturation in some areas',
                'Emotional stress from negative feedback and public criticism',
                'Long hours during breaking news or major events',
                'Ethical dilemmas and pressure for sensational content',
                'Job insecurity with project-based and contract work',
                'Health impacts from screen time and irregular schedules'
            ],
            'career_factors': {
                'skills_needed': 'Writing, video editing, social media, research, interviewing, digital marketing, storytelling',
                'certifications': 'Digital marketing, video production, journalism ethics, social media management, Adobe Creative Suite',
                'experience_levels': 'Intern/Assistant (0-2 years), Reporter/Creator (2-5 years), Senior (5-10 years), Editor/Director (10+ years)',
                'specialization_paths': 'Journalism, Content Creation, Digital Marketing, Broadcasting, Public Relations, Documentary',
                'industry_applications': 'News Media, Digital Agencies, Broadcasting, Publishing, Corporate Communications, Social Media'
            }
        },
        'agriculture_food': {
            'growth_rate': 'Stable (6-10% annually) - Food security focus and agricultural modernization',
            'job_security': 'High - Essential industry with government support and growing food demand',
            'work_life_balance': 'Good - Seasonal work patterns with natural breaks and rural lifestyle',
            'remote_work': 'Limited - Field work and hands-on agricultural activities required',
            'skill_demand': 'Moderate - Traditional knowledge with increasing need for modern techniques',
            'career_progression': 'Steady - Experience-based advancement with opportunities for specialization',
            'education_requirements': 'Varies - Traditional knowledge to agricultural degrees for technical roles',
            'entry_barriers': 'Low to Moderate - Land access and initial capital for farming operations',
            'market_outlook': 'Strong - Food security priorities and export potential for Nepal',
            'advantages': [
                'Direct contribution to food security and national development',
                'Meaningful rural development impact and community building',
                'Opportunity to implement sustainable and organic practices',
                'Strong government support through subsidies and programs',
                'Essential industry with consistent demand and stability',
                'Connection with nature and environmentally conscious work',
                'Potential for organic and premium product markets',
                'Self-sufficiency and food production independence',
                'Seasonal work patterns allowing for rest periods',
                'Opportunity for cooperative farming and community collaboration'
            ],
            'disadvantages': [
                'High dependency on weather patterns and climate conditions',
                'Seasonal income fluctuations affecting financial stability',
                'Physically demanding labor requiring strength and endurance',
                'Market price volatility affecting profitability and planning',
                'Limited technology adoption increasing labor requirements',
                'Pest and disease risks threatening crop yields',
                'Limited access to modern farming equipment and techniques',
                'Difficulty accessing credit and financial services',
                'Climate change impacts on traditional farming practices',
                'Limited value addition and processing opportunities'
            ],
            'career_factors': {
                'skills_needed': 'Agricultural techniques, soil management, pest control, marketing, business planning, sustainability',
                'certifications': 'Organic farming, agricultural extension, food safety, cooperative management, sustainable agriculture',
                'experience_levels': 'Farm Worker (0-3 years), Farm Manager (3-8 years), Agricultural Specialist (8-15 years), Agribusiness Owner (15+ years)',
                'specialization_paths': 'Crop Production, Livestock, Organic Farming, Agricultural Technology, Food Processing, Agribusiness',
                'industry_applications': 'Farming, Food Processing, Agricultural Cooperatives, Government Extension, NGOs, Research'
            }
        },
        'retail_sales': {
            'growth_rate': 'Moderate (8-12% annually) - E-commerce growth and consumer market expansion',
            'job_security': 'Moderate - Economic sensitivity but essential consumer service industry',
            'work_life_balance': 'Moderate - Weekend work required but predictable schedules',
            'remote_work': 'Limited - Customer-facing roles require physical presence in stores',
            'skill_demand': 'Moderate - Customer service and sales skills with digital literacy growing',
            'career_progression': 'Moderate - Clear advancement path from sales to management roles',
            'education_requirements': 'High school minimum, sales training and product knowledge important',
            'entry_barriers': 'Low - Entry-level positions available with on-the-job training',
            'market_outlook': 'Good - Growing consumer market and retail expansion in Nepal',
            'advantages': [
                'Rich customer interaction and relationship building skills',
                'Valuable sales skill development transferable across industries',
                'Deep product knowledge and market understanding',
                'Commission and incentive opportunities for high performers',
                'Clear management pathways and career advancement opportunities',
                'Flexible scheduling options in many retail environments',
                'Development of communication and persuasion skills',
                'Networking opportunities with diverse customer base',
                'Understanding of consumer behavior and market trends',
                'Potential for entrepreneurship and business ownership'
            ],
            'disadvantages': [
                'Required weekend and holiday work during peak shopping periods',
                'Stressful customer service situations and complaint handling',
                'Long hours standing and physical demands of retail work',
                'Pressure from sales targets and performance metrics',
                'Economic sensitivity affecting sales and job security',
                'Dealing with difficult customers and return situations',
                'Limited career growth without management aspirations',
                'Seasonal employment fluctuations in some retail sectors',
                'Competition from online retail affecting traditional stores',
                'Lower base wages with income dependent on commissions'
            ],
            'career_factors': {
                'skills_needed': 'Customer service, sales techniques, product knowledge, communication, problem-solving, cash handling',
                'certifications': 'Sales training, customer service, retail management, digital payment systems, inventory management',
                'experience_levels': 'Sales Associate (0-2 years), Senior Associate (2-4 years), Supervisor (4-7 years), Manager (7+ years)',
                'specialization_paths': 'Sales Management, Visual Merchandising, Inventory Management, Customer Relations, E-commerce',
                'industry_applications': 'Department Stores, Specialty Retail, Electronics, Fashion, Grocery, Online Retail, Automotive'
            }
        },
        'manufacturing_production': {
            'growth_rate': 'Moderate (7-11% annually) - Industrial growth and export manufacturing expansion',
            'job_security': 'High - Essential production roles with growing manufacturing sector',
            'work_life_balance': 'Good - Structured shifts with clear work-life boundaries',
            'remote_work': 'Very Limited - Hands-on production work requires physical presence',
            'skill_demand': 'High - Technical skills and quality control expertise in demand',
            'career_progression': 'Steady - Clear advancement from operator to supervisor to manager',
            'education_requirements': 'Technical training or diploma, specialized certifications beneficial',
            'entry_barriers': 'Moderate - Technical aptitude and safety training required',
            'market_outlook': 'Strong - Manufacturing growth and Made in Nepal initiatives',
            'advantages': [
                'Comprehensive technical skill development and expertise',
                'Excellent job stability with established manufacturing companies',
                'Clear and structured career progression pathways',
                'Valuable industrial experience transferable globally',
                'Export opportunities and international market exposure',
                'Competitive wages with overtime and shift differentials',
                'Comprehensive benefits including health insurance and bonuses',
                'Opportunity to work with advanced machinery and technology',
                'Team-based work environment with strong camaraderie',
                'Contribution to national economic development and exports'
            ],
            'disadvantages': [
                'Significant safety concerns and occupational hazards',
                'Mandatory shift work including nights and weekends',
                'High physical demands and repetitive motion requirements',
                'Environmental factors including noise, heat, and chemical exposure',
                'Automation threats potentially eliminating traditional jobs',
                'Limited creativity and variety in daily work tasks',
                'Strict quality control and production pressure',
                'Health risks from industrial processes and materials',
                'Economic sensitivity affecting production volumes',
                'Limited flexibility in work schedules and locations'
            ],
            'career_factors': {
                'skills_needed': 'Technical operations, quality control, safety protocols, machinery operation, problem-solving, teamwork',
                'certifications': 'Safety training, quality management, machinery operation, lean manufacturing, industrial maintenance',
                'experience_levels': 'Operator (0-3 years), Senior Operator (3-6 years), Supervisor (6-10 years), Production Manager (10+ years)',
                'specialization_paths': 'Quality Control, Maintenance, Production Planning, Safety Management, Process Improvement',
                'industry_applications': 'Textiles, Food Processing, Electronics, Automotive, Pharmaceuticals, Construction Materials'
            }
        },
        'transportation_logistics': {
            'growth_rate': 'High (10-15% annually) - E-commerce boom and cross-border trade expansion',
            'job_security': 'High - Essential services with growing demand for efficient logistics',
            'work_life_balance': 'Moderate - Long hours during peak seasons but structured routes',
            'remote_work': 'Limited - Physical transportation and warehouse operations required',
            'skill_demand': 'High - Logistics planning, technology systems, and route optimization',
            'career_progression': 'Good - Clear advancement from driver to fleet manager to operations director',
            'education_requirements': 'High school to logistics degree, commercial driving licenses required',
            'entry_barriers': 'Moderate - Driving licenses, vehicle knowledge, and safety training required',
            'market_outlook': 'Excellent - E-commerce growth and Nepal\'s strategic trade location',
            'advantages': [
                'Rapidly growing e-commerce demand creating new opportunities',
                'Valuable supply chain expertise applicable across industries',
                'Travel opportunities and geographic mobility',
                'Strong problem-solving skills development',
                'International trade exposure and cross-border experience',
                'High demand for logistics professionals',
                'Opportunity to work with modern technology and tracking systems',
                'Essential service providing job security',
                'Potential for fleet ownership and business development',
                'Competitive wages with overtime and delivery bonuses'
            ],
            'disadvantages': [
                'Intense time-sensitive deadlines and delivery pressure',
                'Challenging traffic and poor road conditions in many areas',
                'High dependency on fuel costs affecting profitability',
                'Complex regulatory compliance and documentation requirements',
                'Significant physical demands and long driving hours',
                'Safety risks from road conditions and cargo handling',
                'Irregular schedules and time away from home',
                'Vehicle maintenance costs and operational expenses',
                'Weather dependency affecting delivery schedules',
                'Stress from customer expectations and delivery commitments'
            ],
            'career_factors': {
                'skills_needed': 'Driving skills, route planning, logistics software, customer service, time management, safety protocols',
                'certifications': 'Commercial driving license, logistics management, safety training, hazardous materials handling',
                'experience_levels': 'Driver (0-3 years), Senior Driver (3-6 years), Supervisor (6-10 years), Operations Manager (10+ years)',
                'specialization_paths': 'Fleet Management, Warehouse Operations, Supply Chain Planning, International Logistics, Last-Mile Delivery',
                'industry_applications': 'E-commerce, Manufacturing, Import/Export, Retail Distribution, Food Delivery, Courier Services'
            }
        },
        'construction_real_estate': {
            'growth_rate': 'High (12-18% annually) - Reconstruction efforts and urban development boom',
            'job_security': 'Moderate - Project-based work with high demand for skilled professionals',
            'work_life_balance': 'Challenging - Long hours during construction seasons and project deadlines',
            'remote_work': 'Very Limited - Hands-on construction and site supervision required',
            'skill_demand': 'High - Technical construction skills and project management expertise',
            'career_progression': 'Good - Merit-based advancement from laborer to contractor to developer',
            'education_requirements': 'Technical training to engineering degrees, trade certifications valuable',
            'entry_barriers': 'Moderate - Physical ability, technical skills, and safety training required',
            'market_outlook': 'Excellent - Post-earthquake reconstruction and rapid urbanization in Nepal',
            'advantages': [
                'Exceptional high earning potential with skilled trades and management',
                'Direct contribution to infrastructure development and national progress',
                'Diverse project variety from residential to commercial to infrastructure',
                'Highly transferable skills applicable globally',
                'Property investment opportunities and real estate knowledge',
                'Strong job demand due to reconstruction and development needs',
                'Opportunity for entrepreneurship and contracting business',
                'Tangible results and satisfaction from completed projects',
                'Team-based work environment with strong professional relationships',
                'Potential for international construction projects and mobility'
            ],
            'disadvantages': [
                'High weather dependency affecting work schedules and productivity',
                'Significant safety risks and occupational hazards',
                'Project-based employment creating income instability',
                'Extremely demanding physical work and manual labor',
                'High sensitivity to economic cycles and market downturns',
                'Long working hours during peak construction seasons',
                'Exposure to dust, noise, and hazardous materials',
                'Pressure from tight project deadlines and budget constraints',
                'Seasonal unemployment during monsoon and winter periods',
                'High competition and bidding pressure for projects'
            ],
            'career_factors': {
                'skills_needed': 'Construction techniques, project management, safety protocols, blueprint reading, equipment operation, quality control',
                'certifications': 'Trade licenses, safety training, project management (PMP), equipment operation, building codes',
                'experience_levels': 'Laborer (0-2 years), Skilled Worker (2-5 years), Supervisor (5-10 years), Contractor/Manager (10+ years)',
                'specialization_paths': 'Residential Construction, Commercial Development, Infrastructure, Project Management, Real Estate Development',
                'industry_applications': 'Residential Building, Commercial Construction, Infrastructure Projects, Real Estate Development, Renovation'
            }
        }
    }
    
    # Calculate statistics
    total_positions = sum(len(positions) for positions in nepal_salary_data.values())
    total_sectors = len(nepal_salary_data)
    
    # Create a flattened list for JavaScript consumption
    positions_for_js = []
    for sector, positions in nepal_salary_data.items():
        if isinstance(positions, dict):
            # Handle dictionary format (like it_tech)
            for position, data in positions.items():
                positions_for_js.append({
                    'sector': sector,
                    'position': position,
                    'minSalary': data['min'],
                    'maxSalary': data['max']
                })
        elif isinstance(positions, list):
            # Handle list format (like healthcare)
            for item in positions:
                # Parse salary range from string format
                salary_range = item['salary_range'].replace(',', '')
                if ' - ' in salary_range:
                    min_sal, max_sal = salary_range.split(' - ')
                    min_salary = int(min_sal)
                    max_salary = int(max_sal.replace('+', ''))
                else:
                    # Single value salary
                    min_salary = max_salary = int(salary_range.replace('+', ''))
                
                positions_for_js.append({
                    'sector': sector,
                    'position': item['position'],
                    'minSalary': min_salary,
                    'maxSalary': max_salary
                })
    
    import json
    positions_json = json.dumps(positions_for_js)
    
    # Convert all data to consistent list format for template rendering
    salary_data_for_template = {}
    for sector, positions in nepal_salary_data.items():
        if isinstance(positions, dict):
            # Convert dictionary format to list format
            salary_data_for_template[sector] = []
            for position, data in positions.items():
                salary_data_for_template[sector].append({
                    'position': position,
                    'salary_range': f"{data['min']:,} - {data['max']:,}",
                    'annual_range': f"{data['min']*12/100000:.1f}L - {data['max']*12/100000:.1f}L"
                })
        else:
            # Keep list format as is
            salary_data_for_template[sector] = positions
    
    # Serialize job insights to JSON for JavaScript consumption
    job_insights_json = json.dumps(job_insights)
    
    context = {
        'nepal_salary_data': nepal_salary_data,
        'salary_data': salary_data_for_template,
        'total_positions': total_positions,
        'total_sectors': total_sectors,
        'positions_json': positions_json,
        'job_insights': job_insights_json,
    }
    
    return render(request, 'accounts/nepal_salary_guide.html', context)


@login_required
@jobseeker_required
def salary_comparison_result(request):
    """Display detailed salary comparison results in a dedicated page"""
    
    # Get comparison data from session or URL parameters
    pos1_data = request.GET.get('pos1', '')
    pos2_data = request.GET.get('pos2', '')
    
    if not pos1_data or not pos2_data:
        messages.error(request, 'Invalid comparison data. Please select positions to compare.')
        return redirect('accounts:nepal_salary_guide')
    
    try:
        # Parse position data
        pos1_parts = pos1_data.split('|')
        pos2_parts = pos2_data.split('|')
        
        if len(pos1_parts) != 4 or len(pos2_parts) != 4:
            raise ValueError("Invalid position data format")
        
        position1 = {
            'sector': pos1_parts[0],
            'position': pos1_parts[1],
            'min_salary': int(pos1_parts[2]),
            'max_salary': int(pos1_parts[3]),
            'avg_salary': (int(pos1_parts[2]) + int(pos1_parts[3])) / 2,
            'salary_range': int(pos1_parts[3]) - int(pos1_parts[2])
        }
        
        position2 = {
            'sector': pos2_parts[0],
            'position': pos2_parts[1],
            'min_salary': int(pos2_parts[2]),
            'max_salary': int(pos2_parts[3]),
            'avg_salary': (int(pos2_parts[2]) + int(pos2_parts[3])) / 2,
            'salary_range': int(pos2_parts[3]) - int(pos2_parts[2])
        }
        
        # Get job insights for both positions
        job_insights = {
            'it_tech': {
                'growth_rate': 'High (15-20% annually)',
                'job_security': 'High',
                'work_life_balance': 'Moderate to Good',
                'remote_work': 'Excellent',
                'career_progression': 'Fast',
                'advantages': ['High salary potential', 'Remote work opportunities', 'Continuous learning', 'Global market access', 'Flexible hours'],
                'disadvantages': ['High stress', 'Constant upskilling needed', 'Long hours', 'Sedentary lifestyle', 'Rapid tech changes']
            },
            'healthcare': {
                'growth_rate': 'Stable (8-12% annually)',
                'job_security': 'Very High',
                'work_life_balance': 'Challenging',
                'remote_work': 'Limited',
                'career_progression': 'Steady',
                'advantages': ['Job security', 'Social impact', 'Respected profession', 'Continuous demand', 'Good benefits'],
                'disadvantages': ['Long hours', 'High stress', 'Emotional demands', 'Limited remote work', 'Extensive education']
            },
            'finance': {
                'growth_rate': 'Moderate (10-15% annually)',
                'job_security': 'High',
                'work_life_balance': 'Moderate',
                'remote_work': 'Good',
                'career_progression': 'Good',
                'advantages': ['High earning potential', 'Career advancement', 'Analytical skills', 'Business acumen', 'Networking'],
                'disadvantages': ['High pressure', 'Market volatility', 'Long hours', 'Regulatory burden', 'Economic dependency']
            },
            'banking_finance': {
                'growth_rate': 'Moderate (10-15% annually)',
                'job_security': 'High',
                'work_life_balance': 'Moderate',
                'remote_work': 'Good',
                'career_progression': 'Good',
                'advantages': ['High earning potential', 'Career advancement', 'Analytical skills', 'Business acumen', 'Networking'],
                'disadvantages': ['High pressure', 'Market volatility', 'Long hours', 'Regulatory burden', 'Economic dependency']
            },
            'education': {
                'growth_rate': 'Stable (5-8% annually)',
                'job_security': 'High',
                'work_life_balance': 'Good',
                'remote_work': 'Moderate',
                'career_progression': 'Slow but Steady',
                'advantages': ['Job satisfaction', 'Work-life balance', 'Vacation time', 'Pension benefits', 'Intellectual stimulation'],
                'disadvantages': ['Lower salary potential', 'Limited mobility', 'Bureaucratic constraints', 'Resource limitations', 'Slow growth']
            },
            'engineering': {
                'growth_rate': 'Moderate (8-12% annually)',
                'job_security': 'High',
                'work_life_balance': 'Good',
                'remote_work': 'Moderate',
                'career_progression': 'Steady',
                'advantages': ['Problem-solving', 'Infrastructure impact', 'Technical expertise', 'Project variety', 'Good progression'],
                'disadvantages': ['Deadline pressure', 'Site work', 'Weather dependency', 'Safety concerns', 'Regulatory compliance']
            },
            'marketing': {
                'growth_rate': 'High (12-18% annually)',
                'job_security': 'Moderate',
                'work_life_balance': 'Moderate',
                'remote_work': 'Excellent',
                'career_progression': 'Fast',
                'advantages': ['Creative work', 'Remote flexibility', 'Networking', 'Diverse projects', 'Digital skills'],
                'disadvantages': ['Campaign pressure', 'Results-driven', 'Market dependency', 'Client stress', 'Irregular hours']
            },
            'government': {
                'growth_rate': 'Stable (3-6% annually)',
                'job_security': 'Very High',
                'work_life_balance': 'Excellent',
                'remote_work': 'Limited',
                'career_progression': 'Slow but Guaranteed',
                'advantages': ['Job security', 'Benefits', 'Work-life balance', 'Pension', 'Social service'],
                'disadvantages': ['Lower salary', 'Bureaucracy', 'Limited innovation', 'Slow progression', 'Political influence']
            },
            'hospitality': {
                'growth_rate': 'Moderate (8-12% annually)',
                'job_security': 'Moderate',
                'work_life_balance': 'Challenging',
                'remote_work': 'Very Limited',
                'career_progression': 'Moderate',
                'advantages': ['People interaction', 'Cultural exposure', 'Travel opportunities', 'Service skills', 'Tip income'],
                'disadvantages': ['Seasonal employment', 'Weekend work', 'Physical demands', 'Customer stress', 'Lower base salaries']
            },
            'media_communications': {
                'growth_rate': 'High (12-15% annually)',
                'job_security': 'Moderate',
                'work_life_balance': 'Moderate',
                'remote_work': 'Good',
                'career_progression': 'Fast',
                'advantages': ['Creative expression', 'Public influence', 'Networking', 'Content creation', 'Digital growth'],
                'disadvantages': ['Deadline pressure', 'Public scrutiny', 'Irregular income', 'Tech adaptation', 'Market competition']
            },
            'agriculture_food': {
                'growth_rate': 'Stable (6-10% annually)',
                'job_security': 'High',
                'work_life_balance': 'Good',
                'remote_work': 'Limited',
                'career_progression': 'Steady',
                'advantages': ['Food security impact', 'Rural development', 'Sustainable practices', 'Government support', 'Essential industry'],
                'disadvantages': ['Weather dependency', 'Seasonal income', 'Physical labor', 'Price volatility', 'Limited tech adoption']
            },
            'retail_sales': {
                'growth_rate': 'Moderate (8-12% annually)',
                'job_security': 'Moderate',
                'work_life_balance': 'Moderate',
                'remote_work': 'Limited',
                'career_progression': 'Moderate',
                'advantages': ['Customer interaction', 'Sales skills', 'Product knowledge', 'Commission opportunities', 'Management paths'],
                'disadvantages': ['Weekend work', 'Customer stress', 'Standing hours', 'Sales pressure', 'Economic sensitivity']
            },
            'manufacturing_production': {
                'growth_rate': 'Moderate (7-11% annually)',
                'job_security': 'High',
                'work_life_balance': 'Good',
                'remote_work': 'Very Limited',
                'career_progression': 'Steady',
                'advantages': ['Technical skills', 'Job stability', 'Clear progression', 'Industrial experience', 'Export opportunities'],
                'disadvantages': ['Safety concerns', 'Shift work', 'Physical demands', 'Environmental factors', 'Automation threats']
            },
            'transportation_logistics': {
                'growth_rate': 'High (10-15% annually)',
                'job_security': 'High',
                'work_life_balance': 'Moderate',
                'remote_work': 'Limited',
                'career_progression': 'Good',
                'advantages': ['E-commerce demand', 'Supply chain expertise', 'Travel opportunities', 'Problem-solving', 'International exposure'],
                'disadvantages': ['Time pressure', 'Traffic conditions', 'Fuel dependency', 'Regulatory compliance', 'Physical demands']
            },
            'construction_real_estate': {
                'growth_rate': 'High (12-18% annually)',
                'job_security': 'Moderate',
                'work_life_balance': 'Challenging',
                'remote_work': 'Very Limited',
                'career_progression': 'Good',
                'advantages': ['High earning potential', 'Infrastructure development', 'Project variety', 'Skill transferability', 'Property investment'],
                'disadvantages': ['Weather dependency', 'Safety risks', 'Project-based employment', 'Physical demands', 'Economic sensitivity']
            }
        }
        
        # Get insights for both positions
        pos1_insights = job_insights.get(position1['sector'], job_insights['it_tech'])
        pos2_insights = job_insights.get(position2['sector'], job_insights['it_tech'])
        
        # Add insights to position data
        position1.update({
            'advantages': pos1_insights['advantages'],
            'disadvantages': pos1_insights['disadvantages'],
            'factors': {
                'growth_rate': pos1_insights['growth_rate'],
                'job_security': pos1_insights['job_security'],
                'work_life_balance': pos1_insights['work_life_balance'],
                'remote_work': pos1_insights['remote_work'],
                'career_progression': pos1_insights['career_progression']
            },
            'factor_scores': {
                'growth': 4 if 'High' in pos1_insights['growth_rate'] else 3 if 'Moderate' in pos1_insights['growth_rate'] else 2,
                'security': 5 if pos1_insights['job_security'] == 'Very High' else 4 if pos1_insights['job_security'] == 'High' else 3,
                'balance': 4 if 'Good' in pos1_insights['work_life_balance'] else 3 if 'Moderate' in pos1_insights['work_life_balance'] else 2,
                'remote': 5 if pos1_insights['remote_work'] == 'Excellent' else 4 if pos1_insights['remote_work'] == 'Good' else 2,
                'progression': 4 if 'Fast' in pos1_insights['career_progression'] else 3 if 'Steady' in pos1_insights['career_progression'] else 2
            }
        })
        
        position2.update({
            'advantages': pos2_insights['advantages'],
            'disadvantages': pos2_insights['disadvantages'],
            'factors': {
                'growth_rate': pos2_insights['growth_rate'],
                'job_security': pos2_insights['job_security'],
                'work_life_balance': pos2_insights['work_life_balance'],
                'remote_work': pos2_insights['remote_work'],
                'career_progression': pos2_insights['career_progression']
            },
            'factor_scores': {
                'growth': 4 if 'High' in pos2_insights['growth_rate'] else 3 if 'Moderate' in pos2_insights['growth_rate'] else 2,
                'security': 5 if pos2_insights['job_security'] == 'Very High' else 4 if pos2_insights['job_security'] == 'High' else 3,
                'balance': 4 if 'Good' in pos2_insights['work_life_balance'] else 3 if 'Moderate' in pos2_insights['work_life_balance'] else 2,
                'remote': 5 if pos2_insights['remote_work'] == 'Excellent' else 4 if pos2_insights['remote_work'] == 'Good' else 2,
                'progression': 4 if 'Fast' in pos2_insights['career_progression'] else 3 if 'Steady' in pos2_insights['career_progression'] else 2
            }
        })
        
        # Calculate salary differences
        difference = abs(position1['avg_salary'] - position2['avg_salary'])
        percent_diff = (difference / min(position1['avg_salary'], position2['avg_salary'])) * 100
        annual_diff = difference * 12
        
        salary_difference = {
            'difference': difference,
            'percent_diff': round(percent_diff, 1),
            'annual_diff': annual_diff
        }
        
        # Generate recommendation
        score1 = sum(position1['factor_scores'].values()) + (position1['avg_salary'] / 10000)
        score2 = sum(position2['factor_scores'].values()) + (position2['avg_salary'] / 10000)
        
        if abs(score1 - score2) < 2:
            winner = "Both positions are equally viable"
            reason = "Both positions offer similar overall value. Consider your personal interests and career goals."
            timeframe = "Take time to evaluate which aligns better with your preferences."
        elif score1 > score2:
            winner = position1['position']
            reason = f"{position1['position']} offers better overall career prospects with higher growth potential and salary."
            timeframe = "Consider this position for better long-term career prospects."
        else:
            winner = position2['position']
            reason = f"{position2['position']} provides better overall value with superior career factors and growth opportunities."
            timeframe = "This position offers better long-term career prospects."
        
        recommendation = {
            'winner': winner,
            'reason': reason,
            'timeframe': timeframe,
            'score': round(max(score1, score2), 1)
        }
        
        comparison_score = round((max(score1, score2) / 20) * 100, 0)
        
        # Prepare chart data for JavaScript
        chart_data = {
            'pos1Salary': position1['avg_salary'],
            'pos2Salary': position2['avg_salary'],
            'pos1Name': position1['position'],
            'pos2Name': position2['position'],
            'pos1Factors': [
                position1['factor_scores']['growth'],
                position1['factor_scores']['security'],
                position1['factor_scores']['balance'],
                position1['factor_scores']['remote'],
                position1['factor_scores']['progression']
            ],
            'pos2Factors': [
                position2['factor_scores']['growth'],
                position2['factor_scores']['security'],
                position2['factor_scores']['balance'],
                position2['factor_scores']['remote'],
                position2['factor_scores']['progression']
            ]
        }
        
        context = {
            'position1': position1,
            'position2': position2,
            'salary_difference': salary_difference,
            'recommendation': recommendation,
            'comparison_score': comparison_score,
            'chart_data': chart_data,
        }
        
        return render(request, 'accounts/salary_comparison_result.html', context)
        
    except (ValueError, IndexError) as e:
        messages.error(request, 'Invalid comparison data format. Please try again.')
        return redirect('accounts:nepal_salary_guide')


def calculate_salary_insights(jobseeker_profile, applications):
    """Calculate salary insights and negotiation data for the user"""
    from django.db.models import Avg, Min, Max
    from jobs.models import JobPost
    
    insights = {
        'market_data': {},
        'negotiation_tips': [],
        'salary_range': {},
        'user_applications': []
    }
    
    if not applications.exists():
        return insights
    
    # Get salary data from user's applications
    app_salaries = []
    for app in applications:
        job = app.job
        if job.min_salary and job.max_salary:
            app_salaries.append({
                'company': job.company.name,
                'position': job.title,
                'min_salary': job.min_salary,
                'max_salary': job.max_salary,
                'currency': job.salary_currency,
                'status': app.status
            })
    
    insights['user_applications'] = app_salaries
    
    # Calculate market data based on similar positions
    if app_salaries:
        # Get job titles from applications
        job_titles = [app.job.title for app in applications]
        
        # Find similar jobs in the market
        similar_jobs = JobPost.objects.filter(
            title__icontains=job_titles[0].split()[0],  # Use first word of first job title
            min_salary__isnull=False,
            max_salary__isnull=False
        ).aggregate(
            avg_min=Avg('min_salary'),
            avg_max=Avg('max_salary'),
            market_min=Min('min_salary'),
            market_max=Max('max_salary')
        )
        
        insights['market_data'] = {
            'average_min': similar_jobs['avg_min'] or 0,
            'average_max': similar_jobs['avg_max'] or 0,
            'market_min': similar_jobs['market_min'] or 0,
            'market_max': similar_jobs['market_max'] or 0,
        }
        
        # Generate negotiation tips based on user's profile
        tips = []
        
        if jobseeker_profile.experience_years >= 3:
            tips.append("Highlight your {} years of experience as a key value proposition".format(jobseeker_profile.experience_years))
        
        if jobseeker_profile.skill_set.count() >= 5:
            tips.append("Emphasize your diverse skill set ({} skills) to justify higher compensation".format(jobseeker_profile.skill_set.count()))
        
        if jobseeker_profile.educations.exists():
            tips.append("Mention your educational background as additional qualification")
        
        # Add general tips
        tips.extend([
            "Research the company's salary range and industry standards",
            "Focus on total compensation package, not just base salary",
            "Be prepared to discuss your achievements and quantifiable results",
            "Consider non-monetary benefits like flexible work arrangements"
        ])
        
        insights['negotiation_tips'] = tips
    
    return insights

@login_required
def career_progress(request, career_id):
    """Detailed career progress view with analytics"""
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    career_path = get_object_or_404(CareerPath, id=career_id)
    user_progress = get_object_or_404(
        UserCareerProgress,
        user_profile=request.user.userprofile,
        career_path=career_path
    )
    
    # Handle milestone progress updates
    if request.method == 'POST':
        action = request.POST.get('action')
        milestone_id = request.POST.get('milestone_id')
        
        if action == 'update_progress' and milestone_id:
            milestone_progress = get_object_or_404(
                MilestoneProgress,
                user_career_progress=user_progress,
                milestone_id=milestone_id
            )
            is_completed = request.POST.get('is_completed') == 'on'
            notes = request.POST.get('notes', '')
            
            # Update milestone progress
            if is_completed and not milestone_progress.is_completed:
                milestone_progress.is_completed = True
                milestone_progress.completed_at = timezone.now()
                milestone_progress.save()
                messages.success(request, f'Congratulations! You completed: {milestone_progress.milestone.title}')
            elif not is_completed and milestone_progress.is_completed:
                milestone_progress.is_completed = False
                milestone_progress.completed_at = None
                milestone_progress.save()
                messages.info(request, 'Milestone marked as not completed.')
            else:
                messages.info(request, 'Progress updated successfully!')
            
            return redirect('accounts:career_progress', career_id=career_id)
    
    # Get milestone progress with analytics
    milestone_progress = user_progress.milestone_progress.select_related('milestone').order_by('milestone__order')
    
    # Calculate analytics
    total_milestones = milestone_progress.count()
    completed_milestones = milestone_progress.filter(is_completed=True).count()
    # For now, we'll consider all non-completed milestones as "not started"
    # since we don't have a progress_percentage field in the current model
    in_progress_milestones = 0
    
    completion_percentage = user_progress.get_completion_percentage()
    next_milestone = user_progress.get_next_milestone()
    
    # Get recent assessments
    recent_assessments = user_progress.assessments.order_by('-taken_at')[:5]
    
    # Calculate skill gaps and recommendations
    required_skills = career_path.required_skills.split(',')
    user_skills = [skill.name.lower() for skill in request.user.userprofile.jobseekerprofile.skill_set.all()]
    skill_gaps = [skill.strip() for skill in required_skills if skill.strip().lower() not in user_skills]
    
    # Generate progress insights
    insights = generate_career_insights(user_progress, milestone_progress)
    
    context = {
        'user': request.user,
        'profile': request.user.userprofile,
        'career_path': career_path,
        'user_progress': user_progress,
        'milestone_progress': milestone_progress,
        'total_milestones': total_milestones,
        'completed_milestones': completed_milestones,
        'in_progress_milestones': in_progress_milestones,
        'completion_percentage': completion_percentage,
        'next_milestone': next_milestone,
        'recent_assessments': recent_assessments,
        'skill_gaps': skill_gaps,
        'insights': insights,
    }
    return render(request, 'accounts/career_progress.html', context)

@login_required
def take_assessment(request, career_id, milestone_id=None):
    """Take a career assessment"""
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    career_path = get_object_or_404(CareerPath, id=career_id)
    user_progress = get_object_or_404(
        UserCareerProgress,
        user_profile=request.user.userprofile,
        career_path=career_path
    )
    
    milestone = None
    if milestone_id:
        milestone = get_object_or_404(CareerMilestone, id=milestone_id)
    
    if request.method == 'POST':
        assessment_type = request.POST.get('assessment_type')
        title = request.POST.get('title')
        description = request.POST.get('description')
        score = request.POST.get('score')
        
        assessment = CareerAssessment.objects.create(
            user_career_progress=user_progress,
            milestone=milestone,
            assessment_type=assessment_type,
            title=title,
            description=description,
            score=int(score) if score else None
        )
        
        # Generate feedback based on score
        if assessment.score:
            if assessment.score >= 90:
                feedback = "Excellent work! You have mastered this area."
            elif assessment.score >= 75:
                feedback = "Good job! You're on the right track with room for improvement."
            elif assessment.score >= 60:
                feedback = "Fair performance. Consider reviewing the materials and practicing more."
            else:
                feedback = "Needs improvement. Focus on strengthening your fundamentals in this area."
            
            assessment.feedback = feedback
            assessment.save()
        
        messages.success(request, 'Assessment completed successfully!')
        return redirect('accounts:career_progress', career_id=career_id)
    
    context = {
        'user': request.user,
        'career_path': career_path,
        'milestone': milestone,
        'user_progress': user_progress,
    }
    return render(request, 'accounts/take_assessment.html', context)

def generate_career_insights(user_progress, milestone_progress):
    """Generate actionable career insights and recommendations"""
    insights = []
    
    # Progress-based insights
    completion_rate = user_progress.get_completion_percentage()
    
    if completion_rate == 0:
        insights.append({
            'type': 'getting_started',
            'title': 'Getting Started',
            'message': 'Welcome to your career journey! Start with the first milestone to build momentum.',
            'action': 'Begin your first milestone'
        })
    elif completion_rate < 25:
        insights.append({
            'type': 'early_progress',
            'title': 'Building Foundation',
            'message': 'Great start! Focus on completing foundational milestones to establish your base knowledge.',
            'action': 'Complete basic skill milestones'
        })
    elif completion_rate < 50:
        insights.append({
            'type': 'steady_progress',
            'title': 'Steady Progress',
            'message': 'You\'re making good progress! Consider taking on more challenging projects.',
            'action': 'Start intermediate-level projects'
        })
    elif completion_rate < 75:
        insights.append({
            'type': 'advanced_progress',
            'title': 'Advanced Development',
            'message': 'Excellent progress! Focus on specialized skills and real-world applications.',
            'action': 'Pursue advanced certifications'
        })
    else:
        insights.append({
            'type': 'expert_level',
            'title': 'Expert Level',
            'message': 'Outstanding! You\'re ready for leadership roles and mentoring others.',
            'action': 'Consider teaching or leading projects'
        })
    
    # Milestone-specific insights
    # Since we don't have progress_percentage, we'll consider all incomplete milestones as potential areas for focus
    stalled_milestones = milestone_progress.filter(is_completed=False).count()
    
    if stalled_milestones > 2:
        insights.append({
            'type': 'stalled_progress',
            'title': 'Focus Needed',
            'message': f'You have {stalled_milestones} milestones in progress. Consider focusing on completing one at a time.',
            'action': 'Complete pending milestones'
        })
    
    # Time-based insights
    from datetime import datetime, timedelta
    recent_activity = milestone_progress.filter(
        completed_at__gte=datetime.now() - timedelta(days=30)
    ).count()
    
    if recent_activity == 0 and completion_rate > 0:
        insights.append({
            'type': 'inactive',
            'title': 'Stay Active',
            'message': 'No recent progress detected. Regular practice is key to career development.',
            'action': 'Set weekly learning goals'
        })
    
    return insights

# Profile sub-pages
@login_required
def profile_skills(request):
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    context = {
        'user': request.user,
        'profile': request.user.userprofile,
    }
    return render(request, 'accounts/profile/skills.html', context)

@login_required
def profile_experience(request):
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    context = {
        'user': request.user,
        'profile': request.user.userprofile,
    }
    return render(request, 'accounts/profile/experience.html', context)

@login_required
def profile_education(request):
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    context = {
        'user': request.user,
        'profile': request.user.userprofile,
    }
    return render(request, 'accounts/profile/education.html', context)

@login_required
def profile_certifications(request):
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    context = {
        'user': request.user,
        'profile': request.user.userprofile,
    }
    return render(request, 'accounts/profile/certifications.html', context)

@login_required
def profile_portfolio(request):
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    context = {
        'user': request.user,
        'profile': request.user.userprofile,
    }
    return render(request, 'accounts/profile/portfolio.html', context)

@login_required
def profile_preferences(request):
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    context = {
        'user': request.user,
        'profile': request.user.userprofile,
    }
    return render(request, 'accounts/profile/preferences.html', context)

# Resume Builder sub-pages
@login_required
def resume_builder_templates(request):
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    context = {
        'user': request.user,
        'profile': request.user.userprofile,
    }
    return render(request, 'accounts/resume/templates.html', context)

@login_required
def resume_preview(request):
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    context = {
        'user': request.user,
        'profile': request.user.userprofile,
    }
    return render(request, 'accounts/resume/preview.html', context)

# Cover Letter sub-pages
@login_required
def cover_letter_templates(request):
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    context = {
        'user': request.user,
        'profile': request.user.userprofile,
    }
    return render(request, 'accounts/cover_letter/templates.html', context)

# Interview Prep sub-pages
@login_required
def interview_practice(request):
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    context = {
        'user': request.user,
        'profile': request.user.userprofile,
    }
    return render(request, 'accounts/interview/practice.html', context)

@login_required
def mock_interview(request):
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    context = {
        'user': request.user,
        'profile': request.user.userprofile,
    }
    return render(request, 'accounts/interview/mock_interview.html', context)

@login_required
def voice_practice(request):
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    context = {
        'user': request.user,
        'profile': request.user.userprofile,
    }
    return render(request, 'accounts/interview_prep/voice_practice.html', context)

# Career Roadmap sub-pages
@login_required
def skills_gap_analysis(request):
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    context = {
        'user': request.user,
        'profile': request.user.userprofile,
    }
    return render(request, 'accounts/career_roadmap/skills_gap.html', context)

# Salary Negotiation sub-pages
@login_required
def salary_calculator(request):
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    context = {
        'user': request.user,
        'profile': request.user.userprofile,
    }
    return render(request, 'accounts/salary_negotiation/calculator.html', context)

@login_required
def negotiation_guides(request):
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    context = {
        'user': request.user,
        'profile': request.user.userprofile,
    }
    return render(request, 'accounts/salary_negotiation/guides.html', context)

# Applications management sub-pages
@login_required
def application_analytics(request):
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    context = {
        'user': request.user,
        'profile': request.user.userprofile,
    }
    return render(request, 'accounts/applications/analytics.html', context)

@login_required
def application_tracker(request):
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    context = {
        'user': request.user,
        'profile': request.user.userprofile,
    }
    return render(request, 'accounts/applications/tracker.html', context)

# Job management sub-pages

@login_required
def manage_job_alerts(request):
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    context = {
        'user': request.user,
        'profile': request.user.userprofile,
    }
    return render(request, 'accounts/jobs/manage_alerts.html', context)

@login_required
def notification_center(request):
    """Notification center page"""
    return render(request, 'notifications/notification_center.html')


# AI-Powered Interview Preparation Views
@login_required
def interview_prep(request):
    """Main interview preparation dashboard"""
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    user_profile = request.user.userprofile
    
    # Get or create analytics
    analytics, created = InterviewAnalytics.objects.get_or_create(user_profile=user_profile)
    if created:
        analytics.update_analytics()
    
    # Recent sessions
    recent_sessions = InterviewSession.objects.filter(
        user_profile=user_profile
    ).order_by('-started_at')[:5]
    
    # Question categories with counts
    question_counts = {}
    for category, label in InterviewQuestion.CATEGORY_CHOICES:
        question_counts[category] = {
            'label': label,
            'count': InterviewQuestion.objects.filter(category=category, is_active=True).count()
        }
    
    context = {
        'analytics': analytics,
        'recent_sessions': recent_sessions,
        'question_counts': question_counts,
        'industries': InterviewQuestion.INDUSTRY_CHOICES,
        'difficulties': InterviewQuestion.DIFFICULTY_CHOICES,
    }
    return render(request, 'accounts/interview_prep.html', context)


@login_required
def start_interview_session(request):
    """Start a new interview session with selected parameters and intelligent question selection"""
    import json
    from django.urls import reverse
    
    if request.user.userprofile.user_type != 'jobseeker':
        return JsonResponse({'error': 'Access denied'}, status=403)
    
    if request.method == 'POST':
        try:
            # Handle both JSON and form data
            if request.content_type == 'application/json':
                data = json.loads(request.body)
            else:
                data = request.POST
            session_type = data.get('session_type', 'practice')
            industry = data.get('industry', 'general')
            difficulty = data.get('difficulty', 'beginner')
            category = data.get('category', 'behavioral')
            question_count = int(data.get('question_count', 5))
            
            user_profile = request.user.userprofile
            
            # Get user analytics for intelligent difficulty adjustment
            analytics, _ = InterviewAnalytics.objects.get_or_create(user_profile=user_profile)
            
            # Adjust difficulty based on user performance
            adjusted_difficulty = adjust_difficulty_based_on_performance(analytics, difficulty, category)
            
            # Create new session
            session = InterviewSession.objects.create(
                user_profile=user_profile,
                session_type=session_type,
                industry_focus=industry,
                difficulty_level=adjusted_difficulty,
                total_questions=question_count,
                status='active'
            )
            
            # Select and randomize questions
            questions = select_randomized_questions(
                category=category,
                industry=industry,
                difficulty=adjusted_difficulty,
                count=question_count,
                user_profile=user_profile
            )
            
            # Check if we have enough questions
            if not questions:
                # Create some default questions if none exist
                from .models import InterviewQuestion
                default_question = InterviewQuestion.objects.create(
                    question_text="Tell me about yourself and your professional background.",
                    category='behavioral',
                    industry='general',
                    difficulty='beginner',
                    is_active=True
                )
                questions = [default_question]
            
            # Store question order in session
            question_ids = [q.id for q in questions]
            session.question_order = question_ids
            session.save()
            
            return JsonResponse({
                'success': True,
                'session_id': session.id,
                'difficulty_adjusted': adjusted_difficulty != difficulty,
                'adjusted_difficulty': adjusted_difficulty,
                'redirect_url': reverse('accounts:interview_session', args=[session.id])
            })
            
        except Exception as e:
            return JsonResponse({'error': f'Failed to create session: {str(e)}'}, status=500)
    
    return JsonResponse({'error': 'Invalid request'}, status=400)


def adjust_difficulty_based_on_performance(analytics, requested_difficulty, category):
    """Intelligently adjust difficulty based on user's historical performance"""
    # Get category-specific performance with safe attribute access
    category_scores = {
        'behavioral': getattr(analytics, 'behavioral_avg_score', 0),
        'technical': getattr(analytics, 'technical_avg_score', 0),
        'communication': getattr(analytics, 'communication_avg_score', 0),
        'leadership': getattr(analytics, 'leadership_avg_score', 0),
        'problem_solving': getattr(analytics, 'problem_solving_avg_score', 0),
        'situational': getattr(analytics, 'situational_avg_score', 0),
    }
    
    category_score = category_scores.get(category, getattr(analytics, 'average_score', 0))
    
    # Difficulty adjustment logic
    if category_score >= 85:
        # High performer - can handle advanced questions
        if requested_difficulty == 'beginner':
            return 'intermediate'
        elif requested_difficulty == 'intermediate':
            return 'advanced'
    elif category_score <= 60:
        # Struggling - keep at easier levels
        if requested_difficulty == 'advanced':
            return 'intermediate'
        elif requested_difficulty == 'intermediate' and category_score <= 50:
            return 'beginner'
    
    return requested_difficulty


def select_randomized_questions(category, industry, difficulty, count, user_profile):
    """Select randomized questions with intelligent filtering"""
    import random
    from django.db.models import Q
    
    # Get previously answered questions to avoid repetition
    answered_question_ids = InterviewAnswer.objects.filter(
        session__user_profile=user_profile
    ).values_list('question_id', flat=True).distinct()
    
    # Build query filters
    filters = Q(category=category, is_active=True)
    
    # Industry filter - prefer specific industry but include general
    if industry != 'general':
        filters &= Q(Q(industry=industry) | Q(industry='general'))
    else:
        filters &= Q(industry='general')
    
    # Difficulty filter - include requested and one level below for variety
    difficulty_levels = ['beginner', 'intermediate', 'advanced']
    if difficulty in difficulty_levels:
        current_index = difficulty_levels.index(difficulty)
        allowed_difficulties = [difficulty]
        if current_index > 0:  # Include easier questions for warm-up
            allowed_difficulties.append(difficulty_levels[current_index - 1])
        filters &= Q(difficulty__in=allowed_difficulties)
    
    # Get available questions
    available_questions = list(InterviewQuestion.objects.filter(filters))
    
    # Separate new and previously answered questions
    new_questions = [q for q in available_questions if q.id not in answered_question_ids]
    old_questions = [q for q in available_questions if q.id in answered_question_ids]
    
    # Prioritize new questions, but include some old ones for practice
    selected_questions = []
    
    # First, try to get mostly new questions
    if len(new_questions) >= count:
        selected_questions = random.sample(new_questions, count)
    else:
        # Use all new questions and fill with old ones
        selected_questions.extend(new_questions)
        remaining_count = count - len(new_questions)
        if old_questions and remaining_count > 0:
            selected_questions.extend(random.sample(old_questions, min(remaining_count, len(old_questions))))
    
    # If still not enough questions, relax filters
    if len(selected_questions) < count:
        # Include all categories if needed
        fallback_filters = Q(is_active=True)
        if industry != 'general':
            fallback_filters &= Q(Q(industry=industry) | Q(industry='general'))
        
        fallback_questions = list(InterviewQuestion.objects.filter(fallback_filters).exclude(
            id__in=[q.id for q in selected_questions]
        ))
        
        remaining_needed = count - len(selected_questions)
        if fallback_questions:
            selected_questions.extend(random.sample(
                fallback_questions, 
                min(remaining_needed, len(fallback_questions))
            ))
    
    # Randomize the final order
    random.shuffle(selected_questions)
    
    return selected_questions[:count]


@login_required
def interview_session(request, session_id):
    """Interactive interview session page"""
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    session = get_object_or_404(InterviewSession, id=session_id, user_profile=request.user.userprofile)
    
    # Get next question
    answered_questions = session.answers.values_list('question_id', flat=True)
    next_question = InterviewQuestion.objects.filter(
        industry__in=[session.industry_focus, 'general'],
        difficulty=session.difficulty_level,
        is_active=True
    ).exclude(id__in=answered_questions).first()
    
    if not next_question and session.questions_answered == 0:
        # Fallback to general questions if no specific industry questions
        next_question = InterviewQuestion.objects.filter(
            industry='general',
            difficulty=session.difficulty_level,
            is_active=True
        ).first()
    
    context = {
        'session': session,
        'question': next_question,
        'progress_percentage': session.completion_percentage,
    }
    return render(request, 'accounts/interview_session.html', context)


@login_required
@require_POST
def submit_answer(request, session_id):
    """Submit answer for evaluation with session-specific logic"""
    try:
        if request.user.userprofile.user_type != 'jobseeker':
            return JsonResponse({'error': 'Access denied'}, status=403)
        
        session = get_object_or_404(InterviewSession, id=session_id, user_profile=request.user.userprofile)
        
        question_id = request.POST.get('question_id')
        answer_text = request.POST.get('answer_text', '').strip()
        input_type = request.POST.get('input_type', 'text')
        response_time = int(request.POST.get('response_time', 0))
        session_type = request.POST.get('session_type', session.session_type)
        
        if not question_id:
            return JsonResponse({'error': 'Question ID is required'}, status=400)
            
        if not answer_text and input_type != 'video':
            return JsonResponse({'error': 'Answer text is required'}, status=400)
        
        question = get_object_or_404(InterviewQuestion, id=question_id)
        
        # Create answer record
        answer = InterviewAnswer.objects.create(
            session=session,
            question=question,
            answer_text=answer_text,
            input_type=input_type,
            response_time=response_time
        )
        
        # Session-specific AI evaluation
        evaluation_result = evaluate_answer_by_session_type(answer, question, session_type, response_time)
        
        # Update answer with evaluation
        answer.content_score = evaluation_result['content_score']
        answer.clarity_score = evaluation_result['clarity_score']
        answer.confidence_score = evaluation_result['confidence_score']
        answer.keyword_score = evaluation_result['keyword_score']
        answer.overall_score = evaluation_result['overall_score']
        answer.strengths = evaluation_result['strengths']
        answer.weaknesses = evaluation_result['weaknesses']
        answer.suggestions = evaluation_result['suggestions']
        answer.alternative_phrases = evaluation_result['alternative_phrases']
        answer.save()
        
        # Update session progress
        session.questions_answered += 1
        if session.questions_answered >= session.total_questions:
            session.status = 'completed'
            session.completed_at = timezone.now()
            session.overall_score = session.answers.aggregate(
                avg_score=models.Avg('overall_score')
            )['avg_score'] or 0.0
        session.save()
        
        # Update analytics
        analytics, created = InterviewAnalytics.objects.get_or_create(user_profile=request.user.userprofile)
        analytics.update_analytics()
        
        return JsonResponse({
            'success': True,
            'evaluation': {
                'overall_score': answer.overall_score,
                'content_score': answer.content_score,
                'clarity_score': answer.clarity_score,
                'confidence_score': answer.confidence_score,
                'keyword_score': answer.keyword_score,
                'strengths': answer.strengths,
                'weaknesses': answer.weaknesses,
                'suggestions': answer.suggestions,
                'alternative_phrases': answer.alternative_phrases,
            },
            'session_completed': session.status == 'completed',
            'next_question_available': session.questions_answered < session.total_questions
        })
        
    except Exception as e:
        return JsonResponse({'error': f'Failed to submit answer: {str(e)}'}, status=500)


@login_required
def interview_analytics(request):
    """Interview preparation analytics dashboard"""
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    user_profile = request.user.userprofile
    analytics, created = InterviewAnalytics.objects.get_or_create(user_profile=user_profile)
    
    # Recent sessions with detailed scores
    recent_sessions = InterviewSession.objects.filter(
        user_profile=user_profile,
        status='completed'
    ).order_by('-completed_at')[:10]
    
    # Performance trends
    performance_data = []
    for session in recent_sessions:
        if session.overall_score:
            performance_data.append({
                'date': session.completed_at.strftime('%Y-%m-%d'),
                'score': session.overall_score,
                'session_type': session.get_session_type_display()
            })
    
    # Category breakdown
    category_performance = {}
    for category, label in InterviewQuestion.CATEGORY_CHOICES:
        answers = InterviewAnswer.objects.filter(
            session__user_profile=user_profile,
            question__category=category
        )
        if answers.exists():
            avg_score = answers.aggregate(avg=Avg('overall_score'))['avg'] or 0.0
            category_performance[category] = {
                'label': label,
                'score': round(avg_score, 1),
                'count': answers.count()
            }
    
    context = {
        'analytics': analytics,
        'recent_sessions': recent_sessions,
        'performance_data': performance_data,
        'category_performance': category_performance,
    }
    return render(request, 'accounts/interview_analytics.html', context)


def evaluate_answer(answer, question):
    """Advanced AI-powered answer evaluation with comprehensive analysis"""
    import re
    import string
    from collections import Counter
    from django.db import models
    
    answer_text = answer.answer_text
    answer_lower = answer_text.lower()
    question_keywords = question.keywords if question.keywords else []
    
    # Initialize scores
    content_score = 60
    clarity_score = 70
    confidence_score = 75
    grammar_score = 80
    structure_score = 70
    
    # Character-level analysis
    char_analysis = analyze_character_level(answer_text)
    
    # Grammar and language analysis
    grammar_analysis = analyze_grammar_and_language(answer_text)
    
    # Sentence structure analysis
    structure_analysis = analyze_sentence_structure(answer_text)
    
    # Word-level analysis
    word_analysis = analyze_word_level(answer_text, question_keywords)
    
    # Content Score Enhancement
    content_score += word_analysis['keyword_bonus']
    content_score += word_analysis['vocabulary_bonus']
    content_score += char_analysis['completeness_bonus']
    
    # Clarity Score Enhancement
    clarity_score += structure_analysis['clarity_bonus']
    clarity_score += grammar_analysis['readability_bonus']
    clarity_score -= char_analysis['repetition_penalty']
    
    # Grammar Score
    grammar_score += grammar_analysis['grammar_bonus']
    grammar_score -= grammar_analysis['error_penalty']
    
    # Structure Score
    structure_score += structure_analysis['organization_bonus']
    structure_score += structure_analysis['flow_bonus']
    
    # Confidence Score Enhancement
    confidence_score += word_analysis['confidence_bonus']
    confidence_score -= word_analysis['uncertainty_penalty']
    
    # Normalize scores (0-100)
    content_score = max(0, min(100, content_score))
    clarity_score = max(0, min(100, clarity_score))
    confidence_score = max(0, min(100, confidence_score))
    grammar_score = max(0, min(100, grammar_score))
    structure_score = max(0, min(100, structure_score))
    
    # Calculate overall score with weighted average
    overall_score = (
        content_score * 0.3 +
        clarity_score * 0.25 +
        grammar_score * 0.2 +
        structure_score * 0.15 +
        confidence_score * 0.1
    )
    
    # STAR method detection (enhanced)
    star_analysis = analyze_star_method(answer_text)
    
    # Generate comprehensive feedback
    feedback = generate_comprehensive_feedback(
        answer_text, word_analysis, grammar_analysis, 
        structure_analysis, char_analysis, star_analysis
    )
    
    return {
        'content_score': round(content_score, 1),
        'clarity_score': round(clarity_score, 1),
        'confidence_score': round(confidence_score, 1),
        'grammar_score': round(grammar_score, 1),
        'structure_score': round(structure_score, 1),
        'keyword_score': round(word_analysis['keyword_score'], 1),
        'overall_score': round(overall_score, 1),
        'strengths': feedback['strengths'],
        'weaknesses': feedback['weaknesses'],
        'suggestions': feedback['suggestions'],
        'alternative_phrases': feedback['alternative_phrases'],
        'detailed_analysis': {
            'character_analysis': char_analysis,
            'grammar_analysis': grammar_analysis,
            'structure_analysis': structure_analysis,
            'word_analysis': word_analysis,
            'star_analysis': star_analysis
        }
    }


def analyze_character_level(text):
    """Analyze text at character level for completeness and patterns"""
    import string
    from collections import Counter
    
    # Character distribution
    char_count = Counter(text.lower())
    total_chars = len(text)
    
    # Calculate character diversity
    unique_chars = len(set(text.lower()))
    char_diversity = unique_chars / max(total_chars, 1) * 100
    
    # Punctuation usage
    punctuation_count = sum(1 for c in text if c in string.punctuation)
    punctuation_ratio = punctuation_count / max(total_chars, 1)
    
    # Repetition detection
    words = text.lower().split()
    word_count = Counter(words)
    repeated_words = sum(1 for count in word_count.values() if count > 2)
    repetition_penalty = min(repeated_words * 2, 15)
    
    # Completeness indicators
    has_proper_ending = text.strip().endswith(('.', '!', '?'))
    has_capital_start = text.strip() and text.strip()[0].isupper()
    
    completeness_bonus = 0
    if has_proper_ending:
        completeness_bonus += 5
    if has_capital_start:
        completeness_bonus += 3
    if punctuation_ratio > 0.02:  # Good punctuation usage
        completeness_bonus += 5
    
    return {
        'char_diversity': char_diversity,
        'punctuation_ratio': punctuation_ratio,
        'repetition_penalty': repetition_penalty,
        'completeness_bonus': completeness_bonus,
        'has_proper_ending': has_proper_ending,
        'has_capital_start': has_capital_start
    }


def analyze_grammar_and_language(text):
    """Analyze grammar, spelling, and language quality"""
    import re
    
    # Common grammar patterns
    grammar_issues = 0
    
    # Subject-verb agreement (simplified check)
    singular_subjects = re.findall(r'\b(he|she|it|this|that)\s+(\w+)', text.lower())
    for subject, verb in singular_subjects:
        if verb.endswith('s') and verb not in ['is', 'was', 'has', 'does']:
            continue  # Likely correct
        elif not verb.endswith('s') and verb in ['are', 'were', 'have', 'do']:
            grammar_issues += 1
    
    # Sentence fragments (very basic check)
    sentences = re.split(r'[.!?]+', text)
    fragments = sum(1 for s in sentences if s.strip() and len(s.split()) < 3)
    
    # Run-on sentences
    long_sentences = sum(1 for s in sentences if len(s.split()) > 30)
    
    # Spelling indicators (basic patterns)
    common_errors = [
        r'\bteh\b', r'\badn\b', r'\brecieve\b', r'\boccured\b',
        r'\bseperate\b', r'\bdefinately\b', r'\bneccessary\b'
    ]
    spelling_errors = sum(len(re.findall(pattern, text.lower())) for pattern in common_errors)
    
    # Readability indicators
    import string
    words = text.split()
    avg_word_length = sum(len(word.strip(string.punctuation)) for word in words) / max(len(words), 1)
    
    # Calculate bonuses and penalties
    grammar_bonus = 0
    error_penalty = 0
    readability_bonus = 0
    
    if grammar_issues == 0:
        grammar_bonus += 10
    else:
        error_penalty += grammar_issues * 3
    
    if fragments == 0:
        grammar_bonus += 5
    else:
        error_penalty += fragments * 2
    
    if long_sentences == 0:
        readability_bonus += 5
    elif long_sentences > 2:
        readability_bonus -= 5
    
    if spelling_errors == 0:
        grammar_bonus += 8
    else:
        error_penalty += spelling_errors * 4
    
    if 4 <= avg_word_length <= 6:
        readability_bonus += 8
    
    return {
        'grammar_issues': grammar_issues,
        'fragments': fragments,
        'long_sentences': long_sentences,
        'spelling_errors': spelling_errors,
        'avg_word_length': avg_word_length,
        'grammar_bonus': grammar_bonus,
        'error_penalty': error_penalty,
        'readability_bonus': readability_bonus
    }


def analyze_sentence_structure(text):
    """Analyze sentence structure and organization"""
    import re
    
    sentences = [s.strip() for s in re.split(r'[.!?]+', text) if s.strip()]
    
    if not sentences:
        return {
            'sentence_count': 0,
            'avg_length': 0,
            'variety_score': 0,
            'clarity_bonus': 0,
            'organization_bonus': 0,
            'flow_bonus': 0
        }
    
    # Sentence length analysis
    sentence_lengths = [len(s.split()) for s in sentences]
    avg_length = sum(sentence_lengths) / len(sentence_lengths)
    
    # Sentence variety (different lengths)
    length_variety = len(set(sentence_lengths)) / len(sentence_lengths)
    variety_score = length_variety * 100
    
    # Transition words and phrases
    transitions = [
        'however', 'therefore', 'furthermore', 'moreover', 'additionally',
        'consequently', 'meanwhile', 'subsequently', 'nevertheless', 'thus',
        'first', 'second', 'finally', 'in conclusion', 'for example'
    ]
    transition_count = sum(1 for word in transitions if word in text.lower())
    
    # Paragraph structure (if multiple sentences)
    has_intro = len(sentences) > 2 and any(
        word in sentences[0].lower() 
        for word in ['i', 'my', 'when', 'during', 'in my experience']
    )
    
    has_conclusion = len(sentences) > 2 and any(
        word in sentences[-1].lower() 
        for word in ['therefore', 'thus', 'in conclusion', 'overall', 'ultimately']
    )
    
    # Calculate bonuses
    clarity_bonus = 0
    organization_bonus = 0
    flow_bonus = 0
    
    if 10 <= avg_length <= 20:
        clarity_bonus += 10
    elif avg_length > 25:
        clarity_bonus -= 5
    
    if variety_score > 50:
        clarity_bonus += 8
    
    if transition_count > 0:
        flow_bonus += min(transition_count * 3, 12)
    
    if has_intro:
        organization_bonus += 8
    if has_conclusion:
        organization_bonus += 8
    
    return {
        'sentence_count': len(sentences),
        'avg_length': avg_length,
        'variety_score': variety_score,
        'transition_count': transition_count,
        'has_intro': has_intro,
        'has_conclusion': has_conclusion,
        'clarity_bonus': clarity_bonus,
        'organization_bonus': organization_bonus,
        'flow_bonus': flow_bonus
    }


def analyze_word_level(text, keywords):
    """Analyze vocabulary, keywords, and word choice"""
    import re
    from collections import Counter
    
    words = re.findall(r'\b\w+\b', text.lower())
    word_count = len(words)
    unique_words = len(set(words))
    
    # Vocabulary diversity
    vocab_diversity = unique_words / max(word_count, 1) * 100
    
    # Keyword analysis
    keyword_matches = sum(1 for keyword in keywords if keyword.lower() in text.lower())
    keyword_score = (keyword_matches / max(len(keywords), 1)) * 100 if keywords else 70
    
    # Advanced vocabulary indicators
    advanced_words = [
        'demonstrate', 'implement', 'facilitate', 'optimize', 'leverage',
        'collaborate', 'innovate', 'strategic', 'comprehensive', 'analytical',
        'proficient', 'expertise', 'methodology', 'initiative', 'leadership'
    ]
    advanced_count = sum(1 for word in advanced_words if word in text.lower())
    
    # Confidence indicators
    confidence_words = [
        'confident', 'experienced', 'skilled', 'successfully', 'achieved',
        'accomplished', 'expert', 'proficient', 'mastered', 'excelled'
    ]
    uncertainty_words = [
        'maybe', 'perhaps', 'might', 'possibly', 'not sure', 'i think',
        'probably', 'somewhat', 'kind of', 'sort of'
    ]
    
    confidence_count = sum(1 for word in confidence_words if word in text.lower())
    uncertainty_count = sum(1 for word in uncertainty_words if word in text.lower())
    
    # Calculate bonuses and penalties
    keyword_bonus = min(keyword_matches * 5, 20)
    vocabulary_bonus = min(advanced_count * 3, 15)
    confidence_bonus = min(confidence_count * 4, 16)
    uncertainty_penalty = min(uncertainty_count * 6, 18)
    
    if vocab_diversity > 70:
        vocabulary_bonus += 8
    elif vocab_diversity < 40:
        vocabulary_bonus -= 5
    
    return {
        'word_count': word_count,
        'unique_words': unique_words,
        'vocab_diversity': vocab_diversity,
        'keyword_matches': keyword_matches,
        'keyword_score': keyword_score,
        'advanced_count': advanced_count,
        'confidence_count': confidence_count,
        'uncertainty_count': uncertainty_count,
        'keyword_bonus': keyword_bonus,
        'vocabulary_bonus': vocabulary_bonus,
        'confidence_bonus': confidence_bonus,
        'uncertainty_penalty': uncertainty_penalty
    }


def analyze_star_method(text):
    """Enhanced STAR method analysis"""
    text_lower = text.lower()
    
    # STAR indicators with more comprehensive patterns
    star_patterns = {
        'situation': [
            'situation', 'when', 'during', 'at my previous job', 'in my role',
            'while working', 'the scenario', 'the context', 'background'
        ],
        'task': [
            'task', 'responsibility', 'needed to', 'had to', 'was required',
            'my role was', 'objective', 'goal', 'assignment'
        ],
        'action': [
            'action', 'i did', 'i implemented', 'i developed', 'i created',
            'i organized', 'i managed', 'steps i took', 'approach'
        ],
        'result': [
            'result', 'outcome', 'achieved', 'accomplished', 'success',
            'impact', 'improvement', 'increased', 'reduced', 'delivered'
        ]
    }
    
    star_scores = {}
    for component, patterns in star_patterns.items():
        score = sum(10 for pattern in patterns if pattern in text_lower)
        star_scores[component] = min(score, 30)  # Cap at 30 per component
    
    total_star_score = sum(star_scores.values())
    has_complete_star = all(score > 0 for score in star_scores.values())
    
    return {
        'star_scores': star_scores,
        'total_score': total_star_score,
        'has_complete_star': has_complete_star,
        'completeness_percentage': (sum(1 for score in star_scores.values() if score > 0) / 4) * 100
    }


def generate_comprehensive_feedback(text, word_analysis, grammar_analysis, 
                                  structure_analysis, char_analysis, star_analysis):
    """Generate detailed feedback based on all analyses"""
    strengths = []
    weaknesses = []
    suggestions = []
    alternative_phrases = []
    
    # Strengths identification
    if word_analysis['keyword_matches'] >= 3:
        strengths.append("Excellent use of relevant keywords and industry terminology")
    
    if grammar_analysis['grammar_issues'] == 0:
        strengths.append("Strong grammar and language mechanics")
    
    if structure_analysis['variety_score'] > 60:
        strengths.append("Good sentence variety and structure")
    
    if char_analysis['has_proper_ending'] and char_analysis['has_capital_start']:
        strengths.append("Well-formatted response with proper capitalization and punctuation")
    
    if star_analysis['has_complete_star']:
        strengths.append("Excellent use of STAR method structure")
    
    if word_analysis['vocab_diversity'] > 70:
        strengths.append("Rich vocabulary and diverse word choice")
    
    # Weaknesses identification
    if word_analysis['keyword_matches'] < 2:
        weaknesses.append("Limited use of relevant keywords")
        suggestions.append("Include more industry-specific terminology and buzzwords")
    
    if grammar_analysis['spelling_errors'] > 0:
        weaknesses.append("Contains spelling errors")
        suggestions.append("Proofread your response for spelling accuracy")
    
    if grammar_analysis['fragments'] > 0:
        weaknesses.append("Contains sentence fragments")
        suggestions.append("Ensure all sentences are complete with subject and verb")
    
    if structure_analysis['avg_length'] > 25:
        weaknesses.append("Sentences are too long and complex")
        suggestions.append("Break down long sentences for better clarity")
    
    if not star_analysis['has_complete_star']:
        weaknesses.append("Response lacks clear STAR method structure")
        suggestions.append("Structure your answer using Situation, Task, Action, Result format")
    
    if word_analysis['uncertainty_count'] > 2:
        weaknesses.append("Contains too many uncertainty words")
        suggestions.append("Use more confident language to demonstrate expertise")
    
    # Alternative phrases based on common words
    text_lower = text.lower()
    if 'good' in text_lower:
        alternative_phrases.append("Replace 'good' with 'excellent', 'outstanding', or 'exceptional'")
    if 'worked' in text_lower:
        alternative_phrases.append("Replace 'worked' with 'collaborated', 'spearheaded', or 'orchestrated'")
    if 'helped' in text_lower:
        alternative_phrases.append("Replace 'helped' with 'facilitated', 'supported', or 'enabled'")
    if 'did' in text_lower:
        alternative_phrases.append("Replace 'did' with 'executed', 'implemented', or 'delivered'")
    if 'made' in text_lower:
        alternative_phrases.append("Replace 'made' with 'created', 'developed', or 'established'")
    
    return {
        'strengths': strengths,
        'weaknesses': weaknesses,
        'suggestions': suggestions,
        'alternative_phrases': alternative_phrases
    }


def evaluate_answer_by_session_type(answer, question, session_type, response_time):
    """
    Session-specific AI-powered answer evaluation
    Returns detailed scoring and feedback based on session type
    """
    answer_text = answer.answer_text.lower()
    
    # Base evaluation
    base_evaluation = evaluate_answer(answer, question)
    
    # Session-specific adjustments
    if session_type == 'practice':
        # More lenient scoring for practice mode
        base_evaluation['overall_score'] = min(100, base_evaluation['overall_score'] + 5)
        base_evaluation['suggestions'].insert(0, "Great practice! Keep working on your responses.")
        
    elif session_type == 'mock':
        # Professional evaluation for mock interviews
        if response_time > 300:  # Over 5 minutes
            base_evaluation['weaknesses'].append("Response time was longer than recommended")
            base_evaluation['suggestions'].append("Practice being more concise while maintaining detail")
        
        # Add professional communication feedback
        if 'um' in answer_text or 'uh' in answer_text:
            base_evaluation['weaknesses'].append("Contains filler words")
            base_evaluation['suggestions'].append("Practice speaking more confidently without filler words")
            
    elif session_type == 'challenge':
        # Quick thinking evaluation for challenge mode
        if response_time <= 60:  # Under 1 minute
            base_evaluation['strengths'].append("Excellent quick thinking and response time")
            base_evaluation['overall_score'] = min(100, base_evaluation['overall_score'] + 10)
        elif response_time > 120:  # Over 2 minutes
            base_evaluation['weaknesses'].append("Response time exceeded optimal range for quick challenges")
            
        # Emphasize conciseness
        word_count = len(answer_text.split())
        if word_count > 100:
            base_evaluation['suggestions'].append("For quick challenges, aim for more concise responses")
    
    return base_evaluation


@login_required
def application_tracking(request):
    """
    Job seeker application tracking dashboard with real-time updates
    """
    if not hasattr(request.user, 'userprofile') or request.user.userprofile.user_type != 'job_seeker':
        messages.error(request, 'Access denied. This page is for job seekers only.')
        return redirect('accounts:login')
    
    from applications.models import Application
    from django.core.paginator import Paginator
    from django.db.models import Count, Q
    
    # Get user's applications with related data
    try:
        job_seeker = JobSeekerProfile.objects.get(user_profile=request.user.userprofile)
    except JobSeekerProfile.DoesNotExist:
        job_seeker = JobSeekerProfile.objects.create(user_profile=request.user.userprofile)
    
    applications = Application.objects.filter(
        applicant=job_seeker
    ).select_related(
        'job', 'job__company'
    ).prefetch_related(
        'status_history', 'interviews', 'messages'
    ).order_by('-applied_at')
    
    # Filter by status if requested
    status_filter = request.GET.get('status')
    if status_filter and status_filter != 'all':
        applications = applications.filter(status=status_filter)
    
    # Calculate statistics
    total_applications = applications.count()
    active_applications = applications.filter(
        status__in=['applied', 'reviewing', 'shortlisted', 'interviewing', 'offered']
    ).count()
    interview_count = applications.filter(
        status__in=['interviewing', 'offered', 'hired']
    ).count()
    
    # Calculate success rate (responses received vs total applications)
    responded_applications = applications.exclude(status='applied').count()
    success_rate = round((responded_applications / total_applications * 100) if total_applications > 0 else 0)
    
    # Count unread notifications
    unread_notifications = 0
    for app in applications:
        unread_notifications += app.messages.filter(is_read=False, sender_type='employer').count()
    
    # Pagination
    paginator = Paginator(applications, 10)
    page_number = request.GET.get('page')
    applications_page = paginator.get_page(page_number)
    
    # Add progress percentage for each application
    for application in applications_page:
        application.get_progress_percentage = get_application_progress_percentage(application.status)
        # Handle messages if they exist
        try:
            application.unread_messages_count = application.messages.filter(
                is_read=False, sender_type='employer'
            ).count()
        except:
            application.unread_messages_count = 0
    
    context = {
        'applications': applications_page,
        'total_applications': total_applications,
        'active_applications': active_applications,
        'interview_count': interview_count,
        'success_rate': success_rate,
        'unread_notifications': unread_notifications,
        'current_status_filter': status_filter or 'all'
    }
    
    return render(request, 'accounts/application_tracking.html', context)


def get_application_progress_percentage(status):
    """Calculate progress percentage based on application status"""
    status_progress = {
        'applied': 20,
        'reviewing': 40,
        'shortlisted': 60,
        'interviewing': 80,
        'offered': 90,
        'hired': 100,
        'rejected': 100,
        'withdrawn': 100
    }
    return status_progress.get(status, 0)


@login_required
@require_POST
def withdraw_application(request, application_id):
    """
    Allow job seekers to withdraw their application
    """
    from applications.models import Application, ApplicationStatus
    from django.http import JsonResponse
    
    try:
        application = Application.objects.get(
            id=application_id,
            applicant=request.user.userprofile
        )
        
        # Check if application can be withdrawn
        if application.status in ['hired', 'rejected', 'withdrawn']:
            return JsonResponse({
                'success': False,
                'error': 'This application cannot be withdrawn.'
            })
        
        # Update application status
        application.status = 'withdrawn'
        application.save()
        
        # Create status history entry
        ApplicationStatus.objects.create(
            application=application,
            status='withdrawn',
            changed_by=request.user,
            notes='Application withdrawn by candidate'
        )
        
        # Send notification to employer
        from applications.models import Notification
        Notification.objects.create(
            user=application.job.posted_by.user,
            title='Application Withdrawn',
            message=f'{request.user.get_full_name()} has withdrawn their application for {application.job.title}',
            notification_type='application_update',
            related_application=application
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Application withdrawn successfully.'
        })
        
    except Application.DoesNotExist:
        return JsonResponse({
            'success': False,
            'error': 'Application not found.'
        })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        })


@login_required
@jobseeker_required
def send_jobseeker_message(request, application_id):
    """Send a message to employer for a specific application"""
    if request.method == 'POST':
        try:
            # Get the application
            application = Application.objects.select_related('job', 'job__employer').get(
                id=application_id,
                applicant=request.user.userprofile
            )
            
            content = request.POST.get('content', '').strip()
            if not content:
                return JsonResponse({
                    'success': False,
                    'error': 'Message content is required'
                })
            
            # Create the message
            from applications.models import Message
            message = Message.objects.create(
                application=application,
                sender=request.user,
                recipient=application.job.employer.user_profile.user,
                subject=f'Message regarding {application.job.title}',
                content=content,
                message_type='application'
            )
            
            # Create notification for employer
            from applications.models import Notification
            Notification.objects.create(
                user=application.job.employer.user_profile.user,
                notification_type='message',
                title=f'New message from {request.user.get_full_name()}',
                message=f'You have received a new message regarding the application for {application.job.title}',
                application=application
            )
            
            return JsonResponse({
                'success': True,
                'message': 'Message sent successfully',
                'sender_name': request.user.get_full_name()
            })
            
        except Application.DoesNotExist:
            return JsonResponse({
                'success': False,
                'error': 'Application not found'
            })
        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': f'Failed to send message: {str(e)}'
            })
    
    return JsonResponse({
        'success': False,
        'error': 'Invalid request method'
    })


@login_required
def check_application_updates(request):
    """
    API endpoint to check for new application updates
    """
    from applications.models import Application
    from django.http import JsonResponse
    
    if not hasattr(request.user, 'userprofile') or request.user.userprofile.user_type != 'job_seeker':
        return JsonResponse({'has_updates': False})
    
    # Check for recent updates (last 5 minutes)
    from django.utils import timezone
    from datetime import timedelta
    
    recent_time = timezone.now() - timedelta(minutes=5)
    
    recent_updates = Application.objects.filter(
        applicant=request.user.userprofile,
        updated_at__gte=recent_time
    ).exists()
    
    # Check for unread messages
    unread_messages = Application.objects.filter(
        applicant=request.user.userprofile,
        messages__is_read=False,
        messages__sender_type='employer',
        messages__created_at__gte=recent_time
    ).exists()
    
    has_updates = recent_updates or unread_messages
    
    return JsonResponse({
        'has_updates': has_updates,
        'recent_updates': recent_updates,
        'unread_messages': unread_messages
    })


@login_required
def notification_center(request):
    """
    Comprehensive notification center for job seekers
    """
    if not hasattr(request.user, 'userprofile') or request.user.userprofile.user_type != 'job_seeker':
        messages.error(request, 'Access denied. This page is for job seekers only.')
        return redirect('accounts:login')
    
    from applications.models import Notification
    from django.core.paginator import Paginator
    from django.utils import timezone
    from datetime import timedelta
    
    # Get notifications for the user
    notifications = Notification.objects.filter(
        user=request.user
    ).order_by('-created_at')
    
    # Filter by type if requested
    filter_type = request.GET.get('filter', 'all')
    if filter_type != 'all':
        if filter_type == 'unread':
            notifications = notifications.filter(is_read=False)
        else:
            notifications = notifications.filter(notification_type=filter_type)
    
    # Calculate statistics
    total_notifications = Notification.objects.filter(user=request.user).count()
    unread_count = Notification.objects.filter(user=request.user, is_read=False).count()
    
    # Today's notifications
    today = timezone.now().date()
    today_count = Notification.objects.filter(
        user=request.user,
        created_at__date=today
    ).count()
    
    # Urgent notifications (unread application updates and interviews)
    urgent_count = Notification.objects.filter(
        user=request.user,
        is_read=False,
        notification_type__in=['application_update', 'interview_scheduled', 'status_change']
    ).count()
    
    # Pagination
    paginator = Paginator(notifications, 15)
    page_number = request.GET.get('page')
    notifications_page = paginator.get_page(page_number)
    
    context = {
        'notifications': notifications_page,
        'total_notifications': total_notifications,
        'unread_count': unread_count,
        'today_count': today_count,
        'urgent_count': urgent_count,
        'current_filter': filter_type
    }
    
    return render(request, 'accounts/notifications.html', context)


@login_required
def interview_results(request, session_id):
    """Display detailed results for completed interview session"""
    if request.user.userprofile.user_type != 'jobseeker':
        return redirect('employer:dashboard')
    
    session = get_object_or_404(InterviewSession, id=session_id, user_profile=request.user.userprofile)
    
    if session.status != 'completed':
        return redirect('accounts:interview_session', session_id=session.id)
    
    answers = session.answers.all()
    
    # Calculate session statistics
    total_response_time = sum(answer.response_time for answer in answers)
    average_response_time = total_response_time / len(answers) if answers else 0
    
    # Calculate average scores
    average_content_score = answers.aggregate(avg=models.Avg('content_score'))['avg'] or 0
    average_clarity_score = answers.aggregate(avg=models.Avg('clarity_score'))['avg'] or 0
    average_confidence_score = answers.aggregate(avg=models.Avg('confidence_score'))['avg'] or 0
    average_keyword_score = answers.aggregate(avg=models.Avg('keyword_score'))['avg'] or 0
    
    # Calculate session duration
    if session.completed_at and session.started_at:
        duration = session.completed_at - session.started_at
        duration_minutes = int(duration.total_seconds() / 60)
    else:
        duration_minutes = 0
    
    # Aggregate feedback
    all_strengths = []
    all_weaknesses = []
    all_suggestions = []
    
    for answer in answers:
        if answer.strengths:
            all_strengths.extend(answer.strengths)
        if answer.weaknesses:
            all_weaknesses.extend(answer.weaknesses)
        if answer.suggestions:
            all_suggestions.extend(answer.suggestions)
    
    # Find common feedback patterns
    from collections import Counter
    common_strengths = [item for item, count in Counter(all_strengths).most_common(3)]
    common_weaknesses = [item for item, count in Counter(all_weaknesses).most_common(3)]
    common_suggestions = [item for item, count in Counter(all_suggestions).most_common(3)]
    
    context = {
        'session': session,
        'answers': answers,
        'average_response_time': average_response_time,
        'average_content_score': average_content_score,
        'average_clarity_score': average_clarity_score,
        'average_confidence_score': average_confidence_score,
        'average_keyword_score': average_keyword_score,
        'duration_minutes': duration_minutes,
        'common_strengths': common_strengths,
        'common_weaknesses': common_weaknesses,
        'common_suggestions': common_suggestions,
    }
    
    return render(request, 'accounts/interview_results.html', context)


@login_required
@require_http_methods(["POST"])
def generate_resume_pdf(request):
    """Generate PDF resume from HTML content"""
    try:
        data = json.loads(request.body)
        html_content = data.get('html_content', '')
        file_name = data.get('file_name', 'resume')
        
        # Create PDF buffer
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=18)
        
        # Get styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1,  # Center alignment
            textColor=colors.HexColor('#2c3e50')
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            spaceAfter=12,
            spaceBefore=12,
            textColor=colors.HexColor('#2c3e50'),
            borderWidth=1,
            borderColor=colors.HexColor('#cccccc'),
            borderPadding=5
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            alignment=4  # Justify
        )
        
        # Parse resume data from request
        resume_data = {
            'firstName': data.get('firstName', ''),
            'lastName': data.get('lastName', ''),
            'email': data.get('email', ''),
            'phone': data.get('phone', ''),
            'location': data.get('location', ''),
            'linkedinUrl': data.get('linkedinUrl', ''),
            'website': data.get('website', ''),
            'summary': data.get('summary', ''),
            'skills': data.get('skills', []),
            'experiences': data.get('experiences', []),
            'education': data.get('education', [])
        }
        
        # Build PDF content
        story = []
        
        # Header
        if resume_data['firstName'] or resume_data['lastName']:
            story.append(Paragraph(f"{resume_data['firstName']} {resume_data['lastName']}", title_style))
        
        # Contact info
        contact_info = []
        if resume_data['location']:
            contact_info.append(resume_data['location'])
        if resume_data['phone']:
            contact_info.append(resume_data['phone'])
        if resume_data['email']:
            contact_info.append(resume_data['email'])
        if resume_data['linkedinUrl']:
            contact_info.append(f"LinkedIn: {resume_data['linkedinUrl']}")
        if resume_data['website']:
            contact_info.append(f"Portfolio: {resume_data['website']}")
        
        if contact_info:
            story.append(Paragraph(" | ".join(contact_info), normal_style))
            story.append(Spacer(1, 12))
        
        # Professional Summary
        if resume_data['summary']:
            story.append(Paragraph("Professional Summary", heading_style))
            story.append(Paragraph(resume_data['summary'], normal_style))
            story.append(Spacer(1, 12))
        
        # Skills
        if resume_data['skills']:
            story.append(Paragraph("Skills", heading_style))
            skills_text = " • ".join([f"{skill['name']} ({skill['level']})" for skill in resume_data['skills']])
            story.append(Paragraph(skills_text, normal_style))
            story.append(Spacer(1, 12))
        
        # Work Experience
        if resume_data['experiences']:
            story.append(Paragraph("Work Experience", heading_style))
            for exp in resume_data['experiences']:
                exp_title = f"<b>{exp['title']}</b>"
                if exp['company']:
                    exp_title += f" – {exp['company']}"
                if exp['period']:
                    exp_title += f" ({exp['period']})"
                story.append(Paragraph(exp_title, normal_style))
                if exp['description']:
                    story.append(Paragraph(exp['description'], normal_style))
                story.append(Spacer(1, 6))
        
        # Education
        if resume_data['education']:
            story.append(Paragraph("Education", heading_style))
            for edu in resume_data['education']:
                edu_text = f"<b>{edu['degree']}</b> – {edu['institution']}"
                if edu['gpa']:
                    edu_text += f", {edu['gpa']}"
                story.append(Paragraph(edu_text, normal_style))
                story.append(Spacer(1, 6))
        
        # Build PDF
        doc.build(story)
        
        # Return PDF response
        buffer.seek(0)
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{file_name}.pdf"'
        return response
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
@require_http_methods(["POST"])
def generate_resume_docx(request):
    """Generate DOCX resume from resume data"""
    try:
        data = json.loads(request.body)
        
        # Create new document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Header - Name
        if data.get('firstName') or data.get('lastName'):
            name_paragraph = doc.add_paragraph()
            name_run = name_paragraph.add_run(f"{data.get('firstName', '')} {data.get('lastName', '')}")
            name_run.font.size = Pt(24)
            name_run.bold = True
            name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact Information
        contact_info = []
        if data.get('location'):
            contact_info.append(data['location'])
        if data.get('phone'):
            contact_info.append(data['phone'])
        if data.get('email'):
            contact_info.append(data['email'])
        if data.get('linkedinUrl'):
            contact_info.append(f"LinkedIn: {data['linkedinUrl']}")
        if data.get('website'):
            contact_info.append(f"Portfolio: {data['website']}")
        
        if contact_info:
            contact_paragraph = doc.add_paragraph(" | ".join(contact_info))
            contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()  # Add space
        
        # Professional Summary
        if data.get('summary'):
            heading = doc.add_paragraph()
            heading_run = heading.add_run("Professional Summary")
            heading_run.font.size = Pt(16)
            heading_run.bold = True
            
            summary_paragraph = doc.add_paragraph(data['summary'])
            summary_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            doc.add_paragraph()  # Add space
        
        # Skills
        if data.get('skills'):
            heading = doc.add_paragraph()
            heading_run = heading.add_run("Skills")
            heading_run.font.size = Pt(16)
            heading_run.bold = True
            
            skills_text = " • ".join([f"{skill['name']} ({skill['level']})" for skill in data['skills']])
            doc.add_paragraph(skills_text)
            doc.add_paragraph()  # Add space
        
        # Work Experience
        if data.get('experiences'):
            heading = doc.add_paragraph()
            heading_run = heading.add_run("Work Experience")
            heading_run.font.size = Pt(16)
            heading_run.bold = True
            
            for exp in data['experiences']:
                exp_paragraph = doc.add_paragraph()
                title_run = exp_paragraph.add_run(exp['title'])
                title_run.bold = True
                
                if exp['company']:
                    exp_paragraph.add_run(f" – {exp['company']}")
                if exp['period']:
                    exp_paragraph.add_run(f" ({exp['period']})")
                
                if exp['description']:
                    desc_paragraph = doc.add_paragraph(exp['description'])
                    desc_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            doc.add_paragraph()  # Add space
        
        # Education
        if data.get('education'):
            heading = doc.add_paragraph()
            heading_run = heading.add_run("Education")
            heading_run.font.size = Pt(16)
            heading_run.bold = True
            
            for edu in data['education']:
                edu_paragraph = doc.add_paragraph()
                degree_run = edu_paragraph.add_run(edu['degree'])
                degree_run.bold = True
                edu_paragraph.add_run(f" – {edu['institution']}")
                if edu['gpa']:
                    edu_paragraph.add_run(f", {edu['gpa']}")
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Return DOCX response
        file_name = f"{data.get('firstName', 'Resume')}_{data.get('lastName', '')}_Resume".replace(' ', '_')
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{file_name}.docx"'
        return response
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
def select_user_type(request):
    """View for social login users to select their account type"""
    if request.method == 'POST':
        user_type = request.POST.get('user_type')
        if user_type in ['jobseeker', 'employer']:
            return handle_social_user_type_selection(request, request.user, user_type)
        else:
            messages.error(request, 'Please select a valid account type.')
    
    return render(request, 'accounts/select_user_type.html')


@login_required
def create_custom_roadmap(request):
    """Create a new custom career roadmap"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            
            roadmap = CustomCareerRoadmap.objects.create(
                user_profile=request.user.userprofile,
                title=data.get('title'),
                target_position=data.get('target_position'),
                description=data.get('description', ''),
                target_salary=data.get('target_salary') if data.get('target_salary') else None,
                target_company=data.get('target_company', ''),
                timeline_months=data.get('timeline_months', 12)
            )
            
            # Create initial steps if provided
            steps_data = data.get('steps', [])
            for i, step_data in enumerate(steps_data):
                CustomRoadmapStep.objects.create(
                    roadmap=roadmap,
                    title=step_data.get('title'),
                    description=step_data.get('description', ''),
                    step_type=step_data.get('step_type', 'other'),
                    priority=step_data.get('priority', 'medium'),
                    order=i,
                    estimated_duration=step_data.get('estimated_duration', ''),
                    resources=step_data.get('resources', ''),
                    cost_estimate=step_data.get('cost_estimate') if step_data.get('cost_estimate') else None
                )
            
            return JsonResponse({
                'success': True,
                'roadmap_id': roadmap.id,
                'message': 'Custom roadmap created successfully!'
            })
            
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    return render(request, 'accounts/create_custom_roadmap.html')


@login_required
def custom_roadmap_detail(request, roadmap_id):
    """View custom roadmap details and progress"""
    roadmap = get_object_or_404(CustomCareerRoadmap, id=roadmap_id, user_profile=request.user.userprofile)
    steps = roadmap.steps.all()
    
    context = {
        'roadmap': roadmap,
        'steps': steps,
        'completion_percentage': roadmap.get_completion_percentage(),
        'current_step': roadmap.get_current_step(),
        'next_milestone': roadmap.get_next_milestone(),
    }
    
    return render(request, 'accounts/custom_roadmap_detail.html', context)


@login_required
def roadmap_detail(request, roadmap_id):
    """Display detailed view of a custom career roadmap"""
    roadmap = get_object_or_404(CustomCareerRoadmap, id=roadmap_id, user=request.user)
    steps = roadmap.steps.all().order_by('order')
    
    context = {
        'roadmap': roadmap,
        'steps': steps
    }
    return render(request, 'accounts/roadmap_detail.html', context)


@login_required
def edit_custom_roadmap(request, roadmap_id):
    """Edit custom roadmap details"""
    roadmap = get_object_or_404(CustomCareerRoadmap, id=roadmap_id, user=request.user)
    
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            
            roadmap.title = data.get('title', roadmap.title)
            roadmap.target_position = data.get('target_position', roadmap.target_position)
            roadmap.description = data.get('description', roadmap.description)
            roadmap.target_salary = data.get('target_salary') if data.get('target_salary') else roadmap.target_salary
            roadmap.target_company = data.get('target_company', roadmap.target_company)
            roadmap.timeline_months = data.get('timeline_months', roadmap.timeline_months)
            roadmap.status = data.get('status', roadmap.status)
            roadmap.save()
            
            return JsonResponse({'success': True, 'message': 'Roadmap updated successfully!'})
            
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    context = {'roadmap': roadmap}
    return render(request, 'accounts/edit_custom_roadmap.html', context)


@login_required
def add_roadmap_step(request, roadmap_id):
    """Add a new step to custom roadmap"""
    roadmap = get_object_or_404(CustomCareerRoadmap, id=roadmap_id, user_profile=request.user.userprofile)
    
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            
            # Get the next order number
            max_order = roadmap.steps.aggregate(models.Max('order'))['order__max'] or 0
            
            step = CustomRoadmapStep.objects.create(
                roadmap=roadmap,
                title=data.get('title'),
                description=data.get('description', ''),
                step_type=data.get('step_type', 'other'),
                priority=data.get('priority', 'medium'),
                order=max_order + 1,
                estimated_duration=data.get('estimated_duration', ''),
                resources=data.get('resources', ''),
                cost_estimate=data.get('cost_estimate') if data.get('cost_estimate') else None
            )
            
            return JsonResponse({
                'success': True,
                'step_id': step.id,
                'message': 'Step added successfully!'
            })
            
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    context = {'roadmap': roadmap}
    return render(request, 'accounts/add_roadmap_step.html', context)


@login_required
@require_POST
def update_step_progress(request, step_id):
    """Update progress for a roadmap step"""
    step = get_object_or_404(CustomRoadmapStep, id=step_id, roadmap__user_profile=request.user.userprofile)
    
    try:
        data = json.loads(request.body)
        
        # Create progress log entry
        RoadmapProgressLog.objects.create(
            step=step,
            user_profile=request.user.userprofile,
            progress_percentage=data.get('progress_percentage', 0),
            notes=data.get('notes', ''),
            evidence_url=data.get('evidence_url', '')
        )
        
        # Update step notes if provided
        if data.get('step_notes'):
            step.notes = data.get('step_notes')
            step.save()
        
        return JsonResponse({'success': True, 'message': 'Progress updated successfully!'})
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})


@login_required
@require_POST
def complete_roadmap_step(request, step_id):
    """Mark a roadmap step as completed"""
    try:
        step = get_object_or_404(CustomRoadmapStep, id=step_id, roadmap__user=request.user)
        
        data = json.loads(request.body)
        notes = data.get('notes', 'Step marked as completed')
        
        # Mark step as completed
        step.is_completed = True
        step.completed_at = timezone.now()
        step.save()
        
        # Create progress log
        RoadmapProgressLog.objects.create(
            step=step,
            user_profile=request.user.userprofile,
            progress_percentage=100,
            notes=notes
        )
        
        return JsonResponse({'success': True})
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})


@login_required
def get_step_data(request, step_id):
    """Get step data for editing"""
    try:
        step = get_object_or_404(CustomRoadmapStep, id=step_id, roadmap__user=request.user)
        
        data = {
            'title': step.title,
            'description': step.description or '',
            'step_type': step.step_type,
            'priority': step.priority,
            'estimated_duration': step.estimated_duration or '',
            'resources': step.resources or ''
        }
        
        return JsonResponse(data)
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})


@login_required
@require_POST
def edit_roadmap_step(request, step_id):
    """Edit a roadmap step"""
    try:
        step = get_object_or_404(CustomRoadmapStep, id=step_id, roadmap__user=request.user)
        
        step.title = request.POST.get('title')
        step.description = request.POST.get('description', '')
        step.step_type = request.POST.get('step_type')
        step.priority = request.POST.get('priority')
        step.estimated_duration = request.POST.get('estimated_duration', '')
        step.resources = request.POST.get('resources', '')
        step.save()
        
        return JsonResponse({'success': True})
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})


@login_required
@require_POST
def delete_roadmap_step(request, step_id):
    """Delete a roadmap step"""
    try:
        step = get_object_or_404(CustomRoadmapStep, id=step_id, roadmap__user=request.user)
        step.delete()
        
        return JsonResponse({'success': True})
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})


@login_required
@require_POST
def update_step_order(request, roadmap_id):
    """Update the order of roadmap steps"""
    try:
        roadmap = get_object_or_404(CustomCareerRoadmap, id=roadmap_id, user=request.user)
        
        data = json.loads(request.body)
        updates = data.get('updates', [])
        
        for update in updates:
            step_id = update.get('step_id')
            new_order = update.get('order')
            
            if step_id and new_order is not None:
                CustomRoadmapStep.objects.filter(
                    id=step_id, 
                    roadmap=roadmap
                ).update(order=new_order)
        
        return JsonResponse({'success': True})
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})


@login_required
@require_POST
def delete_custom_roadmap(request, roadmap_id):
    """Delete a custom roadmap"""
    try:
        roadmap = get_object_or_404(CustomCareerRoadmap, id=roadmap_id, user=request.user)
        roadmap.delete()
        
        messages.success(request, f'Roadmap "{roadmap.title}" has been deleted successfully.')
        return redirect('accounts:career_roadmap')
        
    except Exception as e:
        messages.error(request, f'Error deleting roadmap: {str(e)}')
        return redirect('accounts:career_roadmap')


# Delete functionality views
@login_required
@jobseeker_required
def delete_account(request):
    """Delete job seeker account and all related data"""
    if request.method == 'POST':
        password = request.POST.get('password')
        if request.user.check_password(password):
            with transaction.atomic():
                # Delete all related data
                user = request.user
                
                # Delete applications
                Application.objects.filter(user_profile=user.userprofile).delete()
                
                # Delete saved jobs
                SavedJob.objects.filter(user=user).delete()
                
                # Delete job alerts
                JobAlert.objects.filter(user=user).delete()
                
                # Delete notifications
                Notification.objects.filter(user=user).delete()
                
                # Delete messages
                Message.objects.filter(sender=user).delete()
                Message.objects.filter(recipient=user).delete()
                
                # Delete profile data
                if hasattr(user, 'userprofile') and hasattr(user.userprofile, 'jobseekerprofile'):
                    user.userprofile.jobseekerprofile.delete()
                
                if hasattr(user, 'userprofile'):
                    user.userprofile.delete()
                
                # Finally delete the user
                user.delete()
                
                messages.success(request, 'Your account has been permanently deleted.')
                return redirect('home')
        else:
            messages.error(request, 'Incorrect password. Account deletion cancelled.')
    
    return render(request, 'accounts/delete_account.html')


@login_required
@jobseeker_required
def delete_application(request, application_id):
    """Delete a specific job application"""
    application = get_object_or_404(Application, id=application_id, user_profile=request.user.userprofile)
    
    if request.method == 'POST':
        with transaction.atomic():
            # Delete related messages
            Message.objects.filter(application=application).delete()
            
            # Delete related notifications
            Notification.objects.filter(
                Q(message__contains=f'application #{application.id}') |
                Q(message__contains=application.job.title)
            ).delete()
            
            # Delete the application
            job_title = application.job.title
            company_name = application.job.company.name
            application.delete()
            
            messages.success(request, f'Your application for {job_title} at {company_name} has been deleted.')
        
        return redirect('accounts:applications_list')
    
    return render(request, 'accounts/delete_application.html', {'application': application})


@login_required
@jobseeker_required
def delete_job_alert(request, alert_id):
    """Delete a job alert"""
    alert = get_object_or_404(JobAlert, id=alert_id, user=request.user)
    
    if request.method == 'POST':
        alert_title = alert.title
        alert.delete()
        messages.success(request, f'Job alert "{alert_title}" has been deleted.')
        return redirect('accounts:job_alerts_list')
    
    return render(request, 'accounts/delete_job_alert.html', {'alert': alert})


@login_required
@jobseeker_required
def delete_education(request, education_id):
    """Delete an education entry"""
    try:
        jobseeker_profile = request.user.userprofile.jobseekerprofile
        education = get_object_or_404(Education, id=education_id, jobseeker_profile=jobseeker_profile)
        
        if request.method == 'POST':
            institution = education.institution
            degree = education.degree
            education.delete()
            messages.success(request, f'Education entry for {degree} at {institution} has been deleted.')
            return redirect('accounts:profile_detail')
        
        return render(request, 'accounts/delete_education.html', {'education': education})
    except:
        messages.error(request, 'Education entry not found.')
        return redirect('accounts:profile_detail')


@login_required
@jobseeker_required
def delete_experience(request, experience_id):
    """Delete a work experience entry"""
    try:
        jobseeker_profile = request.user.userprofile.jobseekerprofile
        experience = get_object_or_404(WorkExperience, id=experience_id, jobseeker_profile=jobseeker_profile)
        
        if request.method == 'POST':
            company = experience.company
            position = experience.position
            experience.delete()
            messages.success(request, f'Work experience for {position} at {company} has been deleted.')
            return redirect('accounts:profile_detail')
        
        return render(request, 'accounts/delete_experience.html', {'experience': experience})
    except:
        messages.error(request, 'Work experience entry not found.')
        return redirect('accounts:profile_detail')


@login_required
@jobseeker_required
def delete_skill(request, skill_id):
    """Delete a skill entry"""
    try:
        jobseeker_profile = request.user.userprofile.jobseekerprofile
        skill = get_object_or_404(Skill, id=skill_id, jobseeker_profile=jobseeker_profile)
        
        if request.method == 'POST':
            skill_name = skill.name
            skill.delete()
            messages.success(request, f'Skill "{skill_name}" has been deleted.')
            return redirect('accounts:profile_detail')
        
        return render(request, 'accounts/delete_skill.html', {'skill': skill})
    except:
        messages.error(request, 'Skill not found.')
        return redirect('accounts:profile_detail')


# ============================================================================
# FORGOT PASSWORD VIEWS USING SECURITY QUESTIONS
# ============================================================================

def forgot_password_request(request):
    """Step 1: Request password reset by entering username/email"""
    if request.method == 'POST':
        form = ForgotPasswordRequestForm(request.POST)
        if form.is_valid():
            username_or_email = form.cleaned_data['username_or_email']
            
            # Find the user
            user = None
            if '@' in username_or_email:
                try:
                    user = User.objects.get(email=username_or_email)
                except User.DoesNotExist:
                    pass
            else:
                try:
                    user = User.objects.get(username=username_or_email)
                except User.DoesNotExist:
                    pass
            
            if user and user.security_answers.exists():
                # Create password reset session
                reset_session = PasswordResetSession.create_session(user)
                
                # Redirect to security questions with session token
                return redirect('accounts:forgot_password_questions', token=reset_session.session_token)
    else:
        form = ForgotPasswordRequestForm()
    
    return render(request, 'accounts/forgot_password_request.html', {'form': form})


def forgot_password_questions(request, token):
    """Step 2: Answer security questions"""
    try:
        reset_session = PasswordResetSession.objects.get(session_token=token)
        if not reset_session.is_valid():
            messages.error(request, 'Password reset session has expired. Please start over.')
            return redirect('accounts:forgot_password_request')
        
        user = reset_session.user
        
        if request.method == 'POST':
            form = SecurityQuestionsForm(user, request.POST)
            if form.is_valid():
                # Update session with correct answers count
                reset_session.questions_answered_correctly = form.correct_answers
                reset_session.is_verified = True
                reset_session.save()
                
                messages.success(request, f'You answered {form.correct_answers} question(s) correctly. You can now reset your password.')
                return redirect('accounts:forgot_password_reset', token=token)
        else:
            form = SecurityQuestionsForm(user)
        
        return render(request, 'accounts/forgot_password_questions.html', {
            'form': form,
            'user': user,
            'token': token
        })
    
    except PasswordResetSession.DoesNotExist:
        messages.error(request, 'Invalid password reset session. Please start over.')
        return redirect('accounts:forgot_password_request')


def forgot_password_reset(request, token):
    """Step 3: Set new password after verification"""
    try:
        reset_session = PasswordResetSession.objects.get(session_token=token)
        if not reset_session.is_valid() or not reset_session.is_verified:
            messages.error(request, 'Invalid or expired password reset session. Please start over.')
            return redirect('accounts:forgot_password_request')
        
        user = reset_session.user
        
        if request.method == 'POST':
            form = NewPasswordForm(request.POST)
            if form.is_valid():
                new_password = form.cleaned_data['new_password1']
                
                # Set new password
                user.set_password(new_password)
                user.save()
                
                # Mark session as used
                reset_session.mark_as_used()
                
                messages.success(request, 'Your password has been successfully reset. You can now log in with your new password.')
                return redirect('accounts:login')
        else:
            form = NewPasswordForm()
        
        return render(request, 'accounts/forgot_password_reset.html', {
            'form': form,
            'user': user,
            'token': token
        })
    
    except PasswordResetSession.DoesNotExist:
        messages.error(request, 'Invalid password reset session. Please start over.')
        return redirect('accounts:forgot_password_request')


@login_required
def security_questions_setup(request):
    """Set up security questions for password recovery"""
    user = request.user
    
    # Check if user already has security questions set up
    if user.security_answers.exists():
        messages.info(request, 'You already have security questions set up.')
        return redirect('accounts:profile_detail')
    
    if request.method == 'POST':
        form = SecurityQuestionsSetupForm(request.POST)
        if form.is_valid():
            # Save security answers
            questions = SecurityQuestion.objects.filter(is_active=True)[:3]
            
            with transaction.atomic():
                for question in questions:
                    field_name = f'question_{question.id}'
                    answer = form.cleaned_data.get(field_name)
                    
                    if answer:
                        UserSecurityAnswer.objects.create(
                            user=user,
                            security_question=question,
                            answer=answer
                        )
            
            messages.success(request, 'Security questions have been set up successfully. You can now use them for password recovery.')
            return redirect('accounts:profile_detail')
    else:
        form = SecurityQuestionsSetupForm()
    
    return render(request, 'accounts/security_questions_setup.html', {'form': form})


@login_required
def security_questions_manage(request):
    """Manage existing security questions"""
    user = request.user
    security_answers = user.security_answers.select_related('security_question').all()
    
    if request.method == 'POST':
        # Update security answers
        form = SecurityQuestionsSetupForm(request.POST)
        if form.is_valid():
            questions = SecurityQuestion.objects.filter(is_active=True)[:3]
            
            with transaction.atomic():
                # Delete existing answers
                user.security_answers.all().delete()
                
                # Create new answers
                for question in questions:
                    field_name = f'question_{question.id}'
                    answer = form.cleaned_data.get(field_name)
                    
                    if answer:
                        UserSecurityAnswer.objects.create(
                            user=user,
                            security_question=question,
                            answer=answer
                        )
            
            messages.success(request, 'Security questions have been updated successfully.')
            return redirect('accounts:security_questions_manage')
    else:
        # Pre-populate form with existing answers (for display only)
        form = SecurityQuestionsSetupForm()
    
    return render(request, 'accounts/security_questions_manage.html', {
        'form': form,
        'security_answers': security_answers
    })
