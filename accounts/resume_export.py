"""
Resume Export Engine for Next-Generation ATS Resume Builder
Supports PDF, DOCX, and Plain Text exports with ATS-safe formatting
"""

import os
import json
from io import BytesIO
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from django.conf import settings
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


class ResumeExportEngine:
    """
    Comprehensive resume export engine supporting multiple formats
    with ATS-optimized formatting and cloud storage integration
    """
    
    def __init__(self):
        self.supported_formats = ['pdf', 'docx', 'txt', 'html']
        self.ats_safe_fonts = ['Arial', 'Calibri', 'Times New Roman', 'Helvetica']
        
    def export_resume(self, resume, format_type: str, template_name: str = None) -> Tuple[bytes, str]:
        """
        Export resume in specified format with ATS optimization
        
        Args:
            resume: Resume model instance
            format_type: Export format ('pdf', 'docx', 'txt', 'html')
            template_name: Optional template override
            
        Returns:
            Tuple of (file_content, filename)
        """
        if format_type not in self.supported_formats:
            raise ValueError(f"Unsupported format: {format_type}")
            
        # Generate resume data
        resume_data = self._prepare_resume_data(resume)
        
        # Export based on format
        if format_type == 'pdf':
            return self._export_pdf(resume_data, resume.title)
        elif format_type == 'docx':
            return self._export_docx(resume_data, resume.title)
        elif format_type == 'txt':
            return self._export_txt(resume_data, resume.title)
        elif format_type == 'html':
            return self._export_html(resume_data, resume.title, template_name)
            
    def _prepare_resume_data(self, resume) -> Dict:
        """Prepare structured resume data for export"""
        
        # Get user profile data
        profile = resume.user.jobseekerprofile
        
        data = {
            'personal_info': {
                'name': f"{resume.user.first_name} {resume.user.last_name}",
                'email': resume.user.email,
                'phone': getattr(profile, 'phone', ''),
                'location': f"{getattr(profile, 'city', '')}, {getattr(profile, 'state', '')}",
                'linkedin': getattr(profile, 'linkedin_url', ''),
                'website': getattr(profile, 'website', ''),
            },
            'summary': resume.summary or '',
            'sections': []
        }
        
        # Add work experience
        experiences = resume.resumeworkexperience_set.all().order_by('-start_date')
        if experiences:
            exp_section = {
                'title': 'Professional Experience',
                'type': 'experience',
                'items': []
            }
            
            for exp in experiences:
                end_date = exp.end_date.strftime('%m/%Y') if exp.end_date else 'Present'
                exp_item = {
                    'title': exp.job_title,
                    'company': exp.company,
                    'location': exp.location or '',
                    'duration': f"{exp.start_date.strftime('%m/%Y')} - {end_date}",
                    'description': exp.description or '',
                    'achievements': exp.achievements.split('\n') if exp.achievements else []
                }
                exp_section['items'].append(exp_item)
            
            data['sections'].append(exp_section)
        
        # Add education
        educations = resume.resumeeducation_set.all().order_by('-graduation_year')
        if educations:
            edu_section = {
                'title': 'Education',
                'type': 'education',
                'items': []
            }
            
            for edu in educations:
                edu_item = {
                    'degree': edu.degree,
                    'institution': edu.institution,
                    'year': str(edu.graduation_year) if edu.graduation_year else '',
                    'gpa': f"GPA: {edu.gpa}" if edu.gpa else '',
                    'honors': edu.honors or ''
                }
                edu_section['items'].append(edu_item)
            
            data['sections'].append(edu_section)
        
        # Add skills
        skills = resume.resumeskill_set.all()
        if skills:
            skills_section = {
                'title': 'Skills',
                'type': 'skills',
                'items': [{'name': skill.skill, 'level': skill.proficiency} for skill in skills]
            }
            data['sections'].append(skills_section)
        
        return data
    
    def _export_pdf(self, resume_data: Dict, title: str) -> Tuple[bytes, str]:
        """Export resume as ATS-optimized PDF"""
        if not REPORTLAB_AVAILABLE:
            raise ImportError("ReportLab is required for PDF export")
            
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        
        # Create styles
        styles = getSampleStyleSheet()
        
        # Custom ATS-friendly styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            fontName='Helvetica-Bold',
            textColor=colors.black,
            alignment=1,  # Center
            spaceAfter=12
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=12,
            fontName='Helvetica-Bold',
            textColor=colors.black,
            spaceBefore=12,
            spaceAfter=6
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            fontName='Helvetica',
            textColor=colors.black,
            leftIndent=0,
            spaceAfter=6
        )
        
        # Build PDF content
        story = []
        
        # Header with personal info
        personal = resume_data['personal_info']
        story.append(Paragraph(personal['name'], title_style))
        
        contact_info = []
        if personal['email']:
            contact_info.append(personal['email'])
        if personal['phone']:
            contact_info.append(personal['phone'])
        if personal['location']:
            contact_info.append(personal['location'])
            
        if contact_info:
            story.append(Paragraph(' | '.join(contact_info), normal_style))
        
        if personal['linkedin'] or personal['website']:
            links = []
            if personal['linkedin']:
                links.append(personal['linkedin'])
            if personal['website']:
                links.append(personal['website'])
            story.append(Paragraph(' | '.join(links), normal_style))
        
        story.append(Spacer(1, 12))
        
        # Professional Summary
        if resume_data['summary']:
            story.append(Paragraph('Professional Summary', heading_style))
            story.append(Paragraph(resume_data['summary'], normal_style))
            story.append(Spacer(1, 12))
        
        # Sections
        for section in resume_data['sections']:
            story.append(Paragraph(section['title'], heading_style))
            
            if section['type'] == 'experience':
                for item in section['items']:
                    # Job title and company
                    job_line = f"<b>{item['title']}</b> | {item['company']}"
                    if item['location']:
                        job_line += f" | {item['location']}"
                    story.append(Paragraph(job_line, normal_style))
                    
                    # Duration
                    story.append(Paragraph(item['duration'], normal_style))
                    
                    # Description
                    if item['description']:
                        story.append(Paragraph(item['description'], normal_style))
                    
                    # Achievements
                    for achievement in item['achievements']:
                        if achievement.strip():
                            story.append(Paragraph(f"• {achievement.strip()}", normal_style))
                    
                    story.append(Spacer(1, 6))
            
            elif section['type'] == 'education':
                for item in section['items']:
                    edu_line = f"<b>{item['degree']}</b> | {item['institution']}"
                    if item['year']:
                        edu_line += f" | {item['year']}"
                    story.append(Paragraph(edu_line, normal_style))
                    
                    if item['gpa']:
                        story.append(Paragraph(item['gpa'], normal_style))
                    if item['honors']:
                        story.append(Paragraph(item['honors'], normal_style))
                    
                    story.append(Spacer(1, 6))
            
            elif section['type'] == 'skills':
                skills_text = ', '.join([item['name'] for item in section['items']])
                story.append(Paragraph(skills_text, normal_style))
            
            story.append(Spacer(1, 12))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        filename = f"{title.replace(' ', '_')}_resume.pdf"
        return buffer.getvalue(), filename
    
    def _export_docx(self, resume_data: Dict, title: str) -> Tuple[bytes, str]:
        """Export resume as ATS-optimized DOCX"""
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX export")
            
        doc = Document()
        
        # Set document margins for ATS compatibility
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Header with personal info
        personal = resume_data['personal_info']
        
        # Name (centered, bold, larger font)
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(personal['name'])
        name_run.bold = True
        name_run.font.size = Pt(16)
        name_run.font.name = 'Arial'
        
        # Contact info (centered)
        contact_info = []
        if personal['email']:
            contact_info.append(personal['email'])
        if personal['phone']:
            contact_info.append(personal['phone'])
        if personal['location']:
            contact_info.append(personal['location'])
            
        if contact_info:
            contact_para = doc.add_paragraph(' | '.join(contact_info))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Links
        if personal['linkedin'] or personal['website']:
            links = []
            if personal['linkedin']:
                links.append(personal['linkedin'])
            if personal['website']:
                links.append(personal['website'])
            links_para = doc.add_paragraph(' | '.join(links))
            links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        doc.add_paragraph()
        
        # Professional Summary
        if resume_data['summary']:
            summary_heading = doc.add_paragraph()
            summary_run = summary_heading.add_run('PROFESSIONAL SUMMARY')
            summary_run.bold = True
            summary_run.font.size = Pt(12)
            summary_run.font.name = 'Arial'
            
            doc.add_paragraph(resume_data['summary'])
        
        # Sections
        for section in resume_data['sections']:
            # Section heading
            section_heading = doc.add_paragraph()
            section_run = section_heading.add_run(section['title'].upper())
            section_run.bold = True
            section_run.font.size = Pt(12)
            section_run.font.name = 'Arial'
            
            if section['type'] == 'experience':
                for item in section['items']:
                    # Job title and company (bold)
                    job_para = doc.add_paragraph()
                    job_run = job_para.add_run(f"{item['title']} | {item['company']}")
                    job_run.bold = True
                    
                    if item['location']:
                        job_para.add_run(f" | {item['location']}")
                    
                    # Duration
                    doc.add_paragraph(item['duration'])
                    
                    # Description
                    if item['description']:
                        doc.add_paragraph(item['description'])
                    
                    # Achievements
                    for achievement in item['achievements']:
                        if achievement.strip():
                            doc.add_paragraph(f"• {achievement.strip()}")
            
            elif section['type'] == 'education':
                for item in section['items']:
                    edu_para = doc.add_paragraph()
                    edu_run = edu_para.add_run(f"{item['degree']} | {item['institution']}")
                    edu_run.bold = True
                    
                    if item['year']:
                        edu_para.add_run(f" | {item['year']}")
                    
                    if item['gpa']:
                        doc.add_paragraph(item['gpa'])
                    if item['honors']:
                        doc.add_paragraph(item['honors'])
            
            elif section['type'] == 'skills':
                skills_text = ', '.join([item['name'] for item in section['items']])
                doc.add_paragraph(skills_text)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        filename = f"{title.replace(' ', '_')}_resume.docx"
        return buffer.getvalue(), filename
    
    def _export_txt(self, resume_data: Dict, title: str) -> Tuple[bytes, str]:
        """Export resume as ATS-optimized plain text"""
        lines = []
        
        # Header
        personal = resume_data['personal_info']
        lines.append(personal['name'].upper())
        lines.append('=' * len(personal['name']))
        lines.append('')
        
        # Contact info
        contact_info = []
        if personal['email']:
            contact_info.append(f"Email: {personal['email']}")
        if personal['phone']:
            contact_info.append(f"Phone: {personal['phone']}")
        if personal['location']:
            contact_info.append(f"Location: {personal['location']}")
        
        lines.extend(contact_info)
        
        if personal['linkedin']:
            lines.append(f"LinkedIn: {personal['linkedin']}")
        if personal['website']:
            lines.append(f"Website: {personal['website']}")
        
        lines.append('')
        lines.append('')
        
        # Professional Summary
        if resume_data['summary']:
            lines.append('PROFESSIONAL SUMMARY')
            lines.append('-' * 20)
            lines.append(resume_data['summary'])
            lines.append('')
            lines.append('')
        
        # Sections
        for section in resume_data['sections']:
            lines.append(section['title'].upper())
            lines.append('-' * len(section['title']))
            lines.append('')
            
            if section['type'] == 'experience':
                for item in section['items']:
                    lines.append(f"{item['title']} | {item['company']}")
                    if item['location']:
                        lines.append(f"Location: {item['location']}")
                    lines.append(f"Duration: {item['duration']}")
                    lines.append('')
                    
                    if item['description']:
                        lines.append(item['description'])
                        lines.append('')
                    
                    for achievement in item['achievements']:
                        if achievement.strip():
                            lines.append(f"• {achievement.strip()}")
                    
                    lines.append('')
            
            elif section['type'] == 'education':
                for item in section['items']:
                    lines.append(f"{item['degree']} | {item['institution']}")
                    if item['year']:
                        lines.append(f"Year: {item['year']}")
                    if item['gpa']:
                        lines.append(item['gpa'])
                    if item['honors']:
                        lines.append(item['honors'])
                    lines.append('')
            
            elif section['type'] == 'skills':
                skills_text = ', '.join([item['name'] for item in section['items']])
                lines.append(skills_text)
            
            lines.append('')
        
        content = '\n'.join(lines).encode('utf-8')
        filename = f"{title.replace(' ', '_')}_resume.txt"
        return content, filename
    
    def _export_html(self, resume_data: Dict, title: str, template_name: str = None) -> Tuple[bytes, str]:
        """Export resume as HTML with ATS-friendly styling"""
        template_name = template_name or 'accounts/resume_templates/professional.html'
        
        html_content = render_to_string(template_name, {
            'resume_data': resume_data,
            'title': title
        })
        
        filename = f"{title.replace(' ', '_')}_resume.html"
        return html_content.encode('utf-8'), filename
    
    def save_to_cloud(self, file_content: bytes, filename: str, user_id: int) -> str:
        """
        Save exported resume to cloud storage
        
        Returns:
            Cloud storage URL
        """
        # Create user-specific path
        file_path = f"resumes/user_{user_id}/exports/{filename}"
        
        # Save to default storage (can be configured for AWS S3, etc.)
        file_obj = ContentFile(file_content)
        saved_path = default_storage.save(file_path, file_obj)
        
        # Return URL
        return default_storage.url(saved_path)
    
    def get_export_history(self, user_id: int) -> List[Dict]:
        """Get user's export history from cloud storage"""
        try:
            export_path = f"resumes/user_{user_id}/exports/"
            
            # List files in user's export directory
            dirs, files = default_storage.listdir(export_path)
            
            history = []
            for file in files:
                file_path = os.path.join(export_path, file)
                try:
                    # Get file info
                    file_info = {
                        'filename': file,
                        'url': default_storage.url(file_path),
                        'size': default_storage.size(file_path),
                        'modified': default_storage.get_modified_time(file_path),
                        'format': file.split('.')[-1].upper()
                    }
                    history.append(file_info)
                except:
                    continue
            
            # Sort by modification time (newest first)
            history.sort(key=lambda x: x['modified'], reverse=True)
            return history
            
        except Exception as e:
            return []


# Export utility functions
def create_export_response(file_content: bytes, filename: str, content_type: str) -> HttpResponse:
    """Create HTTP response for file download"""
    response = HttpResponse(file_content, content_type=content_type)
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def get_content_type(format_type: str) -> str:
    """Get MIME type for export format"""
    content_types = {
        'pdf': 'application/pdf',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'txt': 'text/plain',
        'html': 'text/html'
    }
    return content_types.get(format_type, 'application/octet-stream')
